"""Tests for MCP tool schema validation and strict parameter handling."""

import asyncio

from docx import Document

from office_server import OfficeServer
from tools.word_advanced_tools import _get_text_with_track_changes


def _get_tool_schema(tools, name):
    for tool in tools:
        if tool.get("name") == name:
            return tool.get("inputSchema", {})
    return {}


def test_tool_schema_disallows_unknown_properties():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])
    schema = _get_tool_schema(tools, "office_comment")
    assert schema.get("additionalProperties") is False


def test_rejects_unknown_parameters():
    server = OfficeServer()
    response = asyncio.run(
        server.handle_tools_call_async(
            request_id=1,
            params={
                "name": "office_comment",
                "arguments": {
                    "file_path": "dummy.pptx",
                    "comment": "invalid",
                },
            },
        )
    )

    error = response.get("error", {})
    assert "Unrecognized parameter" in error.get("message", "")


def test_markdown_tools_expose_markdown_file_and_oneof():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])

    for tool_name in [
        "word_from_markdown",
        "excel_from_markdown",
        "pptx_from_markdown",
        "word_create_sow_from_markdown",
    ]:
        schema = _get_tool_schema(tools, tool_name)
        props = schema.get("properties", {})

        assert "markdown" in props
        assert "markdown_file" in props

        # Agent-oriented contract: at least one markdown source must be provided
        assert "oneOf" in schema
        assert {"required": ["markdown"]} in schema["oneOf"]
        assert {"required": ["markdown_file"]} in schema["oneOf"]


def test_word_create_sow_schema_marks_template_required():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])
    schema = _get_tool_schema(tools, "word_create_sow_from_markdown")

    required = schema.get("required", [])
    assert "output_path" in required
    assert "template_path" in required


def test_office_comment_supports_reply_operation():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])
    schema = _get_tool_schema(tools, "office_comment")

    operation = schema.get("properties", {}).get("operation", {})
    enum_values = operation.get("enum", [])
    assert "reply" in enum_values


def test_word_insert_at_anchor_schema_present():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])
    schema = _get_tool_schema(tools, "word_insert_at_anchor")

    required = schema.get("required", [])
    props = schema.get("properties", {})
    assert "file_path" in required
    assert "content" in required
    assert "anchor_text" in props
    assert "paragraph_index" in props
    assert "position" in props


def test_office_help_schema_present():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])
    schema = _get_tool_schema(tools, "office_help")

    props = schema.get("properties", {})
    assert "goal" in props
    assert "document_type" in props
    assert "constraints" in props
    assert "task" in props
    assert "format" in props


def test_mutation_tools_expose_mode_parameter():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])

    for tool_name in ["office_patch", "office_table", "word_create_sow_from_markdown"]:
        schema = _get_tool_schema(tools, tool_name)
        assert "mode" in schema.get("properties", {})


def test_anchor_discovery_tools_schema_present():
    server = OfficeServer()
    tools = server.discover_tools().get("tools", [])

    for tool_name in ["word_list_anchors", "word_document_map"]:
        schema = _get_tool_schema(tools, tool_name)
        assert schema

    anchor_schema = _get_tool_schema(tools, "word_list_anchors")
    assert "query" in anchor_schema.get("properties", {})
    assert "include_paragraphs" in anchor_schema.get("properties", {})


def test_server_instructions_mention_core_first_discovery_guidance_and_word_insertion():
    server = OfficeServer()
    instructions = server.get_instructions()

    assert "core-first" in instructions
    assert "office_help" in instructions
    assert "office_read" in instructions
    assert "word_insert_at_anchor" in instructions
    assert "fallback" in instructions
    assert "section:" in instructions
    assert "word_audit_completion" in instructions


def test_word_insert_at_anchor_runtime_via_mcp_handler(temp_dir):
    server = OfficeServer()
    path = temp_dir / "anchor_runtime.docx"
    out = temp_dir / "anchor_runtime_out.docx"

    doc = Document()
    doc.add_paragraph("Intro")
    doc.add_paragraph("Anchor paragraph")
    doc.save(path)

    response = asyncio.run(
        server.handle_tools_call_async(
            request_id=1,
            params={
                "name": "word_insert_at_anchor",
                "arguments": {
                    "file_path": str(path),
                    "output_path": str(out),
                    "anchor_text": "Anchor paragraph",
                    "position": "after",
                    "content": ["Inserted via MCP handler"],
                },
            },
        )
    )

    assert "error" not in response
    text = "".join(block.get("text", "") for block in response["result"].get("content", []))
    assert "Inserted 1 paragraph" in text

    reloaded = Document(out)
    paragraphs = [_get_text_with_track_changes(p).strip() for p in reloaded.paragraphs if _get_text_with_track_changes(p).strip()]
    assert paragraphs == ["Intro", "Anchor paragraph", "Inserted via MCP handler"]
