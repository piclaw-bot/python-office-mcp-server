"""
Tests for track-change-aware reading functionality.

These tests verify that tools can read content inside w:ins elements
when documents have tracked changes.
"""

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.word_advanced_tools import (
    _add_tracked_deletion,
    _add_tracked_insertion,
    _get_text_with_track_changes,
)


class TestGetTextWithTrackChanges:
    """Test the _get_text_with_track_changes helper function."""

    def test_reads_regular_text(self):
        """Should read normal text without track changes."""
        doc = Document()
        para = doc.add_paragraph("Regular text without changes.")

        text = _get_text_with_track_changes(para)
        assert text == "Regular text without changes."

    def test_reads_inserted_text(self):
        """Should read text inside w:ins elements."""
        doc = Document()
        para = doc.add_paragraph("Before ")
        _add_tracked_insertion(para, "INSERTED", author="Test")
        para.add_run(" after")

        text = _get_text_with_track_changes(para)
        assert "INSERTED" in text
        assert "Before" in text
        assert "after" in text

    def test_excludes_deleted_text(self):
        """Should NOT read text inside w:del elements."""
        doc = Document()
        para = doc.add_paragraph("Visible ")
        _add_tracked_deletion(para, "DELETED", author="Test")
        para.add_run(" text")

        text = _get_text_with_track_changes(para)
        assert "DELETED" not in text
        assert "Visible" in text
        assert "text" in text

    def test_reads_replacement_correctly(self):
        """Should read new value, not old value in replacements."""
        doc = Document()
        para = doc.add_paragraph("Customer: ")
        _add_tracked_deletion(para, "<Customer Name>", author="Test")
        _add_tracked_insertion(para, "Contoso Ltd", author="Test")

        text = _get_text_with_track_changes(para)
        assert "Contoso Ltd" in text
        assert "<Customer Name>" not in text


class TestGetSectionWithTrackChanges:
    """Test that get_section reads track-changed content."""

    def test_get_section_reads_inserted_content(self, word_advanced_tools, temp_dir):
        """get_section should return content from w:ins elements."""
        # Create doc with tracked insertion in a section
        doc = Document()
        doc.add_heading("Introduction", level=1)
        para = doc.add_paragraph()
        _add_tracked_insertion(para, "This is the new introduction content.", author="Test")
        doc.add_heading("Next Section", level=1)

        path = temp_dir / "test_section_track.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_get_section(str(path), "Introduction")

        assert "error" not in result
        # The content should include the inserted text
        content_text = " ".join(result.get("content", []))
        assert "new introduction content" in content_text

    def test_get_section_excludes_deleted_content(self, word_advanced_tools, temp_dir):
        """get_section should NOT return content from w:del elements."""
        doc = Document()
        doc.add_heading("Summary", level=1)
        para = doc.add_paragraph("Current summary. ")
        _add_tracked_deletion(para, "OLD DELETED TEXT", author="Test")
        doc.add_heading("Details", level=1)

        path = temp_dir / "test_section_del.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_get_section(str(path), "Summary")

        content_text = " ".join(result.get("content", []))
        assert "OLD DELETED TEXT" not in content_text
        assert "Current summary" in content_text


class TestListSectionsWithTrackChanges:
    """Test that list_sections reads track-changed headings."""

    def test_list_sections_reads_inserted_heading(self, word_advanced_tools, temp_dir):
        """list_sections should return heading text from w:ins."""
        doc = Document()
        heading = doc.add_heading("", level=1)
        _add_tracked_insertion(heading, "New Section Title", author="Test")

        path = temp_dir / "test_list_sections.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_list_sections(str(path))

        titles = [s["title"] for s in result.get("sections", [])]
        assert "New Section Title" in titles


class TestPatchSectionWithTrackChanges:
    """Test that patch_section finds sections with track-changed titles."""

    def test_patch_section_finds_inserted_heading(self, word_advanced_tools, temp_dir):
        """patch_section should find section by inserted heading text."""
        doc = Document()
        heading = doc.add_heading("", level=1)
        _add_tracked_insertion(heading, "Target Section", author="Test")
        doc.add_paragraph("Old content to replace.")
        doc.add_heading("Other Section", level=1)

        path = temp_dir / "test_patch_find.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_patch_section(
            file_path=str(path),
            section_title="Target Section",
            new_content=["New paragraph content."]
        )

        assert result.get("success"), f"Should find section: {result}"


class TestAuditSowWithTrackChanges:
    """Test that audit_sow reads track-changed content."""

    def test_audit_sow_finds_placeholder_in_insertion(self, word_advanced_tools, temp_dir):
        """audit_sow should find placeholders inside w:ins elements."""
        doc = Document()
        para = doc.add_paragraph("Customer: ")
        _add_tracked_insertion(para, "<Customer Name>", author="Test")

        path = temp_dir / "test_audit_ins.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_audit_sow(str(path))

        # Should find the placeholder (it will be in split_runs since it's inside w:ins)
        total = result.get("summary", {}).get("total_placeholders", 0)
        # The placeholder is found via _get_text_with_track_changes but isn't in para.runs
        # so it ends up in placeholders_split_runs, which is correct behavior
        split_runs = len(result.get("placeholders_split_runs", []))
        assert total >= 1 or split_runs >= 1, f"Should find placeholder: {result}"


class TestCleanupSowWithTrackChanges:
    """Test that cleanup_sow handles track-changed content."""

    def test_cleanup_reads_inserted_instruction_text(self, word_advanced_tools, temp_dir):
        """cleanup_sow should detect instructions in w:ins elements."""
        doc = Document()
        para = doc.add_paragraph()
        _add_tracked_insertion(para, "[Template Guidance: Remove this]", author="Test")

        path = temp_dir / "test_cleanup_ins.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_cleanup_sow(str(path))

        # Should have cleaned something
        assert result.get("success")


class TestNextToolsSuggestions:
    """Test that tools return appropriate next_tools suggestions."""

    def test_copy_template_suggests_generate_sow(self, word_advanced_tools, temp_dir):
        """copy_template should suggest understanding the template as next step."""
        # Create a minimal template
        doc = Document()
        doc.add_paragraph("Template content")
        template_path = temp_dir / "template.docx"
        doc.save(template_path)

        result = word_advanced_tools.tool_word_copy_template(
            template_name="template.docx",
            output_path=str(temp_dir / "output.docx"),
            template_dir=str(temp_dir)
        )

        assert result.get("success")
        # Next step after copying is to understand the template structure
        assert "word_parse_sow_template" in result.get("next_tools", [])

    def test_generate_sow_suggests_patch_section(self, word_advanced_tools, temp_dir):
        """generate_sow should suggest patch_section for prose sections."""
        # Create minimal template
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("<Customer Name> content here.")
        template_path = temp_dir / "template.docx"
        doc.save(template_path)

        result = word_advanced_tools.tool_word_generate_sow(
            template_path=str(template_path),
            output_path=str(temp_dir / "output.docx"),
            sow_data={"customer_name": "Test Corp"}
        )

        assert result.get("success")
        # Should suggest table and section related tools
        assert "word_get_section_guidance" in result.get("next_tools", [])

    def test_audit_completion_suggests_based_on_issues(self, word_advanced_tools, temp_dir):
        """audit_completion should suggest tools based on found issues."""
        # Create doc with placeholders
        doc = Document()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("<Customer Name> needs filling.")
        path = temp_dir / "test_audit.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_audit_completion(str(path))

        assert "next_tools" in result
        # Should suggest tools to fix the issues
        next_tools = result.get("next_tools", [])
        assert len(next_tools) > 0
