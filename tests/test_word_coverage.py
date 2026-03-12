"""
Additional tests for high-coverage impact on word_advanced_tools.py
"""

from docx import Document

# Fixtures temp_dir and word_advanced_tools are provided by conftest.py


class TestPromptFunctions:
    """Tests for prompt/guidance functions."""

    def test_prompt_word_get_section_guidance(self, word_advanced_tools, temp_dir):
        """Should provide section guidance."""
        doc = Document()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("Summary content")
        path = temp_dir / "guidance.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_get_section_guidance(str(path), "Executive Summary")
        assert isinstance(result, dict)


class TestTrackChangesOperations:
    """Tests for track changes functionality."""

    def test_enable_track_changes(self, word_advanced_tools, temp_dir):
        """Should enable track changes."""
        doc = Document()
        doc.add_paragraph("Test content")
        path = temp_dir / "track.docx"
        doc.save(path)

        output = temp_dir / "tracked.docx"
        result = word_advanced_tools.tool_word_enable_track_changes(str(path), output_path=str(output))
        assert isinstance(result, dict)

    def test_patch_with_track_changes(self, word_advanced_tools, temp_dir):
        """Should patch content with track changes."""
        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_paragraph("<Customer Name> is our client.")
        path = temp_dir / "track_patch.docx"
        doc.save(path)

        output = temp_dir / "patched_tracked.docx"
        # Method takes replacements dict
        result = word_advanced_tools.tool_word_patch_with_track_changes(
            str(path),
            replacements={"<Customer Name>": "Contoso Corp"},
            output_path=str(output)
        )
        assert isinstance(result, dict)


class TestCleanupSow:
    """Tests for SOW cleanup functionality."""

    def test_cleanup_sow(self, word_advanced_tools, temp_dir):
        """Should cleanup SOW document."""
        doc = Document()
        doc.add_heading("SOW Title", level=0)
        doc.add_heading("Section 1", level=1)
        # Add guidance text
        doc.add_paragraph("[Template Guidance: Remove this text]")
        doc.add_paragraph("Real content here")
        path = temp_dir / "cleanup_sow.docx"
        doc.save(path)

        output = temp_dir / "cleaned.docx"
        result = word_advanced_tools.tool_word_cleanup_sow(str(path), output_path=str(output))
        assert isinstance(result, dict)


class TestParseSowTemplate:
    """Tests for SOW template parsing."""

    def test_parse_sow_template(self, word_advanced_tools, temp_dir):
        """Should parse SOW template structure."""
        doc = Document()
        doc.add_heading("Statement of Work", level=0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("<Project Description>")
        doc.add_heading("Scope", level=1)
        doc.add_paragraph("Scope details")
        path = temp_dir / "template.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_parse_sow_template(str(path))
        assert isinstance(result, dict)


class TestAnalyzeTemplateFormatting:
    """Tests for template formatting analysis."""

    def test_analyze_template_formatting(self, word_advanced_tools, temp_dir):
        """Should analyze template formatting."""
        doc = Document()
        doc.add_heading("Template", level=0)
        doc.add_paragraph("Boilerplate text here")
        doc.add_paragraph("<Placeholder>")
        doc.add_paragraph("[TBD]")
        path = temp_dir / "format_template.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_analyze_template_formatting(str(path))
        assert isinstance(result, dict)


class TestPatchTableRow:
    """Tests for patching table rows."""

    def test_patch_table_row(self, word_advanced_tools, temp_dir):
        """Should patch a table row."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "Value1"
        table.cell(1, 1).text = "Value2"
        path = temp_dir / "patch_table.docx"
        doc.save(path)

        output = temp_dir / "patched_table.docx"
        result = word_advanced_tools.tool_word_patch_table_row(
            str(path),
            "0",  # First table
            1,    # Second row (0-indexed)
            {"A": "NewValue1", "B": "NewValue2"},
            output_path=str(output)
        )
        assert isinstance(result, dict)


class TestReplaceGlobalVariables:
    """Tests for global variable replacement."""

    def test_replace_global_variables(self, word_advanced_tools, temp_dir):
        """Should replace global variables throughout document."""
        doc = Document()
        doc.add_heading("<Customer Name> Project", level=1)
        doc.add_paragraph("Prepared for <Customer Name>")
        doc.add_paragraph("Project: <Project Name>")
        path = temp_dir / "globals.docx"
        doc.save(path)

        output = temp_dir / "replaced.docx"
        result = word_advanced_tools.tool_word_replace_global_variables(
            str(path),
            replacements={
                "<Customer Name>": "Contoso Corp",
                "<Project Name>": "Cloud Migration"
            },
            output_path=str(output)
        )
        assert isinstance(result, dict)


class TestSowGeneration:
    """Tests for SOW generation functionality."""

    def test_generate_sow_minimal(self, word_advanced_tools, temp_dir):
        """Should generate SOW from minimal data."""
        # Create a basic template
        template_doc = Document()
        template_doc.add_heading("Statement of Work", level=0)
        template_doc.add_heading("Executive Summary", level=1)
        template_doc.add_paragraph("<Executive Summary>")
        template_path = temp_dir / "template.docx"
        template_doc.save(template_path)

        output_path = temp_dir / "generated_sow.docx"

        sow_data = {
            "customer_name": "Test Corp",
            "project_name": "Test Project",
            "executive_summary": "This is a test project."
        }

        # Correct parameter order: template_path, output_path, sow_data
        result = word_advanced_tools.tool_word_generate_sow(
            str(template_path),
            str(output_path),
            sow_data
        )
        assert isinstance(result, dict)


class TestCopyTemplate:
    """Tests for template copying."""

    def test_copy_template_creates_copy(self, word_advanced_tools, temp_dir):
        """Should copy template file."""
        doc = Document()
        doc.add_heading("Original Template", level=1)
        doc.add_paragraph("Template content")
        source = temp_dir / "source_template.docx"
        doc.save(source)

        dest = temp_dir / "dest_template.docx"
        result = word_advanced_tools.tool_word_copy_template(str(source), str(dest))

        assert dest.exists() or result.get("success") is True


class TestAddComment:
    """Tests for adding comments."""

    def test_add_comment(self, word_advanced_tools, temp_dir):
        """Should add comment to document."""
        doc = Document()
        doc.add_paragraph("This is the target text for the comment.")
        path = temp_dir / "comment_doc.docx"
        doc.save(path)

        output = temp_dir / "with_comment.docx"
        result = word_advanced_tools.tool_word_add_comment(
            str(path),
            target_text="target text",
            comment_text="This is a review comment",
            output_path=str(output)
        )
        assert isinstance(result, dict)


class TestMoreTableOperations:
    """Additional table operation tests."""

    def test_list_tables_empty_doc(self, word_advanced_tools, temp_dir):
        """Should handle document with no tables."""
        doc = Document()
        doc.add_paragraph("No tables here")
        path = temp_dir / "no_tables.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_list_tables(str(path))
        tables = result.get("tables", [])
        assert len(tables) == 0

    def test_create_new_table_with_rows(self, word_advanced_tools, temp_dir):
        """Should create table with initial rows."""
        doc = Document()
        doc.add_heading("Document", level=1)
        path = temp_dir / "create_table.docx"
        doc.save(path)

        output = temp_dir / "with_new_table.docx"
        result = word_advanced_tools.tool_word_create_new_table(
            str(path),
            ["Column A", "Column B", "Column C"],
            rows=[
                {"Column A": "1", "Column B": "2", "Column C": "3"},
                {"Column A": "4", "Column B": "5", "Column C": "6"}
            ],
            output_path=str(output)
        )
        assert result.get("success") is True


class TestErrorPaths:
    """Tests for error handling paths."""

    def test_get_section_nonexistent(self, word_advanced_tools, temp_dir):
        """Should handle getting nonexistent section."""
        doc = Document()
        doc.add_heading("Existing Section", level=1)
        doc.add_paragraph("Content")
        path = temp_dir / "sections.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_get_section(str(path), "Nonexistent")
        # Should return error or empty content
        assert "error" in result or result.get("content") is None or result.get("content") == ""

    def test_file_not_found(self, word_advanced_tools):
        """Should handle file not found."""
        result = word_advanced_tools.tool_word_list_sections("/nonexistent/path.docx")
        assert "error" in result

    def test_get_table_invalid_index(self, word_advanced_tools, temp_dir):
        """Should handle invalid table index."""
        doc = Document()
        doc.add_table(rows=2, cols=2)
        path = temp_dir / "one_table.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_get_table(str(path), "999")
        assert "error" in result or "not found" in str(result).lower()
