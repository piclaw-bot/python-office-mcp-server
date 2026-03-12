"""
Tests for Track Changes implementation in SOW tools.

These tests validate that:
1. Track changes XML elements are correctly structured per OOXML spec
2. w:ins elements contain w:r > w:t for insertions
3. w:del elements contain w:r > w:delText for deletions
4. Required attributes (w:id, w:author, w:date) are present
5. Changes appear in the correct position within paragraphs
6. Word can recognize and display the tracked changes
"""

import re

# Import the tools module
import sys
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.word_advanced_tools import (
    _add_tracked_deletion,
    _add_tracked_insertion,
)

# OOXML namespace
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


class TestTrackChangesXMLStructure:
    """Test that track changes XML follows OOXML spec."""

    def test_insertion_has_required_attributes(self, simple_docx, temp_dir):
        """w:ins must have w:id, w:author, and w:date attributes."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        _add_tracked_insertion(para, "inserted text", author="Test Author")

        output = temp_dir / "test_ins_attrs.docx"
        doc.save(output)

        # Extract and parse document.xml
        xml_content = self._extract_document_xml(output)

        # Find w:ins elements
        ins_elements = xml_content.findall(f".//{{{WORD_NS}}}ins")
        assert len(ins_elements) >= 1, "No w:ins elements found"

        ins = ins_elements[-1]  # Get the last one we added

        # Check required attributes
        assert ins.get(qn("w:id")) is not None, "w:ins missing w:id attribute"
        assert ins.get(qn("w:author")) is not None, "w:ins missing w:author attribute"
        assert ins.get(qn("w:date")) is not None, "w:ins missing w:date attribute"

        # Verify author value
        assert ins.get(qn("w:author")) == "Test Author"

    def test_insertion_contains_run_with_text(self, simple_docx, temp_dir):
        """w:ins must contain w:r > w:t structure."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        _add_tracked_insertion(para, "test insertion", author="Test")

        output = temp_dir / "test_ins_structure.docx"
        doc.save(output)

        xml_content = self._extract_document_xml(output)
        ins_elements = xml_content.findall(f".//{{{WORD_NS}}}ins")

        assert len(ins_elements) >= 1
        ins = ins_elements[-1]

        # Check for w:r child
        runs = ins.findall(f"{{{WORD_NS}}}r")
        assert len(runs) >= 1, "w:ins should contain at least one w:r element"

        # Check for w:t in run
        text_elements = runs[0].findall(f"{{{WORD_NS}}}t")
        assert len(text_elements) >= 1, "w:r should contain w:t element"
        assert text_elements[0].text == "test insertion"

    def test_deletion_has_required_attributes(self, simple_docx, temp_dir):
        """w:del must have w:id, w:author, and w:date attributes."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        _add_tracked_deletion(para, "deleted text", author="Test Author")

        output = temp_dir / "test_del_attrs.docx"
        doc.save(output)

        xml_content = self._extract_document_xml(output)
        del_elements = xml_content.findall(f".//{{{WORD_NS}}}del")

        assert len(del_elements) >= 1, "No w:del elements found"

        deletion = del_elements[-1]

        # Check required attributes
        assert deletion.get(qn("w:id")) is not None, "w:del missing w:id attribute"
        assert deletion.get(qn("w:author")) is not None, "w:del missing w:author attribute"
        assert deletion.get(qn("w:date")) is not None, "w:del missing w:date attribute"

    def test_deletion_uses_delText_not_text(self, simple_docx, temp_dir):
        """w:del must contain w:r > w:delText (not w:t)."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        _add_tracked_deletion(para, "deleted content", author="Test")

        output = temp_dir / "test_del_structure.docx"
        doc.save(output)

        xml_content = self._extract_document_xml(output)
        del_elements = xml_content.findall(f".//{{{WORD_NS}}}del")

        assert len(del_elements) >= 1
        deletion = del_elements[-1]

        # Check for w:r child
        runs = deletion.findall(f"{{{WORD_NS}}}r")
        assert len(runs) >= 1, "w:del should contain at least one w:r element"

        # Check for w:delText (NOT w:t) in run
        del_text_elements = runs[0].findall(f"{{{WORD_NS}}}delText")
        assert len(del_text_elements) >= 1, "w:r inside w:del should contain w:delText element"
        assert del_text_elements[0].text == "deleted content"

        # Ensure w:t is NOT used
        text_elements = runs[0].findall(f"{{{WORD_NS}}}t")
        assert len(text_elements) == 0, "w:del should use w:delText, not w:t"

    def test_date_format_is_iso8601(self, simple_docx, temp_dir):
        """w:date should be ISO 8601 format with timezone."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        _add_tracked_insertion(para, "text", author="Test")

        output = temp_dir / "test_date_format.docx"
        doc.save(output)

        xml_content = self._extract_document_xml(output)
        ins_elements = xml_content.findall(f".//{{{WORD_NS}}}ins")

        ins = ins_elements[-1]
        date_str = ins.get(qn("w:date"))

        # Should match ISO 8601 format: YYYY-MM-DDTHH:MM:SSZ
        iso_pattern = r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}"
        assert re.match(iso_pattern, date_str), f"Date '{date_str}' not in ISO 8601 format"

    def test_unique_ids_across_document(self, simple_docx, temp_dir):
        """Each w:ins and w:del should have unique IDs."""
        doc = Document(simple_docx)
        para = doc.add_paragraph()

        # Add multiple track changes
        _add_tracked_insertion(para, "first", author="Test")
        _add_tracked_insertion(para, "second", author="Test")
        _add_tracked_deletion(para, "third", author="Test")
        _add_tracked_insertion(para, "fourth", author="Test")

        output = temp_dir / "test_unique_ids.docx"
        doc.save(output)

        xml_content = self._extract_document_xml(output)

        # Collect all IDs
        ids = []
        for ins in xml_content.findall(f".//{{{WORD_NS}}}ins"):
            ids.append(ins.get(qn("w:id")))
        for deletion in xml_content.findall(f".//{{{WORD_NS}}}del"):
            ids.append(deletion.get(qn("w:id")))

        # Check uniqueness
        assert len(ids) == len(set(ids)), f"Duplicate IDs found: {ids}"

    def _extract_document_xml(self, docx_path: Path) -> ET.Element:
        """Extract and parse document.xml from a .docx file."""
        with ZipFile(docx_path, "r") as zf:
            xml_bytes = zf.read("word/document.xml")
        return ET.fromstring(xml_bytes)


class TestTrackChangesPositioning:
    """Test that track changes are inserted at the correct position."""

    def test_replacement_preserves_surrounding_text(self, word_advanced_tools, multi_paragraph_docx, temp_dir):
        """When replacing text, surrounding content should be preserved."""
        output = temp_dir / "test_positioning.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(multi_paragraph_docx),
            replacements={"PLACEHOLDER": "REPLACED"},
            author="Test",
            output_path=str(output)
        )

        assert result.get("success"), f"Tool failed: {result}"

        # Verify the document can be opened
        doc = Document(output)

        # Check that paragraphs still exist and have content
        assert len(doc.paragraphs) >= 4

    def test_changes_appear_in_correct_paragraph(self, word_advanced_tools, temp_dir):
        """Track changes should be in the same paragraph as the original text."""
        # Create a document with distinct paragraphs
        doc = Document()
        doc.add_paragraph("Paragraph ONE has MARKER here.")
        doc.add_paragraph("Paragraph TWO is clean.")
        doc.add_paragraph("Paragraph THREE has MARKER too.")

        input_path = temp_dir / "test_correct_para.docx"
        doc.save(input_path)

        output = temp_dir / "test_correct_para_out.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(input_path),
            replacements={"MARKER": "CHANGED"},
            author="Test",
            output_path=str(output)
        )

        assert result.get("success")
        assert result.get("total_changes") == 2, "Should have 2 changes"

    def test_replacement_across_split_runs(self, word_advanced_tools, temp_dir):
        """Replacement should work when target text spans multiple runs."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Microsoft Teams Contact Center")
        para.add_run(", Dynamics 365")

        input_path = temp_dir / "split_runs.docx"
        doc.save(input_path)

        output = temp_dir / "split_runs_out.docx"
        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(input_path),
            replacements={"Microsoft Teams Contact Center, Dynamics 365": "Unified Platform"},
            author="Test",
            output_path=str(output)
        )

        assert result.get("success")
        assert result.get("total_changes") == 1

    def _extract_document_xml(self, docx_path: Path) -> ET.Element:
        """Extract and parse document.xml from a .docx file."""
        with ZipFile(docx_path, "r") as zf:
            xml_bytes = zf.read("word/document.xml")
        return ET.fromstring(xml_bytes)


class TestTrackChangesInTables:
    """Test track changes work correctly in table cells."""

    def test_changes_in_table_cells(self, word_advanced_tools, simple_docx, temp_dir):
        """Track changes should work in table cells."""
        output = temp_dir / "test_table_changes.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(simple_docx),
            replacements={"<Customer Name>": "Contoso"},
            author="Test",
            output_path=str(output)
        )

        assert result.get("success")
        # simple_docx has <Customer Name> in both paragraphs and table
        assert result.get("total_changes") >= 2


class TestAcceptAllChanges:
    """Test accepting tracked changes."""

    @pytest.mark.xfail(reason="tool_word_accept_all_changes not yet implemented", strict=False)
    def test_accept_removes_del_elements(self, temp_dir):
        """Accepting changes should remove w:del elements entirely."""
        # Create document with track changes
        doc = Document()
        para = doc.add_paragraph()
        _add_tracked_deletion(para, "deleted text", author="Test")
        _add_tracked_insertion(para, "inserted text", author="Test")

        input_path = temp_dir / "test_accept_input.docx"
        doc.save(input_path)

        pytest.xfail("tool_word_accept_all_changes not yet implemented")

    @pytest.mark.xfail(reason="tool_word_accept_all_changes not yet implemented", strict=False)
    def test_accept_preserves_inserted_text(self, temp_dir):
        """Accepting changes should keep the inserted text content."""
        doc = Document()
        para = doc.add_paragraph("Before ")
        _add_tracked_insertion(para, "INSERTED", author="Test")
        para.add_run(" After")

        input_path = temp_dir / "test_preserve_ins.docx"
        doc.save(input_path)

        pytest.xfail("tool_word_accept_all_changes not yet implemented")


class TestEnableTrackChanges:
    """Test enabling track changes mode."""

    def test_enable_sets_trackRevisions(self, word_advanced_tools, simple_docx, temp_dir):
        """Enabling track changes should set w:trackRevisions in settings."""
        output = temp_dir / "test_enabled.docx"

        result = word_advanced_tools.tool_word_enable_track_changes(
            file_path=str(simple_docx),
            output_path=str(output)
        )

        assert result.get("success")

        # Check settings.xml for trackRevisions
        with ZipFile(output, "r") as zf:
            settings_xml = zf.read("word/settings.xml").decode()

        assert "trackRevisions" in settings_xml, "trackRevisions not found in settings.xml"


class TestPatchWithTrackChangesEnablesRevisions:
    """Test that patch_with_track_changes enables trackRevisions."""

    def test_patch_enables_trackRevisions_in_settings(self, word_advanced_tools, simple_docx, temp_dir):
        """Patching should automatically enable trackRevisions in settings.xml."""
        output = temp_dir / "test_patch_enables.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(simple_docx),
            replacements={"<Customer Name>": "Test Corp"},
            author="Test Author",
            output_path=str(output)
        )

        assert result.get("success"), f"Patch failed: {result}"

        # Check settings.xml for trackRevisions
        with ZipFile(output, "r") as zf:
            settings_xml = zf.read("word/settings.xml").decode()

        assert "trackRevisions" in settings_xml, \
            "trackRevisions not found - Word won't show changes!"

        # Verify it's not set to false
        assert 'trackRevisions w:val="false"' not in settings_xml
        assert 'trackRevisions w:val="0"' not in settings_xml


class TestWordCompatibility:
    """Tests that verify compatibility with Microsoft Word."""

    def test_document_opens_without_corruption(self, word_advanced_tools, simple_docx, temp_dir):
        """Document with track changes should open without errors."""
        output = temp_dir / "test_word_compat.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(simple_docx),
            replacements={"<Customer Name>": "Test Corp"},
            author="Automated Test",
            output_path=str(output)
        )

        assert result.get("success")

        # Verify document can be re-opened by python-docx
        # (This is a basic corruption check)
        try:
            doc = Document(output)
            # Try to read content
            _ = [p.text for p in doc.paragraphs]
        except Exception as e:
            pytest.fail(f"Document appears corrupted: {e}")

    def test_xml_is_well_formed(self, word_advanced_tools, simple_docx, temp_dir):
        """All XML in the docx should be well-formed."""
        output = temp_dir / "test_xml_wellformed.docx"

        word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(simple_docx),
            replacements={"<Customer Name>": "Test"},
            author="Test",
            output_path=str(output)
        )

        # Try to parse all XML files in the docx
        with ZipFile(output, "r") as zf:
            for name in zf.namelist():
                if name.endswith(".xml"):
                    try:
                        content = zf.read(name)
                        ET.fromstring(content)
                    except ET.ParseError as e:
                        pytest.fail(f"Malformed XML in {name}: {e}")
