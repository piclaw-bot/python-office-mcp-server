"""Tests for best_effort/safe/strict/dry_run mutation modes."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.table import Table, TableStyleInfo

from tools.office_unified_tools import OfficeUnifiedTools
from tools.word_advanced_tools import WordAdvancedTools
from tools.excel_advanced_tools import ExcelAdvancedTools


class UnifiedTools(OfficeUnifiedTools, WordAdvancedTools, ExcelAdvancedTools):
    pass


def _build_word_doc(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Current intro")
    doc.save(path)
    return path


def _build_word_template(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("<Customer Name>")
    doc.add_heading("Delivery approach", level=1)
    doc.add_paragraph("[Template Guidance: add delivery approach]")
    doc.save(path)
    return path


def _build_excel_table(path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Role", "Count"])
    ws.append(["Architect", 1])
    table = Table(displayName="Staffing", ref="A1:B2")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    ws.add_table(table)
    wb.save(path)
    return path


class TestMutationModes:
    def test_office_patch_dry_run_does_not_modify_word_file(self, temp_dir):
        path = _build_word_doc(temp_dir / "doc.docx")
        before = path.read_bytes()

        result = UnifiedTools().tool_office_patch(
            file_path=str(path),
            changes=[{"target": "section:Introduction", "value": "New intro"}],
            mode="dry_run",
        )

        assert result["mode"] == "dry_run"
        assert result["success"] is True
        assert path.read_bytes() == before

    def test_office_patch_safe_requires_distinct_output_path(self, temp_dir):
        path = _build_word_doc(temp_dir / "doc.docx")

        result = UnifiedTools().tool_office_patch(
            file_path=str(path),
            changes=[{"target": "section:Introduction", "value": "New intro"}],
            mode="safe",
        )

        assert result["success"] is False
        assert result["mode"] == "safe"
        assert result["status"] == "failed"

    def test_word_create_sow_strict_rejects_unmapped_sections_without_writing(self, temp_dir):
        template = _build_word_template(temp_dir / "template.docx")
        output = temp_dir / "strict-output.docx"
        markdown = """# Sample SOW

Customer: Contoso
Project: Platform Review
Provider: Microsoft

## Introduction
Architecture overview text.

## Assumptions
Customer will provide access.
"""

        result = WordAdvancedTools().tool_word_create_sow_from_markdown(
            output_path=str(output),
            template_path=str(template),
            markdown=markdown,
            mode="strict",
        )

        assert result["success"] is False
        assert result["mode"] == "strict"
        assert result["status"] == "failed"
        assert not output.exists()

    def test_office_table_excel_dry_run_does_not_write(self, temp_dir):
        path = _build_excel_table(temp_dir / "table.xlsx")
        before = path.read_bytes()

        result = UnifiedTools().tool_office_table(
            file_path=str(path),
            operation="add_row",
            table_id="Staffing",
            data={"Role": "Engineer", "Count": 2},
            mode="dry_run",
        )

        assert result["success"] is True
        assert result["mode"] == "dry_run"
        assert path.read_bytes() == before

    def test_office_table_excel_safe_requires_output_path(self, temp_dir):
        path = _build_excel_table(temp_dir / "table.xlsx")

        result = UnifiedTools().tool_office_table(
            file_path=str(path),
            operation="add_row",
            table_id="Staffing",
            data={"Role": "Engineer", "Count": 2},
            mode="safe",
        )

        assert result["success"] is False
        assert result["mode"] == "safe"

    def test_best_effort_preserves_existing_successful_path(self, temp_dir):
        path = _build_excel_table(temp_dir / "table.xlsx")
        output = temp_dir / "table-out.xlsx"

        result = UnifiedTools().tool_office_table(
            file_path=str(path),
            operation="add_row",
            table_id="Staffing",
            data={"Role": "Engineer", "Count": 2},
            output_path=str(output),
            mode="best_effort",
        )

        assert result["success"] is True
        assert result["mode"] == "best_effort"
        assert output.exists()
        wb = load_workbook(output)
        try:
            ws = wb.active
            assert ws["A3"].value == "Engineer"
        finally:
            wb.close()
