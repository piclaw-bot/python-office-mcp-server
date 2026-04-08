"""Tests for persistent Office template metadata caching."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from office_server import OfficeServer
from tools.metadata_cache import (
    CACHE_SCHEMA_VERSION,
    load_cached_metadata,
    metadata_cache_key,
    metadata_cache_path,
    store_cached_metadata,
)
from tools.word_advanced_tools import WordAdvancedTools


def _build_template(path: Path, extra_paragraph: str | None = None) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("<Customer Name> overview")
    doc.add_heading("Delivery approach", level=1)
    doc.add_paragraph("[Template Guidance: add implementation approach]")
    if extra_paragraph:
        doc.add_paragraph(extra_paragraph)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Count"
    table.rows[1].cells[0].text = "Architect"
    table.rows[1].cells[1].text = "1"
    doc.save(path)
    return path


class TestMetadataCacheHelpers:
    def test_cache_key_stable_for_same_file_identity(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = temp_dir / "template.docx"
        _build_template(source)

        key1 = metadata_cache_key(source, "word", "template_metadata")
        key2 = metadata_cache_key(source, "word", "template_metadata")

        assert key1 == key2

    def test_cache_read_write_roundtrip(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = temp_dir / "template.docx"
        _build_template(source)

        metadata = {"sections": [{"title": "Intro"}], "warnings": []}
        store_cached_metadata(source, "word", "template_metadata", metadata)
        cached, info = load_cached_metadata(source, "word", "template_metadata")

        assert info["hit"] is True
        assert cached == metadata

    def test_cache_invalidation_on_file_change(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = temp_dir / "template.docx"
        _build_template(source)

        metadata = {"sections": [{"title": "Intro"}], "warnings": []}
        store_cached_metadata(source, "word", "template_metadata", metadata)

        _build_template(source, extra_paragraph="New content to change the file size")
        cached, info = load_cached_metadata(source, "word", "template_metadata")

        assert cached is None
        assert info["reason"] == "stale"

    def test_corrupt_cache_falls_back_cleanly(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = temp_dir / "template.docx"
        _build_template(source)

        cache_path = metadata_cache_path(source, "word", "template_metadata")
        cache_path.write_text("{not valid json", encoding="utf-8")

        cached, info = load_cached_metadata(source, "word", "template_metadata")

        assert cached is None
        assert info["reason"] == "corrupt"

    def test_schema_mismatch_falls_back_cleanly(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = temp_dir / "template.docx"
        _build_template(source)

        cache_path = metadata_cache_path(source, "word", "template_metadata")
        cache_path.write_text(
            json.dumps(
                {
                    "schemaVersion": CACHE_SCHEMA_VERSION + 1,
                    "documentType": "word",
                    "analysisType": "template_metadata",
                    "source": {
                        "path": str(source.resolve()),
                        "mtime_ns": source.stat().st_mtime_ns,
                        "size": source.stat().st_size,
                    },
                    "generatedAt": "2026-04-08T00:00:00+00:00",
                    "metadata": {"sections": []},
                }
            ),
            encoding="utf-8",
        )

        cached, info = load_cached_metadata(source, "word", "template_metadata")

        assert cached is None
        assert info["reason"] == "schema_mismatch"


class TestMetadataCacheIntegration:
    def test_word_parse_sow_template_uses_cache_on_second_call(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = _build_template(temp_dir / "template.docx")
        tool = WordAdvancedTools()

        first = tool.tool_word_parse_sow_template(str(source))
        second = tool.tool_word_parse_sow_template(str(source))

        assert first["cache"]["reason"] == "stored"
        assert second["cache"]["hit"] is True
        assert second["cache"]["reason"] == "hit"
        assert second["section_count"] == first["section_count"]
        assert second["anchors"]
        assert second["warnings"]

    def test_word_parse_sow_template_regenerates_after_template_change(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = _build_template(temp_dir / "template.docx")
        tool = WordAdvancedTools()

        first = tool.tool_word_parse_sow_template(str(source))
        _build_template(source, extra_paragraph="Architecture decisions")
        second = tool.tool_word_parse_sow_template(str(source))

        assert first["cache"]["reason"] == "stored"
        assert second["cache"]["reason"] == "stored"
        assert any(anchor["text"].startswith("Architecture decisions") for anchor in second["anchors"])

    def test_office_template_analyze_exposes_cached_template_metadata(self, temp_dir, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_METADATA_CACHE_DIR", str(temp_dir / "cache"))
        source = _build_template(temp_dir / "template.docx")
        server = OfficeServer()

        first = server.tool_office_template(
            source_path=str(source),
            destination_path="",
            operation="analyze",
        )
        second = server.tool_office_template(
            source_path=str(source),
            destination_path="",
            operation="analyze",
        )

        assert "template_metadata" in first
        assert first["template_metadata"]["cache"]["reason"] == "stored"
        assert second["template_metadata"]["cache"]["reason"] == "hit"
        assert second["template_metadata"]["tables"][0]["purpose"] == "staffing"
