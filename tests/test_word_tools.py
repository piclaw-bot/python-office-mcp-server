"""
Tests for word_tools.py - Word document processing

Tests cover:
- Word document extraction
- Markdown to Word conversion
- Word to Markdown conversion
"""

import pytest
from docx import Document

# Fixtures temp_dir, word_tools, and sample_docx are provided by conftest.py


class TestWordExtract:
    """Tests for tool_word_extract."""

    def test_extracts_content(self, word_tools, sample_docx):
        """Should extract document content."""
        result = word_tools.tool_word_extract(str(sample_docx))
        assert "content" in result or "paragraphs" in result or "text" in str(result)

    def test_extracts_headings(self, word_tools, sample_docx):
        """Should extract headings."""
        result = word_tools.tool_word_extract(str(sample_docx))
        # Content should contain headings
        content = str(result)
        assert "Test Document" in content or "Section 1" in content


class TestWordToMarkdown:
    """Tests for tool_word_to_markdown."""

    def test_converts_to_markdown(self, word_tools, sample_docx):
        """Should convert Word to markdown."""
        result = word_tools.tool_word_to_markdown(str(sample_docx))
        # Returns string directly, not dict
        assert isinstance(result, str)
        assert "Test Document" in result

    def test_converts_headings(self, word_tools, sample_docx):
        """Should convert headings to markdown format."""
        result = word_tools.tool_word_to_markdown(str(sample_docx))
        # Should have heading markers
        assert "#" in result

    def test_converts_tables(self, word_tools, sample_docx):
        """Should convert tables to markdown format."""
        result = word_tools.tool_word_to_markdown(str(sample_docx))
        # Should have table syntax
        assert "|" in result


class TestWordFileNotFound:
    """Tests for missing file handling across Word tools."""

    @pytest.mark.parametrize("method,path", [
        ("tool_word_extract", "/nonexistent.docx"),
        ("tool_word_to_markdown", "/nonexistent.docx"),
    ])
    def test_file_not_found(self, word_tools, method, path):
        """Should handle missing files."""
        result = getattr(word_tools, method)(path)
        assert "error" in str(result).lower()


class TestWordFromMarkdown:
    """Tests for tool_word_from_markdown."""

    def test_creates_document(self, word_tools, temp_dir):
        """Should create Word document from markdown."""
        md = """# Main Title

This is a paragraph.

## Section 1

Content here.
"""
        output = temp_dir / "created.docx"
        result = word_tools.tool_word_from_markdown(str(output), md)
        assert result.get("success") is True
        assert output.exists()

    def test_creates_headings(self, word_tools, temp_dir):
        """Should create headings from markdown."""
        md = """# Title
## Section
### Subsection
"""
        output = temp_dir / "headings.docx"
        word_tools.tool_word_from_markdown(str(output), md)

        doc = Document(output)
        # Check for heading styles
        heading_styles = [p.style.name for p in doc.paragraphs if "Heading" in p.style.name]
        assert len(heading_styles) >= 2

    def test_creates_paragraphs(self, word_tools, temp_dir):
        """Should create paragraphs from markdown."""
        md = """# Title

First paragraph.

Second paragraph.
"""
        output = temp_dir / "paragraphs.docx"
        word_tools.tool_word_from_markdown(str(output), md)

        doc = Document(output)
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "First paragraph" in " ".join(para_texts)

    def test_creates_tables(self, word_tools, temp_dir):
        """Should create tables from markdown."""
        md = """# Data

| Col1 | Col2 |
|------|------|
| A    | B    |
| C    | D    |
"""
        output = temp_dir / "tables.docx"
        word_tools.tool_word_from_markdown(str(output), md)

        doc = Document(output)
        assert len(doc.tables) >= 1

    def test_creates_blockquotes(self, word_tools, temp_dir):
        """Should render blockquotes as paragraphs."""
        md = """# Quote Section

> This is a blockquote
> spanning multiple lines

Regular paragraph.
"""
        output = temp_dir / "blockquotes.docx"
        word_tools.tool_word_from_markdown(str(output), md)

        doc = Document(output)
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("blockquote" in t for t in para_texts)

    def test_creates_bullet_lists(self, word_tools, temp_dir):
        """Should create bullet lists from markdown."""
        md = """# List

- Item 1
- Item 2
- Item 3
"""
        output = temp_dir / "bullets.docx"
        word_tools.tool_word_from_markdown(str(output), md)

        doc = Document(output)
        # Check for list items
        list_items = [p.text for p in doc.paragraphs if p.text.strip() and "Item" in p.text]
        assert len(list_items) >= 2


class TestWordDocumentStructure:
    """Tests for document structure handling."""

    def test_preserves_formatting(self, word_tools, temp_dir):
        """Should preserve basic formatting."""
        md = """# Title

**Bold text** and *italic text*.
"""
        output = temp_dir / "formatting.docx"
        result = word_tools.tool_word_from_markdown(str(output), md)
        assert result.get("success") is True

    def test_handles_empty_document(self, word_tools, temp_dir):
        """Should handle empty markdown."""
        md = ""
        output = temp_dir / "empty.docx"
        result = word_tools.tool_word_from_markdown(str(output), md)
        # Should succeed even with empty content
        assert output.exists() or "success" in result

    def test_handles_special_characters(self, word_tools, temp_dir):
        """Should handle special characters."""
        md = """# Special Characters

Text with <angle brackets> and & ampersand.
"""
        output = temp_dir / "special.docx"
        result = word_tools.tool_word_from_markdown(str(output), md)
        assert result.get("success") is True
