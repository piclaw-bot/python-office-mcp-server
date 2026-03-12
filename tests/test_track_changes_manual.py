"""
Manual/visual test for track changes.

This script creates a test document with track changes that can be
opened in Microsoft Word to visually verify the changes appear correctly.

Run with: python -m pytest tests/test_track_changes_manual.py -v -s
"""

import sys
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

# Add parent to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.word_advanced_tools import (
    _add_tracked_deletion,
    _add_tracked_insertion,
)

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class TestManualVerification:
    """Tests that create files for manual Word inspection."""

    def test_create_sample_with_track_changes(self, tmp_path):
        """
        Create a sample document with various track changes for manual verification.

        After running this test, open the output file in Word and verify:
        1. Track Changes panel shows the changes
        2. Deletions appear as red strikethrough
        3. Insertions appear as green/blue underline
        4. Author and date are shown correctly
        """
        doc = Document()

        # Title
        doc.add_heading("Track Changes Test Document", 0)
        doc.add_paragraph("Open this in Word and check Review > Track Changes panel.")
        doc.add_paragraph("")

        # Test 1: Simple insertion
        doc.add_heading("Test 1: Insertion Only", level=1)
        para1 = doc.add_paragraph("This paragraph has an insertion: ")
        _add_tracked_insertion(para1, "INSERTED TEXT", author="Test Author")
        para1.add_run(" - end of paragraph.")

        # Test 2: Simple deletion
        doc.add_heading("Test 2: Deletion Only", level=1)
        para2 = doc.add_paragraph("This paragraph has a deletion: ")
        _add_tracked_deletion(para2, "DELETED TEXT", author="Test Author")
        para2.add_run(" - end of paragraph.")

        # Test 3: Replacement (deletion + insertion)
        doc.add_heading("Test 3: Replacement", level=1)
        para3 = doc.add_paragraph("Before replacement ")
        _add_tracked_deletion(para3, "OLD VALUE", author="Test Author")
        _add_tracked_insertion(para3, "NEW VALUE", author="Test Author")
        para3.add_run(" after replacement.")

        # Test 4: Multiple changes by different authors
        doc.add_heading("Test 4: Multiple Authors", level=1)
        para4 = doc.add_paragraph("Changes by different people: ")
        _add_tracked_insertion(para4, "Alice's addition", author="Alice Smith")
        para4.add_run(" and ")
        _add_tracked_insertion(para4, "Bob's addition", author="Bob Jones")
        para4.add_run(".")

        # Save
        output_path = tmp_path / "track_changes_manual_test.docx"
        doc.save(output_path)

        print(f"\n\n{'='*60}")
        print("MANUAL VERIFICATION REQUIRED")
        print(f"{'='*60}")
        print("Open this file in Microsoft Word:")
        print(f"  {output_path}")
        print("")
        print("Verify:")
        print("  1. Go to Review tab > Track Changes > All Markup")
        print("  2. Check that insertions show as underlined (colored)")
        print("  3. Check that deletions show as strikethrough")
        print("  4. Check the Revisions pane shows all changes")
        print(f"{'='*60}\n")

        # Basic validation - file should be openable
        doc2 = Document(output_path)
        assert len(doc2.paragraphs) > 0

    def test_analyze_current_implementation(self, tmp_path):
        """
        Create a test doc and dump the XML to analyze the structure.
        """
        doc = Document()
        para = doc.add_paragraph("Start ")
        _add_tracked_deletion(para, "OLD", author="Tester")
        _add_tracked_insertion(para, "NEW", author="Tester")
        para.add_run(" End")

        output_path = tmp_path / "analyze_structure.docx"
        doc.save(output_path)

        # Extract and print the paragraph XML
        with ZipFile(output_path, "r") as zf:
            xml_content = zf.read("word/document.xml")

        root = ET.fromstring(xml_content)

        # Find the paragraph we added (should be the last one with content)
        body = root.find(f".//{{{WORD_NS}}}body")
        paragraphs = body.findall(f"{{{WORD_NS}}}p")

        print("\n\n" + "="*60)
        print("XML STRUCTURE ANALYSIS")
        print("="*60)

        for i, para in enumerate(paragraphs):
            para_text = "".join(para.itertext())
            if para_text.strip():
                print(f"\nParagraph {i}: '{para_text[:50]}...'")

                # Check for track change elements
                ins_count = len(para.findall(f".//{{{WORD_NS}}}ins"))
                del_count = len(para.findall(f".//{{{WORD_NS}}}del"))

                if ins_count or del_count:
                    print(f"  - Found {ins_count} insertions, {del_count} deletions")

                    # Print structure of ins elements
                    for ins in para.findall(f".//{{{WORD_NS}}}ins"):
                        print(f"  - w:ins id={ins.get(qn('w:id'))} author={ins.get(qn('w:author'))}")
                        for child in ins:
                            tag = child.tag.replace(f"{{{WORD_NS}}}", "w:")
                            print(f"    - Child: {tag}")
                            for subchild in child:
                                subtag = subchild.tag.replace(f"{{{WORD_NS}}}", "w:")
                                print(f"      - {subtag}: '{subchild.text}'")

                    for deletion in para.findall(f".//{{{WORD_NS}}}del"):
                        print(f"  - w:del id={deletion.get(qn('w:id'))} author={deletion.get(qn('w:author'))}")
                        for child in deletion:
                            tag = child.tag.replace(f"{{{WORD_NS}}}", "w:")
                            print(f"    - Child: {tag}")
                            for subchild in child:
                                subtag = subchild.tag.replace(f"{{{WORD_NS}}}", "w:")
                                print(f"      - {subtag}: '{subchild.text}'")

        print("\n" + "="*60 + "\n")

    def test_compare_with_word_generated(self):
        """
        Show what proper Word-generated track changes XML looks like.

        According to OOXML spec (ISO/IEC 29500-1):

        For insertions:
        <w:ins w:id="0" w:author="Joe Smith" w:date="2006-03-31T12:50:00Z">
          <w:r>
            <w:t>inserted text</w:t>
          </w:r>
        </w:ins>

        For deletions:
        <w:del w:id="0" w:author="Joe Smith" w:date="2006-03-31T12:50:00Z">
          <w:r>
            <w:delText>deleted text</w:delText>
          </w:r>
        </w:del>
        """
        print("\n\n" + "="*60)
        print("EXPECTED XML STRUCTURE (per OOXML spec)")
        print("="*60)
        print("""
For a replacement of "OLD" with "NEW":

<w:p>
  <w:r>
    <w:t>Before </w:t>
  </w:r>
  <w:del w:id="0" w:author="Author" w:date="2026-01-20T10:00:00Z">
    <w:r>
      <w:delText>OLD</w:delText>
    </w:r>
  </w:del>
  <w:ins w:id="1" w:author="Author" w:date="2026-01-20T10:00:00Z">
    <w:r>
      <w:t>NEW</w:t>
    </w:r>
  </w:ins>
  <w:r>
    <w:t> After</w:t>
  </w:r>
</w:p>

Key points:
1. w:del and w:ins are SIBLING elements at paragraph level
2. They must be placed IN ORDER where the change occurs
3. w:del uses w:delText, NOT w:t
4. w:ins uses regular w:t
5. Each needs unique w:id
6. Date must be ISO 8601 format
""")
        print("="*60 + "\n")


class TestToolIntegration:
    """Test the high-level tool functions."""

    def test_patch_with_track_changes_creates_changes(self, word_advanced_tools, tmp_path):
        """Verify patch_with_track_changes creates proper revision marks."""
        # Create input doc
        doc = Document()
        doc.add_paragraph("Hello PLACEHOLDER world.")
        input_path = tmp_path / "input.docx"
        doc.save(input_path)

        # Apply changes
        output_path = tmp_path / "output.docx"

        result = word_advanced_tools.tool_word_patch_with_track_changes(
            file_path=str(input_path),
            replacements={"PLACEHOLDER": "REPLACED"},
            author="Integration Test",
            output_path=str(output_path)
        )

        print(f"\nTool result: {result}")

        assert result.get("success"), f"Tool failed: {result}"
        assert result.get("total_changes") == 1

        # Verify XML contains track change elements
        with ZipFile(output_path, "r") as zf:
            xml_content = zf.read("word/document.xml").decode()

        # Should have both ins and del
        has_ins = "<w:ins " in xml_content or "w:ins " in xml_content
        has_del = "<w:del " in xml_content or "w:del " in xml_content

        print(f"Has w:ins: {has_ins}")
        print(f"Has w:del: {has_del}")

        # Print relevant section
        if "w:ins" in xml_content:
            start = xml_content.find("<w:ins")
            if start == -1:
                start = xml_content.find("w:ins")
            end = xml_content.find("</w:ins>", start) + 10
            print(f"\nInsertion XML snippet:\n{xml_content[start:end]}")

        assert has_ins or has_del, "No track change elements found in output"
