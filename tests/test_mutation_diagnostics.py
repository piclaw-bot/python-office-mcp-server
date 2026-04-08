"""Tests for standardized mutation diagnostics across Word and Excel flows."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo

from tools.excel_advanced_tools import ExcelAdvancedTools
from tools.office_unified_tools import OfficeUnifiedTools
from tools.word_advanced_tools import WordAdvancedTools


class UnifiedOfficeTools(OfficeUnifiedTools, WordAdvancedTools, ExcelAdvancedTools):
    pass


def _build_word_template(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("<Customer Name>")
    doc.add_heading("Delivery approach", level=1)
    doc.add_paragraph("[Template Guidance: add delivery approach]")
    doc.save(path)
    return path


def _build_excel_table_workbook(path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Role", "Count"])
    ws.append(["Architect", 1])
    ws.append(["PM", 1])
    table = Table(displayName="Staffing", ref="A1:B3")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    ws.add_table(table)
    wb.save(path)
    return path


class TestMutationDiagnostics:
    def test_word_patch_section_emits_standard_fields(self, temp_dir):
        path = temp_dir / "section.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Old intro")
        doc.save(path)

        result = WordAdvancedTools().tool_word_patch_section(
            file_path=str(path),
            section_title="Introduction",
            new_content=["New intro"],
        )

        assert result["success"] is True
        assert result["status"] == "success"
        assert result["matched_targets"][0]["target"] == "section:Introduction"
        assert result["unmatched_targets"] == []
        assert "diagnostics" in result
        assert "word_insert_at_anchor" in result["next_tools"]

    def test_word_create_sow_from_markdown_surfaces_partial_success(self, temp_dir):
        template = _build_word_template(temp_dir / "template.docx")
        output = temp_dir / "output.docx"
        markdown = """# Sample SOW

Customer: Contoso
Project: Platform Review
Provider: Microsoft

## Introduction
Architecture overview text.

## Assumptions
Customer will provide access.
"""

        result = WordAdvancedTools().tool_word_create_sow_from_markdown(
            output_path=str(output),
            template_path=str(template),
            markdown=markdown,
        )

        assert result["success"] is True
        assert result["status"] == "partial_success"
        assert any(item["target"] == "section:assumptions" for item in result["unmatched_targets"])
        assert result["diagnostics"]["unmapped_sections"]
        assert "word_insert_at_anchor" in result["next_tools"]

    def test_office_patch_word_all_miss_reports_failed(self, temp_dir):
        path = temp_dir / "placeholder.docx"
        doc = Document()
        doc.add_paragraph("No placeholders here")
        doc.save(path)

        result = UnifiedOfficeTools().tool_office_patch(
            file_path=str(path),
            changes=[{"target": "<Customer Name>", "value": "Contoso"}],
        )

        assert result["success"] is False
        assert result["status"] in {"failed", "skipped"}
        assert result["matched_targets"] == []
        assert result["skipped_targets"]
        assert "office_inspect" in result["next_tools"]

    def test_office_patch_excel_all_miss_reports_failed(self, temp_dir):
        path = temp_dir / "book.xlsx"
        wb = Workbook()
        wb.save(path)

        result = UnifiedOfficeTools().tool_office_patch(
            file_path=str(path),
            changes=[{"target": "MissingSheet!A1", "value": "Contoso"}],
        )

        assert result["success"] is False
        assert result["status"] == "failed"
        assert result["matched_targets"] == []
        assert result["unmatched_targets"][0]["target"] == "MissingSheet!A1"
        assert result["edited_sheets"] == []
        assert result["preserved_parts_summary"]["strategy"] == "merge_original_package_with_edited_sheets"

    def test_excel_table_mutations_emit_standard_diagnostics(self, temp_dir):
        path = _build_excel_table_workbook(temp_dir / "table.xlsx")
        tool = ExcelAdvancedTools()

        append_result = tool.tool_excel_append_table_row(
            file_path=str(path),
            table_name="Staffing",
            row_data={"Role": "Engineer", "Count": 2, "Missing": "ignored"},
        )
        update_result = tool.tool_excel_update_table_row(
            file_path=str(path),
            table_name="Staffing",
            row_index=1,
            row_data={"Count": 3, "Unknown": "ignored"},
        )

        assert append_result["success"] is True
        assert append_result["status"] == "partial_success"
        assert append_result["matched_targets"]
        assert append_result["unmatched_targets"][0]["target"] == "column:Missing"
        assert "office_table" in append_result["next_tools"]

        assert update_result["success"] is True
        assert update_result["status"] == "partial_success"
        assert update_result["matched_targets"]
        assert update_result["unmatched_targets"][0]["target"] == "column:Unknown"
        assert update_result["diagnostics"]["updates"]
