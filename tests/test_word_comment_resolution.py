"""Tests for Word comment metadata enrichment and resolve/reopen flows."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID_NS = "http://schemas.microsoft.com/office/word/2016/wordml/cid"


def _read_parts(docx_path: str | Path) -> dict[str, bytes]:
    with zipfile.ZipFile(docx_path, "r") as zf:
        return {name: zf.read(name) for name in zf.namelist()}


def _write_parts(docx_path: str | Path, parts: dict[str, bytes]) -> None:
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, payload in parts.items():
            zf.writestr(name, payload)


def _comment_para_id(comments_root, comment_id: str) -> str | None:
    comment = comments_root.find(f".//{{{W_NS}}}comment[@w:id='{comment_id}']", namespaces={"w": W_NS})
    if comment is None:
        return None
    para = comment.find(f"{{{W_NS}}}p")
    if para is None:
        return None
    return para.get(f"{{{W14_NS}}}paraId")


def _comment_ex_done(comments_ex_root, para_id: str) -> str | None:
    node = comments_ex_root.find(
        f".//{{{W15_NS}}}commentEx[@w15:paraId='{para_id}']",
        namespaces={"w15": W15_NS},
    )
    if node is None:
        return None
    return node.get(f"{{{W15_NS}}}done")


def _create_two_comment_doc(temp_dir, word_advanced_tools):
    path = temp_dir / "resolve_base.docx"
    doc = Document()
    doc.add_paragraph("Alpha target")
    doc.add_paragraph("Beta target")
    doc.save(path)

    add1 = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Alpha target",
        comment_text="Comment alpha",
        author="Manuel",
    )
    add2 = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Beta target",
        comment_text="Comment beta",
        author="Rui Carmo",
    )
    assert add1.get("success") is True
    assert add2.get("success") is True
    return path


def test_t1_get_comments_includes_mixed_done_state(temp_dir, word_advanced_tools, word_tools):
    """T1: word_get_comments returns done=true/false for mixed states."""
    path = _create_two_comment_doc(temp_dir, word_advanced_tools)
    initial = word_tools.tool_word_get_comments(str(path))
    ids = [c["id"] for c in initial["comments"]]

    resolved = word_tools.tool_word_resolve_comment(
        file_path=str(path),
        comment_id=ids[0],
        resolved=True,
    )
    assert resolved.get("success") is True

    got = word_tools.tool_word_get_comments(str(path))
    done_by_id = {c["id"]: c["done"] for c in got["comments"]}
    assert done_by_id[ids[0]] is True
    assert done_by_id[ids[1]] is False


def test_t2_resolve_comment_sets_done_and_rereads_true(temp_dir, word_advanced_tools, word_tools):
    """T2: resolve updates commentsExtended and get() reflects done=true."""
    path = _create_two_comment_doc(temp_dir, word_advanced_tools)
    got = word_tools.tool_word_get_comments(str(path))
    target = got["comments"][0]["id"]

    res = word_tools.tool_word_resolve_comment(str(path), comment_id=target, resolved=True)
    assert res.get("success") is True
    assert res.get("done") is True

    reread = word_tools.tool_word_get_comments(str(path), filter="resolved")
    assert any(c["id"] == target and c.get("done") is True for c in reread.get("comments", []))


def test_t3_reopen_comment_sets_done_false(temp_dir, word_advanced_tools, word_tools):
    """T3: reopen sets w15:done back to 0."""
    path = _create_two_comment_doc(temp_dir, word_advanced_tools)
    got = word_tools.tool_word_get_comments(str(path))
    target = got["comments"][0]["id"]

    assert word_tools.tool_word_resolve_comment(str(path), target, True).get("success") is True
    reopened = word_tools.tool_word_resolve_comment(str(path), target, False)
    assert reopened.get("success") is True
    assert reopened.get("done") is False

    reread = word_tools.tool_word_get_comments(str(path), filter="open")
    assert any(c["id"] == target and c.get("done") is False for c in reread.get("comments", []))


def test_t4_resolve_reply_updates_root_thread_state(temp_dir, word_advanced_tools, word_tools):
    """T4: resolving a reply comment marks the root thread as done."""
    path = temp_dir / "resolve_reply.docx"
    doc = Document()
    doc.add_paragraph("Resolve this thread")
    doc.save(path)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Resolve this thread",
        comment_text="Root comment",
    )
    assert add.get("success") is True

    comments = word_tools.tool_word_get_comments(str(path))
    root_id = comments["comments"][0]["id"]

    reply = word_tools.tool_word_reply_to_comment(
        file_path=str(path),
        comment_id=root_id,
        text="Reply comment",
    )
    assert reply.get("success") is True
    reply_id = reply["reply_comment_id"]

    resolved = word_tools.tool_word_resolve_comment(str(path), comment_id=reply_id, resolved=True)
    assert resolved.get("success") is True
    assert resolved.get("thread_root_comment_id") == root_id

    got = word_tools.tool_word_get_comments(str(path))
    by_id = {c["id"]: c for c in got["comments"]}
    assert by_id[root_id]["done"] is True


def test_t5_fallback_to_comments_ids_when_paraid_missing(temp_dir, word_advanced_tools, word_tools):
    """T5: missing w14:paraId uses commentsIds.xml fallback."""
    path = temp_dir / "comments_ids_fallback.docx"
    doc = Document()
    doc.add_paragraph("Legacy mapping target")
    doc.save(path)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Legacy mapping target",
        comment_text="Legacy style comment",
    )
    assert add.get("success") is True

    current = word_tools.tool_word_get_comments(str(path))
    target_id = current["comments"][0]["id"]

    parts = _read_parts(path)
    comments_root = etree.fromstring(parts["word/comments.xml"])

    para_id = _comment_para_id(comments_root, target_id) or "0F0E0D0C"

    comment_node = comments_root.find(f".//{{{W_NS}}}comment[@w:id='{target_id}']", namespaces={"w": W_NS})
    para = comment_node.find(f"{{{W_NS}}}p")
    para.attrib.pop(f"{{{W14_NS}}}paraId", None)
    parts["word/comments.xml"] = etree.tostring(
        comments_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    ids_root = etree.Element(f"{{{W16CID_NS}}}commentsIds", nsmap={"w16cid": W16CID_NS})
    ids_node = etree.SubElement(ids_root, f"{{{W16CID_NS}}}commentId")
    ids_node.set(f"{{{W16CID_NS}}}id", str(target_id))
    ids_node.set(f"{{{W16CID_NS}}}paraId", str(para_id))
    parts["word/commentsIds.xml"] = etree.tostring(
        ids_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    _write_parts(path, parts)

    got = word_tools.tool_word_get_comments(str(path))
    entry = next(c for c in got["comments"] if c["id"] == target_id)
    assert entry.get("para_id") == para_id


def test_t6_create_comments_extended_entry_when_missing(temp_dir, word_advanced_tools, word_tools):
    """T6: resolve creates commentsExtended part and entry when absent."""
    path = temp_dir / "create_comments_extended.docx"
    doc = Document()
    doc.add_paragraph("No commentsExtended yet")
    doc.save(path)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="No commentsExtended yet",
        comment_text="Needs follow-up",
    )
    assert add.get("success") is True

    got = word_tools.tool_word_get_comments(str(path))
    cid = got["comments"][0]["id"]

    parts_before = _read_parts(path)
    assert "word/commentsExtended.xml" not in parts_before

    resolved = word_tools.tool_word_resolve_comment(str(path), cid, True)
    assert resolved.get("success") is True

    refreshed = word_tools.tool_word_get_comments(str(path))
    para_id = next(c for c in refreshed["comments"] if c["id"] == cid)["para_id"]
    assert para_id

    parts_after = _read_parts(path)
    assert "word/commentsExtended.xml" in parts_after
    ex_root = etree.fromstring(parts_after["word/commentsExtended.xml"])
    assert _comment_ex_done(ex_root, para_id) == "1"


def test_t7_roundtrip_resolve_then_reopen(temp_dir, word_advanced_tools, word_tools):
    """T7: done state toggles 1 -> 0 on resolve/reopen round-trip."""
    path = _create_two_comment_doc(temp_dir, word_advanced_tools)
    got = word_tools.tool_word_get_comments(str(path))
    cid = got["comments"][0]["id"]

    assert word_tools.tool_word_resolve_comment(str(path), cid, True).get("success") is True
    assert word_tools.tool_word_get_comments(str(path), filter="resolved").get("comment_count", 0) >= 1

    assert word_tools.tool_word_resolve_comment(str(path), cid, False).get("success") is True
    reopened = word_tools.tool_word_get_comments(str(path), filter="open")
    assert any(c["id"] == cid and c.get("done") is False for c in reopened.get("comments", []))


def test_word_get_comments_exposes_new_metadata_and_filters(temp_dir, word_advanced_tools, word_tools):
    path = _create_two_comment_doc(temp_dir, word_advanced_tools)

    got = word_tools.tool_word_get_comments(str(path))
    entry = got["comments"][0]
    assert {"done", "is_reply", "parent_id", "para_id"}.issubset(entry.keys())

    resolved_id = got["comments"][0]["id"]
    assert word_tools.tool_word_resolve_comment(str(path), resolved_id, True).get("success") is True

    open_only = word_tools.tool_word_get_comments(str(path), filter="open")
    assert all(c.get("done") is False for c in open_only.get("comments", []))

    resolved_only = word_tools.tool_word_get_comments(str(path), filter="resolved")
    assert all(c.get("done") is True for c in resolved_only.get("comments", []))

    mine = word_tools.tool_word_get_comments(str(path), filter="mine", author="Rui Carmo")
    assert all(str(c.get("author", "")).lower() == "rui carmo" for c in mine.get("comments", []))


def test_word_get_comments_threaded_format_groups_replies(temp_dir, word_advanced_tools, word_tools):
    path = temp_dir / "threaded_format.docx"
    doc = Document()
    doc.add_paragraph("Thread me")
    doc.save(path)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Thread me",
        comment_text="Root",
    )
    assert add.get("success") is True

    root = word_tools.tool_word_get_comments(str(path))["comments"][0]["id"]
    reply = word_tools.tool_word_reply_to_comment(str(path), root, "Reply")
    assert reply.get("success") is True

    threaded = word_tools.tool_word_get_comments(str(path), format="threaded")
    assert threaded.get("thread_count", 0) >= 1
    assert "threads" in threaded
    first = threaded["threads"][0]
    assert first["root"]["id"] == root
    assert any(item["id"] == reply["reply_comment_id"] for item in first["replies"])


def test_word_reply_auto_resolve_marks_thread_done(temp_dir, word_advanced_tools, word_tools):
    path = temp_dir / "reply_auto_resolve.docx"
    doc = Document()
    doc.add_paragraph("Auto resolve target")
    doc.save(path)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Auto resolve target",
        comment_text="Needs action",
    )
    assert add.get("success") is True

    root_id = word_tools.tool_word_get_comments(str(path))["comments"][0]["id"]
    reply = word_tools.tool_word_reply_to_comment(
        file_path=str(path),
        comment_id=root_id,
        text="Done now",
        auto_resolve=True,
    )
    assert reply.get("success") is True
    assert reply.get("resolved") is True

    refreshed = word_tools.tool_word_get_comments(str(path), filter="resolved")
    assert any(c["id"] == root_id and c.get("done") is True for c in refreshed.get("comments", []))
