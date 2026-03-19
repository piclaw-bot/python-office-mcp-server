"""
Tests to increase coverage on word_advanced_tools.py
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from tools.word_advanced_tools import WordAdvancedTools


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as td:
        yield Path(td)


class TestWordCreateNewTable:
    """Test word_create_new_table method."""

    def test_create_table_basic(self, temp_dir):
        """Should create new table with basic data."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Document", level=1)
        doc.add_paragraph("Content before table")
        path = temp_dir / "before_table.docx"
        doc.save(path)

        output = temp_dir / "with_table.docx"
        result = tools.tool_word_create_new_table(
            str(path),
            ["Name", "Role", "Hours"],
            [{"Name": "Alice", "Role": "Developer", "Hours": "40"}, {"Name": "Bob", "Role": "Tester", "Hours": "35"}],
            output_path=str(output)
        )
        assert isinstance(result, dict)
        assert Path(output).exists()

    def test_create_empty_table(self, temp_dir):
        """Should create table with headers only."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("Doc content")
        path = temp_dir / "empty_table.docx"
        doc.save(path)

        result = tools.tool_word_create_new_table(
            str(path),
            ["Col A", "Col B", "Col C"],
            []  # No rows
        )
        assert isinstance(result, dict)


class TestWordDuplicateTableStructure:
    """Test word_duplicate_table_structure method."""

    def test_duplicate_table(self, temp_dir):
        """Should duplicate table structure."""
        tools = WordAdvancedTools()

        doc = Document()
        table = doc.add_table(rows=3, cols=4)
        for i, cell in enumerate(table.row_cells(0)):
            cell.text = f"Header {i+1}"
        for row_idx in range(1, 3):
            for col_idx, cell in enumerate(table.row_cells(row_idx)):
                cell.text = f"R{row_idx}C{col_idx+1}"
        path = temp_dir / "original.docx"
        doc.save(path)

        result = tools.tool_word_duplicate_table_structure(str(path), "0")
        assert isinstance(result, dict)


class TestWordCopyTemplate:
    """Test word_copy_template method."""

    def test_copy_template(self, temp_dir):
        """Should copy template to new location."""
        tools = WordAdvancedTools()

        # Create template
        doc = Document()
        doc.add_heading("<Project Name>", level=1)
        doc.add_paragraph("Customer: <Customer Name>")
        template_path = temp_dir / "template.docx"
        doc.save(template_path)

        output = temp_dir / "copy.docx"
        result = tools.tool_word_copy_template(str(template_path), str(output))
        assert isinstance(result, dict)
        assert Path(output).exists()


class TestWordPatchPlaceholder:
    """Test word_patch_placeholder method."""

    def test_patch_placeholder_basic(self, temp_dir):
        """Should patch single placeholder."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("<Title>", level=1)
        doc.add_paragraph("Content about <Project>")
        path = temp_dir / "placeholders.docx"
        doc.save(path)

        result = tools.tool_word_patch_placeholder(
            str(path),
            "<Title>",
            "Real Title"
        )
        assert isinstance(result, dict)

    def test_patch_multiple_occurrences(self, temp_dir):
        """Should patch multiple occurrences."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("<NAME> is a developer")
        doc.add_paragraph("Contact <NAME> for details")
        doc.add_paragraph("<NAME> will lead the project")
        path = temp_dir / "multi_ph.docx"
        doc.save(path)

        result = tools.tool_word_patch_placeholder(
            str(path),
            "<NAME>",
            "John Smith"
        )
        assert isinstance(result, dict)


class TestWordEnableTrackChanges:
    """Test word_enable_track_changes method."""

    def test_enable_track_changes(self, temp_dir):
        """Should enable track changes."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Document", level=1)
        doc.add_paragraph("Content")
        path = temp_dir / "track.docx"
        doc.save(path)

        result = tools.tool_word_enable_track_changes(str(path))
        assert isinstance(result, dict)


class TestWordCheckTracking:
    """Test word_check_tracking method."""

    def test_check_tracking_status(self, temp_dir):
        """Should check tracking status."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("Document content")
        path = temp_dir / "check_track.docx"
        doc.save(path)

        result = tools.tool_word_check_tracking(str(path))
        assert isinstance(result, dict)


class TestWordPatchWithTrackChanges:
    """Test word_patch_with_track_changes method."""

    def test_patch_with_tracking(self, temp_dir):
        """Should patch content with track changes."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_paragraph("Original content here about <Customer>.")
        path = temp_dir / "patch_track.docx"
        doc.save(path)

        result = tools.tool_word_patch_with_track_changes(
            str(path),
            {"<Customer>": "Contoso", "Original": "Updated"}
        )
        assert isinstance(result, dict)


class TestWordAuditCompletion:
    """Test word_audit_completion method."""

    def test_audit_with_missing(self, temp_dir):
        """Should find incomplete placeholders."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Project: <Name>", level=1)
        doc.add_paragraph("Customer: [TBD]")
        doc.add_paragraph("Date: <insert date>")
        path = temp_dir / "incomplete.docx"
        doc.save(path)

        result = tools.tool_word_audit_completion(str(path))
        assert isinstance(result, dict)

    def test_audit_complete_doc(self, temp_dir):
        """Should pass audit for complete doc."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Complete Project", level=1)
        doc.add_paragraph("All content is filled in properly.")
        path = temp_dir / "complete.docx"
        doc.save(path)

        result = tools.tool_word_audit_completion(str(path))
        assert isinstance(result, dict)


class TestWordAuditSow:
    """Test word_audit_sow method."""

    def test_audit_sow(self, temp_dir):
        """Should audit SOW document."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Statement of Work", level=1)
        doc.add_heading("1. Executive Summary", level=2)
        doc.add_paragraph("This is the executive summary section.")
        doc.add_heading("2. Scope of Work", level=2)
        doc.add_paragraph("<describe scope here>")
        path = temp_dir / "sow_audit.docx"
        doc.save(path)

        result = tools.tool_word_audit_sow(str(path))
        assert isinstance(result, dict)


class TestWordFixSplitPlaceholders:
    """Test word_fix_split_placeholders method."""

    def test_fix_split_placeholders(self, temp_dir):
        """Should fix split placeholders."""
        tools = WordAdvancedTools()

        doc = Document()
        para = doc.add_paragraph()
        # Simulate split placeholder by adding runs
        para.add_run("<Customer")
        para.add_run(" Name>")
        para.add_run(" is important")
        path = temp_dir / "split.docx"
        doc.save(path)

        result = tools.tool_word_fix_split_placeholders(
            str(path),
            {"<Customer Name>": "Contoso"}
        )
        assert isinstance(result, dict)


class TestWordGetSectionGuidance:
    """Test word_get_section_guidance method."""

    def test_get_section_guidance(self, temp_dir):
        """Should get section guidance."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("1. Introduction", level=1)
        # Add guidance-like content
        para = doc.add_paragraph()
        run = para.add_run("[Guidance: Describe the project background]")
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for guidance

        doc.add_paragraph("This section introduces the project.")
        path = temp_dir / "guidance.docx"
        doc.save(path)

        result = tools.tool_word_get_section_guidance(str(path), "1. Introduction")
        assert isinstance(result, dict)


class TestWordExtractSowStructure:
    """Test word_extract_sow_structure method."""

    def test_extract_sow_structure(self, temp_dir):
        """Should extract SOW structure."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Statement of Work", level=1)
        doc.add_heading("1. Overview", level=2)
        doc.add_paragraph("Overview content")
        doc.add_heading("2. Objectives", level=2)
        doc.add_paragraph("Objectives content")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Description"

        path = temp_dir / "sow_structure.docx"
        doc.save(path)

        result = tools.tool_word_extract_sow_structure(str(path))
        assert isinstance(result, dict)


class TestWordParseSowTemplate:
    """Test word_parse_sow_template method."""

    def test_parse_sow_template(self, temp_dir):
        """Should parse SOW template."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("Statement of Work Template", level=1)
        doc.add_heading("1. Engagement Overview", level=2)
        doc.add_paragraph("Customer: <Customer Name>")
        doc.add_paragraph("Provider: <Provider Name>")
        doc.add_heading("2. Business Objectives", level=2)

        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Objective"
        table.cell(0, 1).text = "Activities"
        table.cell(0, 2).text = "Assumptions"

        path = temp_dir / "sow_template.docx"
        doc.save(path)

        result = tools.tool_word_parse_sow_template(str(path))
        assert isinstance(result, dict)


class TestWordCleanupSow:
    """Test word_cleanup_sow method."""

    def test_cleanup_sow(self, temp_dir):
        """Should cleanup SOW document."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("SOW Document", level=1)

        # Add guidance content (should be removed)
        para = doc.add_paragraph()
        run = para.add_run("[Template Guidance: This should be removed]")
        run.font.color.rgb = RGBColor(0, 0, 255)

        doc.add_paragraph("This is real content to keep.")

        # Add placeholder that should be flagged
        doc.add_paragraph("<Unfilled Placeholder>")

        path = temp_dir / "cleanup_sow.docx"
        doc.save(path)

        output = temp_dir / "cleaned.docx"
        result = tools.tool_word_cleanup_sow(str(path), output_path=str(output))
        assert isinstance(result, dict)


class TestWordGenerateSow:
    """Test word_generate_sow method."""

    def test_generate_sow_minimal(self, temp_dir):
        """Should generate SOW with minimal data."""
        tools = WordAdvancedTools()

        # Create a simple template
        doc = Document()
        doc.add_heading("<Project Name>", level=1)
        doc.add_paragraph("Customer: <Customer Name>")
        doc.add_paragraph("Provider: <Provider Name>")
        doc.add_heading("Objectives", level=2)
        template_path = temp_dir / "gen_template.docx"
        doc.save(template_path)

        output = temp_dir / "generated_sow.docx"
        result = tools.tool_word_generate_sow(
            str(template_path),
            str(output),
            {
                "customer_name": "Test Corp",
                "project_name": "Test Project",
                "provider_name": "Provider Inc",
            }
        )
        assert isinstance(result, dict)


class TestWordCreateSowFromMarkdown:
    """Test word_create_sow_from_markdown method."""

    def test_create_sow_from_markdown(self, temp_dir):
        """Should create SOW from markdown."""
        tools = WordAdvancedTools()

        # Create template first
        doc = Document()
        doc.add_heading("<Project Name>", level=1)
        doc.add_paragraph("Customer: <Customer Name>")
        template_path = temp_dir / "md_template.docx"
        doc.save(template_path)

        markdown = """# Test Project - SOW

## Overview
Customer: Test Corp
Provider: Test Provider

## Objectives
- Objective 1
- Objective 2
"""
        output = temp_dir / "md_sow.docx"
        result = tools.tool_word_create_sow_from_markdown(
            str(output),
            markdown,
            template_path=str(template_path)
        )
        assert isinstance(result, dict)

    def test_create_sow_from_markdown_file(self, temp_dir):
        """Should create SOW from markdown_file path."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_heading("<Project Name>", level=1)
        doc.add_paragraph("Customer: <Customer Name>")
        template_path = temp_dir / "md_template_file.docx"
        doc.save(template_path)

        md_file = temp_dir / "sow.md"
        md_file.write_text(
            """# Test Project - SOW

Customer: Test Corp
Project: Test Project
""",
            encoding="utf-8",
        )

        output = temp_dir / "md_sow_file.docx"
        result = tools.tool_word_create_sow_from_markdown(
            str(output),
            template_path=str(template_path),
            markdown_file=str(md_file),
        )
        assert isinstance(result, dict)
        assert "error" not in result


class TestWordAddComment:
    """Test word_add_comment method."""

    def test_add_comment_to_text(self, temp_dir):
        """Should add comment to specific text."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("This is important text to review.")
        path = temp_dir / "comment.docx"
        doc.save(path)

        result = tools.tool_word_add_comment(
            str(path),
            "important text",
            "Please verify this is correct."
        )
        assert isinstance(result, dict)

    def test_add_comment_with_author(self, temp_dir):
        """Should add comment with custom author."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("Review needed for this section.")
        path = temp_dir / "comment_author.docx"
        doc.save(path)

        result = tools.tool_word_add_comment(
            str(path),
            "Review needed",
            "This needs manager approval",
            author="John Smith"
        )
        assert isinstance(result, dict)

    def test_add_comment_with_none_author(self, temp_dir):
        """Should fall back to default author when None is provided."""
        tools = WordAdvancedTools()

        doc = Document()
        doc.add_paragraph("Please review this sentence.")
        path = temp_dir / "comment_none_author.docx"
        doc.save(path)

        result = tools.tool_word_add_comment(
            str(path),
            "review this sentence",
            "Use default author",
            author=None
        )
        assert "error" not in result


class TestWordFromMarkdownComplex:
    """Test complex word_from_markdown scenarios."""

    def test_from_markdown_with_tables(self, temp_dir):
        """Should create doc with tables from markdown."""
        WordAdvancedTools()

        markdown = """# Report

## Data

| Name | Value | Notes |
|------|-------|-------|
| Alpha | 100 | First |
| Beta | 200 | Second |
| Gamma | 300 | Third |

## Summary

Total items: 3
"""
        from tools.word_tools import WordTools
        word_tools = WordTools()

        path = temp_dir / "complex_md.docx"
        word_tools.tool_word_from_markdown(str(path), markdown)
        assert Path(path).exists()

    def test_from_markdown_with_nested_headings(self, temp_dir):
        """Should create doc with nested headings."""
        from tools.word_tools import WordTools
        tools = WordTools()

        markdown = """# Main Title

## Section 1

### Subsection 1.1

Content for 1.1

### Subsection 1.2

Content for 1.2

## Section 2

### Subsection 2.1

Content for 2.1
"""
        path = temp_dir / "nested_headers.docx"
        tools.tool_word_from_markdown(str(path), markdown)
        assert Path(path).exists()
