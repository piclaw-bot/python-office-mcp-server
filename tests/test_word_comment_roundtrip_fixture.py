"""End-to-end round-trip fixture tests for Word comment workflows."""

from __future__ import annotations

import pytest
from docx import Document

from tools import TOOL_CLASSES


@pytest.fixture
def combined_tools():
    class CombinedTools(*TOOL_CLASSES):
        pass

    return CombinedTools()


def _build_roundtrip_doc(temp_dir):
    path = temp_dir / "comment_roundtrip.docx"
    doc = Document()
    doc.add_paragraph("Roundtrip scope item")
    doc.add_paragraph("Secondary note")
    doc.save(path)
    return path


def test_word_comment_roundtrip_direct_tools(temp_dir, word_advanced_tools, word_tools):
    """Direct tools round-trip: add -> reply+resolve -> threaded get -> reopen."""
    path = _build_roundtrip_doc(temp_dir)

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(path),
        target_text="Roundtrip scope item",
        comment_text="Please confirm scope wording",
        author="Manuel",
    )
    assert add.get("success") is True

    initial = word_tools.tool_word_get_comments(str(path), format="threaded")
    assert initial.get("thread_count", 0) == 1
    root = initial["threads"][0]["root"]
    root_id = root["id"]
    assert root.get("done") is False

    reply = word_tools.tool_word_reply_to_comment(
        file_path=str(path),
        comment_id=root_id,
        text="Done — wording updated",
        author="Rui Carmo",
        auto_resolve=True,
    )
    assert reply.get("success") is True
    assert reply.get("resolved") is True

    after_reply = word_tools.tool_word_get_comments(str(path), format="threaded")
    thread = after_reply["threads"][0]
    assert thread["root"]["id"] == root_id
    assert thread["root"]["done"] is True
    assert any(r["id"] == reply.get("reply_comment_id") for r in thread["replies"])

    reopen = word_tools.tool_word_resolve_comment(
        file_path=str(path),
        comment_id=reply.get("reply_comment_id"),
        resolved=False,
    )
    assert reopen.get("success") is True
    assert reopen.get("thread_root_comment_id") == root_id

    open_only = word_tools.tool_word_get_comments(str(path), filter="open")
    assert any(c["id"] == root_id and c.get("done") is False for c in open_only.get("comments", []))


def test_word_comment_roundtrip_unified_tool(temp_dir, combined_tools):
    """Unified round-trip: add/get(threaded)/reply/resolve/reopen/delete."""
    path = _build_roundtrip_doc(temp_dir)

    add = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="add",
        target="Roundtrip scope item",
        text="Initial review note",
        author="Reviewer",
    )
    assert add.get("success") is True

    got = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="get",
        format="threaded",
    )
    assert got.get("thread_count", 0) == 1
    root_id = got["threads"][0]["root"]["id"]

    reply = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="reply",
        target=str(root_id),
        text="Acknowledged",
        author="Rui Carmo",
    )
    assert reply.get("success") is True

    resolved = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="resolve",
        target=str(root_id),
    )
    assert resolved.get("success") is True
    assert resolved.get("done") is True

    resolved_view = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="get",
        filter="resolved",
    )
    assert any(c["id"] == str(root_id) and c.get("done") is True for c in resolved_view.get("comments", []))

    reopened = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="reopen",
        target=str(root_id),
    )
    assert reopened.get("success") is True
    assert reopened.get("done") is False

    open_view = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="get",
        filter="open",
    )
    assert any(c["id"] == str(root_id) and c.get("done") is False for c in open_view.get("comments", []))

    deleted = combined_tools.tool_office_comment(
        file_path=str(path),
        operation="delete",
        target=str(root_id),
    )
    assert deleted.get("success") is True


def test_word_resolve_roundtrip_with_output_path(temp_dir, word_advanced_tools, word_tools):
    """Resolve to output copy keeps source unchanged and toggles copy state."""
    source = _build_roundtrip_doc(temp_dir)
    output = temp_dir / "comment_roundtrip_out.docx"

    add = word_advanced_tools.tool_word_add_comment(
        file_path=str(source),
        target_text="Roundtrip scope item",
        comment_text="Track this",
    )
    assert add.get("success") is True

    source_get = word_tools.tool_word_get_comments(str(source))
    cid = source_get["comments"][0]["id"]
    assert source_get["comments"][0]["done"] is False

    resolved = word_tools.tool_word_resolve_comment(
        file_path=str(source),
        comment_id=cid,
        resolved=True,
        output_path=str(output),
    )
    assert resolved.get("success") is True

    source_after = word_tools.tool_word_get_comments(str(source))
    assert any(c["id"] == cid and c.get("done") is False for c in source_after.get("comments", []))

    output_after = word_tools.tool_word_get_comments(str(output))
    assert any(c["id"] == cid and c.get("done") is True for c in output_after.get("comments", []))

    reopened = word_tools.tool_word_resolve_comment(
        file_path=str(output),
        comment_id=cid,
        resolved=False,
    )
    assert reopened.get("success") is True

    output_reopen = word_tools.tool_word_get_comments(str(output), filter="open")
    assert any(c["id"] == cid and c.get("done") is False for c in output_reopen.get("comments", []))
