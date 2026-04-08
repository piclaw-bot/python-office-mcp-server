"""Tests for Word anchor discovery and document maps."""

from __future__ import annotations

from docx import Document

from tools.word_advanced_tools import WordAdvancedTools, _get_text_with_track_changes


class TestWordAnchorDiscovery:
    def test_headings_and_paragraphs_surface_as_anchors(self, temp_dir):
        path = temp_dir / "anchors.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Customer context paragraph")
        doc.add_heading("Delivery approach", level=1)
        doc.add_paragraph("Use iterative delivery")
        doc.save(path)

        result = WordAdvancedTools().tool_word_list_anchors(str(path))

        assert result["count"] >= 4
        assert any(item["type"] == "section_heading" and item["anchor_text"] == "Introduction" for item in result["anchors"])
        assert any(item["type"] == "paragraph" and "Customer context" in item["anchor_text"] for item in result["anchors"])

    def test_anchor_text_filtering_works(self, temp_dir):
        path = temp_dir / "anchors.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Customer context paragraph")
        doc.add_heading("Delivery approach", level=1)
        doc.add_paragraph("Use iterative delivery")
        doc.save(path)

        result = WordAdvancedTools().tool_word_list_anchors(str(path), query="delivery")

        assert result["count"] >= 1
        assert all("delivery" in item["anchor_text"].lower() for item in result["anchors"])

    def test_document_map_includes_sections_tables_placeholders_and_anchors(self, temp_dir):
        path = temp_dir / "map.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("<Customer Name>")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Role"
        table.rows[0].cells[1].text = "Count"
        table.rows[1].cells[0].text = "Architect"
        table.rows[1].cells[1].text = "1"
        doc.save(path)

        result = WordAdvancedTools().tool_word_document_map(str(path))

        assert result["counts"]["sections"] == 1
        assert result["counts"]["tables"] == 1
        assert result["counts"]["placeholders"] >= 1
        assert result["counts"]["anchors"] >= 2
        assert result["anchors"]

    def test_discover_anchor_then_insert(self, temp_dir):
        path = temp_dir / "insert.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Current intro")
        doc.add_heading("Delivery approach", level=1)
        doc.add_paragraph("Current delivery")
        doc.save(path)

        tools = WordAdvancedTools()
        anchors = tools.tool_word_list_anchors(str(path), query="delivery")
        anchor = next(item for item in anchors["anchors"] if item["anchor_text"] == "Delivery approach")

        result = tools.tool_word_insert_at_anchor(
            file_path=str(path),
            anchor_text=anchor["anchor_text"],
            content="Inserted after discovered anchor",
            position="after",
        )

        assert result["success"] is True

        doc_after = Document(path)
        texts = [_get_text_with_track_changes(p).strip() for p in doc_after.paragraphs]
        idx = texts.index("Delivery approach")
        assert texts[idx + 1] == "Inserted after discovered anchor"

    def test_section_listing_points_to_anchor_discovery(self, temp_dir):
        path = temp_dir / "sections.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Current intro")
        doc.save(path)

        result = WordAdvancedTools().tool_word_list_sections(str(path))

        assert "word_list_anchors" in result["next_tools"]
        assert "word_document_map" in result["next_tools"]
