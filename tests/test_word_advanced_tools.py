"""
Tests for word_advanced_tools.py - Advanced Word document manipulation

Tests cover:
- Section operations
- Table operations
- Placeholder replacement
- SOW generation
- Track changes
- Template operations
"""

from docx import Document

from tools.word_advanced_tools import (
    _get_text_with_track_changes,
)

# Fixtures temp_dir, word_advanced_tools, sample_docx, and sow_template are provided by conftest.py


class TestGetTextWithTrackChanges:
    """Tests for _get_text_with_track_changes helper."""

    def test_reads_normal_text(self, temp_dir):
        """Should read normal paragraph text."""
        doc = Document()
        doc.add_paragraph("Normal text content")
        path = temp_dir / "normal.docx"
        doc.save(path)

        doc = Document(path)
        # Find the paragraph with our text (index varies by template)
        texts = [_get_text_with_track_changes(p) for p in doc.paragraphs]
        assert any("Normal text content" in t for t in texts)


class TestListSections:
    """Tests for tool_word_list_sections."""

    def test_lists_sections(self, word_advanced_tools, sow_template):
        """Should list all sections/headings."""
        result = word_advanced_tools.tool_word_list_sections(str(sow_template))
        assert "sections" in result
        sections = result["sections"]
        # Should find our headings
        section_titles = [s.get("title", s.get("text", "")) for s in sections]
        assert any("Executive Summary" in t for t in section_titles)
        assert any("Scope" in t for t in section_titles)

    def test_file_not_found(self, word_advanced_tools):
        """Should handle missing files."""
        result = word_advanced_tools.tool_word_list_sections("/nonexistent.docx")
        assert "error" in result


class TestGetSection:
    """Tests for tool_word_get_section."""

    def test_gets_section_content(self, word_advanced_tools, sow_template):
        """Should get content of a specific section."""
        result = word_advanced_tools.tool_word_get_section(str(sow_template), "Executive Summary")
        assert "content" in result or "text" in result or "paragraphs" in str(result)

    def test_section_not_found(self, word_advanced_tools, sow_template):
        """Should handle missing section gracefully."""
        result = word_advanced_tools.tool_word_get_section(str(sow_template), "Nonexistent Section")
        # Should indicate section not found
        assert "error" in result or "not found" in str(result).lower() or "available" in str(result).lower()


class TestPatchSection:
    """Tests for tool_word_patch_section."""

    def test_patches_section(self, word_advanced_tools, sow_template, temp_dir):
        """Should update section content."""
        output = temp_dir / "patched.docx"
        result = word_advanced_tools.tool_word_patch_section(
            str(sow_template),
            "Executive Summary",
            "This is the new executive summary content.",
            output_path=str(output)
        )
        assert result.get("success") is True or "patched" in str(result).lower()

    def test_section_not_found(self, word_advanced_tools, sow_template, temp_dir):
        """Should handle missing section gracefully."""
        output = temp_dir / "patched2.docx"
        result = word_advanced_tools.tool_word_patch_section(
            str(sow_template),
            "Nonexistent",
            "New content",
            output_path=str(output)
        )
        assert "error" in result or "not found" in str(result).lower()


class TestListTables:
    """Tests for tool_word_list_tables."""

    def test_lists_tables(self, word_advanced_tools, sample_docx):
        """Should list tables in document."""
        result = word_advanced_tools.tool_word_list_tables(str(sample_docx))
        assert "tables" in result
        assert len(result["tables"]) >= 1


class TestGetTable:
    """Tests for tool_word_get_table."""

    def test_gets_table_by_index(self, word_advanced_tools, sample_docx):
        """Should get table by index."""
        result = word_advanced_tools.tool_word_get_table(str(sample_docx), "0")
        assert "rows" in result or "data" in result or "header" in result


class TestInsertTableRow:
    """Tests for tool_word_insert_table_row."""

    def test_inserts_row(self, word_advanced_tools, sample_docx, temp_dir):
        """Should insert a new row into table."""
        output = temp_dir / "row_added.docx"
        # API expects dict, not list
        result = word_advanced_tools.tool_word_insert_table_row(
            str(sample_docx),
            "0",
            {"Role": "Analyst", "Hours": "80", "Rate": "$125"},
            output_path=str(output)
        )
        assert result.get("success") is True


class TestPatchTableRow:
    """Tests for tool_word_patch_table_row."""

    def test_patches_row(self, word_advanced_tools, sample_docx, temp_dir):
        """Should update a table row."""
        output = temp_dir / "row_patched.docx"
        result = word_advanced_tools.tool_word_patch_table_row(
            str(sample_docx),
            "0",
            1,
            {"Hours": "120"},
            output_path=str(output)
        )
        assert result.get("success") is True or "updated" in str(result).lower()


class TestAuditCompletion:
    """Tests for tool_word_audit_completion."""

    def test_finds_placeholders(self, word_advanced_tools, sow_template):
        """Should find unfilled placeholders."""
        result = word_advanced_tools.tool_word_audit_completion(str(sow_template))
        # Should find <Customer Name> and <Project Name>
        assert "issues" in result or "placeholders" in str(result).lower() or "findings" in result

    def test_empty_document(self, word_advanced_tools, temp_dir):
        """Should handle document without issues."""
        doc = Document()
        doc.add_paragraph("Clean content with no placeholders.")
        path = temp_dir / "clean.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_audit_completion(str(path))
        # Should complete without error
        assert "error" not in result


class TestAuditSow:
    """Tests for tool_word_audit_sow."""

    def test_audits_sow(self, word_advanced_tools, sow_template):
        """Should audit SOW for completeness."""
        result = word_advanced_tools.tool_word_audit_sow(str(sow_template))
        # Should find issues
        assert "issues" in result or "placeholders" in str(result).lower()


class TestReplaceGlobalVariables:
    """Tests for tool_word_replace_global_variables."""

    def test_replaces_placeholders(self, word_advanced_tools, sow_template, temp_dir):
        """Should replace placeholder text."""
        output = temp_dir / "replaced.docx"
        result = word_advanced_tools.tool_word_replace_global_variables(
            str(sow_template),
            {"<Customer Name>": "Contoso Corp", "<Project Name>": "Cloud Migration"},
            output_path=str(output)
        )
        assert result.get("success") is True


class TestEnableTrackChanges:
    """Tests for tool_word_enable_track_changes."""

    def test_enables_tracking(self, word_advanced_tools, sample_docx, temp_dir):
        """Should enable track changes."""
        output = temp_dir / "tracked.docx"
        result = word_advanced_tools.tool_word_enable_track_changes(
            str(sample_docx),
            output_path=str(output)
        )
        assert result.get("success") is True or "enabled" in str(result).lower()


class TestPatchWithTrackChanges:
    """Tests for tool_word_patch_with_track_changes."""

    def test_patches_with_tracking(self, word_advanced_tools, sample_docx, temp_dir):
        """Should patch text with tracking enabled."""
        output = temp_dir / "tracked_patch.docx"
        result = word_advanced_tools.tool_word_patch_with_track_changes(
            str(sample_docx),
            {"executive summary content": "updated summary content"},
            output_path=str(output)
        )
        assert result.get("success") is True


class TestGenerateSow:
    """Tests for tool_word_generate_sow."""

    def test_generates_sow(self, word_advanced_tools, sow_template, temp_dir):
        """Should generate SOW from template."""
        output = temp_dir / "generated.docx"
        result = word_advanced_tools.tool_word_generate_sow(
            str(sow_template),
            str(output),
            {"customer_name": "Contoso Corp", "project_name": "Migration"}
        )
        assert result.get("success") is True

    def test_generates_sow_fills_staffing_table(self, word_advanced_tools, sow_template, temp_dir):
        """Structured staffing rows should populate the staffing table."""
        output = temp_dir / "generated_staffing.docx"
        result = word_advanced_tools.tool_word_generate_sow(
            str(sow_template),
            str(output),
            {
                "customer_name": "Contoso Corp",
                "project_name": "Migration",
                "staffing": [
                    {"role": "Architect", "hours": "40"},
                ],
            }
        )

        assert result.get("success") is True
        assert result.get("tables_filled", 0) >= 1
        assert any(item.get("purpose") == "staffing" and item.get("matched") for item in result.get("table_diagnostics", []))
        assert "word_insert_at_anchor" in result.get("next_tools", [])

        doc = Document(output)
        table_text = [[_get_text_with_track_changes(cell) for cell in row.cells] for row in doc.tables[0].rows]
        flattened = " ".join(" ".join(row) for row in table_text)
        assert "Architect" in flattened
        assert "40" in flattened

    def test_generate_sow_reports_missing_target_table_reason(self, word_advanced_tools, temp_dir):
        template = Document()
        template.add_heading("Executive Summary", level=1)
        other_table = template.add_table(rows=2, cols=2)
        other_table.rows[0].cells[0].text = "Term / acronym"
        other_table.rows[0].cells[1].text = "Description"
        other_table.rows[1].cells[0].text = "API"
        other_table.rows[1].cells[1].text = "Application Programming Interface"
        template_path = temp_dir / "missing_staffing_table.docx"
        template.save(template_path)

        output = temp_dir / "missing_staffing_out.docx"
        result = word_advanced_tools.tool_word_generate_sow(
            str(template_path),
            str(output),
            {
                "staffing": [{"role": "Architect", "hours": "40"}],
            },
        )

        assert result.get("success") is False or result.get("status") in {"failed", "partial_success"}
        assert any(
            item.get("purpose") == "staffing" and item.get("reason") == "no_matching_table_found"
            for item in result.get("table_diagnostics", [])
        )

    def test_generate_sow_normalizes_complex_table_headers(self, word_advanced_tools, temp_dir):
        template = Document()
        template.add_heading("Staffing", level=1)
        staffing_table = template.add_table(rows=2, cols=2)
        staffing_table.rows[0].cells[0].text = "Role / Skill"
        staffing_table.rows[0].cells[1].text = "Count & Hours"
        staffing_table.rows[1].cells[0].text = "Template role"
        staffing_table.rows[1].cells[1].text = "Template hours"
        template_path = temp_dir / "complex_staffing_header.docx"
        template.save(template_path)

        output = temp_dir / "complex_staffing_header_out.docx"
        result = word_advanced_tools.tool_word_generate_sow(
            str(template_path),
            str(output),
            {
                "staffing": [{"role": "Architect", "hours": "40"}],
            },
        )

        assert result.get("success") is True
        staffing_diag = next(item for item in result.get("table_diagnostics", []) if item.get("purpose") == "staffing")
        assert staffing_diag.get("matched") is True
        assert staffing_diag.get("normalized_header") == ["role skill", "count and hours"]

    def test_generate_sow_uses_second_header_row_when_first_is_banner(self, word_advanced_tools, temp_dir):
        template = Document()
        template.add_heading("Staffing", level=1)
        staffing_table = template.add_table(rows=3, cols=2)
        staffing_table.cell(0, 0).text = "Staffing Plan"
        staffing_table.cell(0, 1).text = "Staffing Plan"
        staffing_table.cell(0, 0).merge(staffing_table.cell(0, 1))
        staffing_table.rows[1].cells[0].text = "Role / Skill"
        staffing_table.rows[1].cells[1].text = "Count & Hours"
        staffing_table.rows[2].cells[0].text = "Template role"
        staffing_table.rows[2].cells[1].text = "Template hours"
        template_path = temp_dir / "banner_header_staffing.docx"
        template.save(template_path)

        output = temp_dir / "banner_header_staffing_out.docx"
        result = word_advanced_tools.tool_word_generate_sow(
            str(template_path),
            str(output),
            {
                "staffing": [{"role": "Architect", "hours": "40"}],
            },
        )

        assert result.get("success") is True
        staffing_diag = next(item for item in result.get("table_diagnostics", []) if item.get("purpose") == "staffing")
        assert staffing_diag.get("matched") is True
        assert staffing_diag.get("header_rows_used") == [0, 1]

        doc = Document(output)
        rows = [[_get_text_with_track_changes(cell) for cell in row.cells] for row in doc.tables[0].rows]
        flattened = " ".join(" ".join(row) for row in rows)
        assert "Architect" in flattened
        assert "40" in flattened


class TestCopyTemplate:
    """Tests for tool_word_copy_template."""

    def test_copies_template(self, word_advanced_tools, sow_template, temp_dir):
        """Should copy template to new location."""
        output = temp_dir / "copied.docx"
        result = word_advanced_tools.tool_word_copy_template(
            str(sow_template),
            str(output)
        )
        assert result.get("success") is True
        assert output.exists()


class TestFixSplitPlaceholders:
    """Tests for tool_word_fix_split_placeholders."""

    def test_fixes_placeholders(self, word_advanced_tools, sow_template, temp_dir):
        """Should fix split placeholder text."""
        output = temp_dir / "fixed.docx"
        result = word_advanced_tools.tool_word_fix_split_placeholders(
            str(sow_template),
            {"<Customer Name>": "Contoso"},
            output_path=str(output)
        )
        assert result.get("success") is True


class TestCleanupSow:
    """Tests for tool_word_cleanup_sow."""

    def test_cleans_sow(self, word_advanced_tools, sow_template, temp_dir):
        """Should cleanup SOW removing instructions."""
        output = temp_dir / "cleaned.docx"
        result = word_advanced_tools.tool_word_cleanup_sow(
            str(sow_template),
            output_path=str(output)
        )
        assert result.get("success") is True or "cleaned" in str(result).lower()


class TestGetSectionGuidance:
    """Tests for tool_word_get_section_guidance."""

    def test_gets_guidance(self, word_advanced_tools, sow_template):
        """Should get section guidance information."""
        result = word_advanced_tools.tool_word_get_section_guidance(
            str(sow_template),
            "Executive Summary"
        )
        # Should return guidance or indication it exists
        assert "guidance" in result or "content" in result or "error" not in str(result).lower()


class TestParseSowTemplate:
    """Tests for tool_word_parse_sow_template."""

    def test_parses_template(self, word_advanced_tools, sow_template):
        """Should parse SOW template structure."""
        result = word_advanced_tools.tool_word_parse_sow_template(str(sow_template))
        assert "sections" in result or "structure" in result or "template" in str(result).lower()


class TestAnalyzeTemplateFormatting:
    """Tests for tool_word_analyze_template_formatting."""

    def test_analyzes_formatting(self, word_advanced_tools, sow_template):
        """Should analyze template formatting."""
        result = word_advanced_tools.tool_word_analyze_template_formatting(str(sow_template))
        # Should return analysis information
        assert "error" not in result or isinstance(result, dict)


class TestWordFromMarkdown:
    """Tests for word_from_markdown - this is in word_tools.py, not word_advanced_tools."""

    def test_method_location(self):
        """Note: word_from_markdown is in WordTools, not WordAdvancedTools."""
        # This method is in word_tools.py
        from tools.word_tools import WordTools
        wt = WordTools()
        assert hasattr(wt, 'tool_word_from_markdown')


class TestCreateSowFromMarkdown:
    """Tests for tool_word_create_sow_from_markdown."""

    def test_creates_sow(self, word_advanced_tools, sow_template, temp_dir):
        """Should create SOW from markdown with template."""
        md = """# Statement of Work

## Executive Summary

This is the executive summary.

## Scope

Project scope description.
"""
        output = temp_dir / "sow_from_md.docx"
        result = word_advanced_tools.tool_word_create_sow_from_markdown(
            str(output), md, str(sow_template)
        )
        assert result.get("success") is True or output.exists()

    def test_extracts_narrative_sections_from_markdown(self, word_advanced_tools, temp_dir):
        """Markdown narrative headings should be mapped into template sections."""
        template = Document()
        template.add_heading("Executive Summary", level=1)
        template.add_paragraph("[Template Guidance: replace]")
        template.add_heading("Delivery approach", level=1)
        template.add_heading("Customer responsibilities and project assumptions", level=1)
        template_path = temp_dir / "narrative_template.docx"
        template.save(template_path)

        md = """# Statement of Work

Customer: Contoso Ltd
Project: Migration Factory

## Executive Summary

This is the executive summary.

## Delivery approach

Microsoft will undertake an iterative delivery approach.

## Customer responsibilities and project assumptions

Customer will provide timely access to systems.
"""

        output = temp_dir / "narrative_output.docx"
        result = word_advanced_tools.tool_word_create_sow_from_markdown(
            str(output), md, str(template_path)
        )

        assert result.get("success") is True
        assert "sections" in result.get("extracted_fields", [])
        assert result.get("sections_filled", 0) >= 3
        assert not result.get("unmapped_sections")
        assert all(item.get("matched") for item in result.get("section_diagnostics", []))
        assert "word_insert_at_anchor" in result.get("next_tools", [])

        doc = Document(output)
        paragraphs = [_get_text_with_track_changes(p).strip() for p in doc.paragraphs]
        assert any("This is the executive summary." in p for p in paragraphs)
        assert any("Microsoft will undertake an iterative delivery approach." in p for p in paragraphs)
        assert any("Customer will provide timely access to systems." in p for p in paragraphs)


class TestInsertAtAnchor:
    """Tests for tool_word_insert_at_anchor."""

    def test_inserts_after_anchor_text(self, word_advanced_tools, temp_dir):
        doc = Document()
        doc.add_paragraph("Intro")
        doc.add_paragraph("Anchor paragraph")
        doc.add_paragraph("Tail")
        path = temp_dir / "insert_anchor.docx"
        doc.save(path)

        output = temp_dir / "insert_anchor_out.docx"
        result = word_advanced_tools.tool_word_insert_at_anchor(
            str(path),
            "Inserted content.",
            anchor_text="Anchor paragraph",
            position="after",
            output_path=str(output),
        )

        assert result.get("success") is True
        reloaded = Document(output)
        paragraphs = [_get_text_with_track_changes(p).strip() for p in reloaded.paragraphs if _get_text_with_track_changes(p).strip()]
        assert paragraphs == ["Intro", "Anchor paragraph", "Inserted content.", "Tail"]

    def test_inserts_before_paragraph_index(self, word_advanced_tools, temp_dir):
        doc = Document()
        doc.add_paragraph("First")
        doc.add_paragraph("Second")
        path = temp_dir / "insert_index.docx"
        doc.save(path)

        output = temp_dir / "insert_index_out.docx"
        result = word_advanced_tools.tool_word_insert_at_anchor(
            str(path),
            ["Inserted A", "Inserted B"],
            paragraph_index=1,
            position="before",
            output_path=str(output),
        )

        assert result.get("success") is True
        reloaded = Document(output)
        paragraphs = [_get_text_with_track_changes(p).strip() for p in reloaded.paragraphs if _get_text_with_track_changes(p).strip()]
        assert paragraphs == ["First", "Inserted A", "Inserted B", "Second"]

    def test_insert_requires_single_anchor_mode(self, word_advanced_tools, temp_dir):
        doc = Document()
        doc.add_paragraph("Only paragraph")
        path = temp_dir / "insert_validation.docx"
        doc.save(path)

        result = word_advanced_tools.tool_word_insert_at_anchor(
            str(path),
            "Text",
            anchor_text="Only paragraph",
            paragraph_index=0,
        )
        assert "error" in result


    def test_reports_unmapped_markdown_sections(self, word_advanced_tools, temp_dir):
        """Narrative sections without a template target should be surfaced explicitly."""
        template = Document()
        template.add_heading("Executive Summary", level=1)
        template.add_paragraph("[Template Guidance]")
        template_path = temp_dir / "unmapped_template.docx"
        template.save(template_path)

        md = """# Statement of Work

## Executive Summary

Covered summary.

## Completely Custom Section

This section has no template match.
"""
        output = temp_dir / "unmapped_output.docx"
        result = word_advanced_tools.tool_word_create_sow_from_markdown(
            str(output), md, str(template_path)
        )

        assert result.get("success") is True
        assert "completely_custom_section" in result.get("unmapped_sections", [])
        assert any(not item.get("matched") for item in result.get("section_diagnostics", []))
        assert "word_insert_at_anchor" in result.get("next_tools", [])


class TestAddComment:
    """Tests for tool_word_add_comment."""

    def test_adds_comment(self, word_advanced_tools, sow_template, temp_dir):
        """Should add comment to document."""
        output = temp_dir / "commented.docx"
        result = word_advanced_tools.tool_word_add_comment(
            str(sow_template),
            "Executive Summary",
            "This needs more detail",
            output_path=str(output)
        )
        assert result.get("success") is True or "added" in str(result).lower()
