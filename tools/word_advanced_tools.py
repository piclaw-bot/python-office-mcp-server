#!/usr/bin/env python3
"""
word_advanced_tools.py - Advanced MCP tools for Word document manipulation

Provides comprehensive tools to:
- Parse and manipulate document structure (sections, tables)
- Fill templates with structured data
- Generate and audit deliverable documents
- Track changes and add comments

Quality bar: All tools preserve document structure by editing templates
rather than creating new documents from scratch.
"""

import contextlib
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .save_utils import open_docx_with_retries, resolve_office_path, safe_save_docx

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False



# Author name for tracked changes and comments (from environment or default)
DEFAULT_AUTHOR = os.environ.get("MCP_AUTHOR", "Solution Architect Agent")

# Track changes revision ID counter (global for document consistency)
_revision_id = 0


def _next_revision_id() -> str:
    """Get next unique revision ID for track changes."""
    global _revision_id
    _revision_id += 1
    return str(_revision_id)


def _enable_track_revisions(doc) -> None:
    """Enable track revisions in a document's settings.

    This sets the w:trackRevisions element in settings.xml which tells
    Word to display tracked changes. Without this, Word will silently
    accept all changes when opening the document.

    Per OOXML spec, an empty <w:trackRevisions/> element means "true".
    The w:val attribute is only needed to explicitly disable it.
    """
    settings = doc.settings.element

    # Check if trackRevisions already exists
    existing = settings.find(qn("w:trackRevisions"))
    if existing is not None:
        # Remove any w:val attribute to ensure it's enabled
        if existing.get(qn("w:val")) == "false" or existing.get(qn("w:val")) == "0":
            existing.attrib.pop(qn("w:val"), None)
        return

    # Create and add trackRevisions element (empty = true)
    track_revisions = OxmlElement("w:trackRevisions")
    settings.append(track_revisions)


def _add_tracked_insertion(paragraph, text: str, author: str = "MCP Server") -> None:
    """Add text to a paragraph as a tracked insertion (green underline in Word)."""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), datetime.utcnow().isoformat() + 'Z')
    ins.set(qn('w:id'), _next_revision_id())

    # Create run with text
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    # Preserve spaces
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    ins.append(r)

    paragraph._p.append(ins)


def _add_tracked_deletion(paragraph, text: str, author: str = "MCP Server") -> None:
    """Mark text as deleted in a paragraph (red strikethrough in Word)."""
    deletion = OxmlElement('w:del')
    deletion.set(qn('w:author'), author)
    deletion.set(qn('w:date'), datetime.utcnow().isoformat() + 'Z')
    deletion.set(qn('w:id'), _next_revision_id())

    # Create run with deleted text
    r = OxmlElement('w:r')
    # Use w:delText for deleted content
    t = OxmlElement('w:delText')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    deletion.append(r)

    paragraph._p.append(deletion)


def _set_cell_text_with_tracking(cell, text: str, author: str = "MCP Server", old_text: str = None) -> None:
    """Set cell text with track changes.

    This is the canonical way to set table cell content with change tracking.
    It ensures the changes are visible in Word's Track Changes view.

    Args:
        cell: A python-docx table cell
        text: The new text to set
        author: Author name for the tracked change
        old_text: Optional old text to mark as deleted (if replacing)
    """
    # Get or create the first paragraph in the cell
    if cell.paragraphs:
        para = cell.paragraphs[0]
        # Clear any existing runs
        for run in para.runs:
            run.text = ""
    else:
        para = cell.add_paragraph()

    # If there's old content to track as deleted
    if old_text and old_text.strip():
        _add_tracked_deletion(para, old_text, author)

    # Add new content as tracked insertion
    if text and str(text).strip():
        _add_tracked_insertion(para, str(text), author)


def _get_text_with_track_changes(element) -> str:
    """Extract text from an element including content inside track change elements.

    python-docx's .text property only gets text from regular runs, missing
    content inside <w:ins> (insertions) and <w:del> (deletions) elements.
    This function extracts all visible text (regular + insertions), excluding
    deleted text which shouldn't be shown as current content.

    Args:
        element: A python-docx element (paragraph, cell, etc.) with an underlying
                 XML element accessible via ._element or ._tc

    Returns:
        Combined text from regular runs and tracked insertions
    """
    # Get the underlying XML element
    if hasattr(element, '_tc'):  # Table cell
        xml_element = element._tc
    elif hasattr(element, '_p'):  # Paragraph
        xml_element = element._p
    elif hasattr(element, '_element'):  # Generic
        xml_element = element._element
    else:
        # Fallback to .text if we can't access XML
        return element.text if hasattr(element, 'text') else str(element)

    # Collect text from all w:t elements (regular text)
    # and w:t elements inside w:ins (tracked insertions)
    # but NOT w:delText elements (tracked deletions - should not be visible)
    text_parts = []

    # Find all text elements
    for t in xml_element.iter(qn('w:t')):
        # Check if this is inside a deletion (w:del) - skip those
        parent = t.getparent()
        in_deletion = False
        while parent is not None:
            if parent.tag == qn('w:del'):
                in_deletion = True
                break
            parent = parent.getparent()

        if not in_deletion and t.text:
            text_parts.append(t.text)

    return ''.join(text_parts)


def _get_cell_text(cell) -> str:
    """Get text from a table cell, including content in track change elements."""
    return _get_text_with_track_changes(cell).strip()


def _replace_with_track_changes(
    paragraph,
    old_text: str,
    new_text: str,
    author: str = "MCP Server"
) -> bool:
    """Replace text in a paragraph with track changes markup.

    Creates a deletion mark for old_text and insertion mark for new_text.
    Returns True if replacement was made.
    """
    full_text = _get_text_with_track_changes(paragraph)
    if old_text not in full_text:
        return False

    # Find position of old text
    idx = full_text.find(old_text)
    before = full_text[:idx]
    after = full_text[idx + len(old_text):]

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Rebuild paragraph with track changes
    if before:
        if paragraph.runs:
            paragraph.runs[0].text = before
        else:
            paragraph.add_run(before)

    # Add tracked deletion (old text)
    _add_tracked_deletion(paragraph, old_text, author)

    # Add tracked insertion (new text)
    _add_tracked_insertion(paragraph, new_text, author)

    # Add remaining text
    if after:
        paragraph.add_run(after)

    return True


class WordAdvancedTools:
    """MCP tool mixin for advanced Word document manipulation."""

    # Patterns for template instructions to strip
    INSTRUCTION_PATTERNS = [
        r'\[Template [Gg]uidance:.*?\]',
        r'\[Template [Gg]uidance:.*?(?=\])\]',
        r'\[PLEASE NOTE:.*?\]',
        r'Template Guidance:.*?(?=\n|$)',
    ]

    # Placeholder patterns
    PLACEHOLDER_PATTERNS = [
        (r'<Customer Name>', 'customer_name'),
        (r'<Customer short name if any>', 'customer_short_name'),
        (r'<Project Name>', 'project_name'),
        (r'<Microsoft OR Partner name>', 'provider_name'),
        (r'\[Customer\]', 'customer_name'),
        (r'\[insert WO number\]', 'work_order_number'),
        (r'\[insert language\]', 'language'),
    ]

    def tool_word_parse_sow_template(self, template_path: str) -> dict[str, Any]:
        """Parse a SOW template to extract its structure.

        Analyzes a Word document template to identify:
        - Document sections and headings
        - Tables and their purposes
        - Placeholder variables that need filling
        - Instructional text to be stripped

        Example:
            parse_sow_template(
                template_path=".github/skills/statement-of-work/templates/Agile.docx"
            )

        Args:
            template_path: Path to the .docx template file

        Returns:
            Dictionary with template structure analysis
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(template_path)
        if not path.exists():
            return {"error": f"Template not found: {template_path}"}

        try:
            doc = Document(template_path)
        except Exception as e:
            return {"error": f"Failed to open template: {e}"}

        # Extract structure
        sections = []
        tables_info = []
        placeholders_found = set()
        instructions_count = 0

        current_section = None

        for para in doc.paragraphs:
            text = _get_text_with_track_changes(para).strip()
            style_name = para.style.name if para.style else "Normal"

            # Track sections by heading styles
            if "Heading" in style_name and text:
                level = 1
                if "Heading 1" in style_name:
                    level = 1
                elif "Heading 2" in style_name:
                    level = 2
                elif "Heading 3" in style_name:
                    level = 3

                current_section = {
                    "title": text,
                    "level": level,
                    "style": style_name
                }
                sections.append(current_section)

            # Find placeholders
            for pattern, var_name in self.PLACEHOLDER_PATTERNS:
                if re.search(pattern, text, re.IGNORECASE):
                    placeholders_found.add(var_name)

            # Count instructional text
            for pattern in self.INSTRUCTION_PATTERNS:
                if re.search(pattern, text, re.IGNORECASE | re.DOTALL):
                    instructions_count += 1

        # Analyze tables
        for idx, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0

            # Get header row - use track-change-aware text extraction
            header = []
            if rows > 0:
                header = [_get_cell_text(cell)[:50] for cell in table.rows[0].cells]

            # Identify table purpose from headers
            purpose = self._identify_table_purpose(header)

            tables_info.append({
                "index": idx,
                "rows": rows,
                "columns": cols,
                "header": header,
                "purpose": purpose
            })

        return {
            "file": path.name,
            "sections": sections,
            "section_count": len(sections),
            "tables": tables_info,
            "table_count": len(tables_info),
            "placeholders": list(placeholders_found),
            "instruction_blocks": instructions_count,
            "message": f"Template has {len(sections)} sections, {len(tables_info)} tables, {len(placeholders_found)} placeholder types",
            "next_tools": ["word_copy_template", "word_list_tables", "word_get_section_guidance"]
        }

    def _identify_table_purpose(self, header: list[str]) -> str:
        """Identify the purpose of a table from its header row."""
        header_lower = " ".join(header).lower()

        if "business objective" in header_lower or "desired" in header_lower:
            return "business_objectives"
        elif "epic" in header_lower:
            return "epics"
        elif "out of scope" in header_lower or "areas out" in header_lower:
            return "out_of_scope"
        elif "technology" in header_lower and "requirement" in header_lower:
            return "technology_requirements"
        elif "environment" in header_lower and "requirement" in header_lower:
            return "environment_requirements"
        elif "staffing" in header_lower or "role" in header_lower:
            return "staffing"
        elif "term" in header_lower and ("acronym" in header_lower or "description" in header_lower):
            return "definitions"
        elif "test" in header_lower and "type" in header_lower:
            return "testing"
        elif "priority" in header_lower or "severity" in header_lower:
            return "defect_definitions"
        elif "phase" in header_lower or "activities" in header_lower:
            return "delivery_approach"
        else:
            return "unknown"

    def tool_word_generate_sow(
        self,
        template_path: str,
        output_path: str,
        sow_data: dict[str, Any]
    ) -> dict[str, Any]:
        """Generate a SOW document from a template and structured data.

        IMPORTANT: This tool REQUIRES a template document. It fills placeholders
        and tables but does NOT generate prose sections (Introduction, Executive
        Summary, etc.). After using this tool, use patch_section to add narrative
        content to key sections.

        Template-based workflow:
        1. Use copy_template to copy .github/skills/statement-of-work/templates/Agile.docx
        2. Use generate_sow to fill placeholders and tables
        3. Use patch_section to add Introduction, Business Context, etc.
        4. Use audit_completion to verify completeness
        5. Use cleanup_sow to remove template artifacts

        Takes a SOW template and fills it with actual engagement data,
        stripping instructional boilerplate and replacing placeholders.

        Example:
            generate_sow(
                template_path=".github/skills/statement-of-work/templates/Agile.docx",
                output_path="04. Artifacts/contoso-sow.docx",
                sow_data={
                    "customer_name": "Contoso",
                    "customer_short_name": "Contoso",
                    "project_name": "Cloud Migration",
                    "provider_name": "Microsoft",
                    "work_order_number": "WO-2026-001",
                    "language": "English",
                    "business_objectives": [
                        {"objective": "Migrate 15 apps to Azure", "activities": "Assessment, migration planning", "assumptions": "Apps are containerizable"}
                    ],
                    "epics": [
                        {"name": "Infrastructure Setup", "description": "Set up Azure landing zone", "assumptions": "Subscription available"}
                    ],
                    "out_of_scope": [
                        {"area": "Data migration", "description": "Historical data migration not included"}
                    ],
                    "technology_requirements": [
                        {"item": "Azure subscription", "version": "N/A", "ready_by": "Project start"}
                    ],
                    "assumptions": [
                        "Customer will provide access to existing systems",
                        "Dedicated product owner available full-time"
                    ]
                }
            )

        Args:
            template_path: Path to the .docx template (REQUIRED - use copy_template first)
            output_path: Path for the output .docx file
            sow_data: Dictionary containing SOW content

        Returns:
            Status dictionary with file path and next_tools suggestions
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(template_path)
        if not path.exists():
            return {"error": f"Template not found: {template_path}"}

        try:
            doc = Document(template_path)
        except Exception as e:
            return {"error": f"Failed to open template: {e}"}

        # Enable track revisions for all changes
        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0
        author = DEFAULT_AUTHOR

        # Step 1: Replace placeholders in all paragraphs
        replacements_made = 0
        for para in doc.paragraphs:
            original = para.text
            new_text = self._replace_placeholders(para.text, sow_data)
            if new_text != original:
                # Use track changes for placeholder replacement
                _replace_with_track_changes(para, original, new_text, author)
                replacements_made += 1

        # Step 2: Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original = para.text
                        new_text = self._replace_placeholders(para.text, sow_data)
                        if new_text != original:
                            _replace_with_track_changes(para, original, new_text, author)
                            replacements_made += 1

        # Step 3: Fill data tables with actual content (with track changes)
        tables_filled = self._fill_data_tables(doc, sow_data, author)

        # Step 4: Fill section content if provided
        sections_filled = 0
        if "sections" in sow_data and isinstance(sow_data["sections"], dict):
            sections_filled = self._fill_sections(doc, sow_data["sections"])

        # Step 5: Strip instructional text
        instructions_removed = self._strip_instructions(doc)

        # Save the document
        doc.save(output_path)

        # Step 6: Post-process at XML level to handle SDTs and ensure persistence
        # This neutralizes content control placeholders that can reset when opened in Word
        global_replacements = {}
        if "customer_name" in sow_data:
            global_replacements["<Customer Name>"] = sow_data["customer_name"]
        if "customer_short_name" in sow_data:
            global_replacements["<Customer Short Name>"] = sow_data["customer_short_name"]
        if "project_name" in sow_data:
            global_replacements["<Project Name>"] = sow_data["project_name"]
        if "provider_name" in sow_data:
            global_replacements["<Microsoft OR Partner name>"] = sow_data["provider_name"]
            global_replacements["Microsoft OR Partner name"] = sow_data["provider_name"]

        sdts_neutralized = 0
        if global_replacements:
            xml_result = self.tool_word_replace_global_variables(output_path, global_replacements)
            if xml_result.get("success"):
                sdts_neutralized = xml_result.get("sdts_neutralized", 0)
                replacements_made += xml_result.get("total_replacements", 0)

        return {
            "success": True,
            "file": output_path,
            "replacements": replacements_made,
            "tables_filled": tables_filled,
            "sections_filled": sections_filled,
            "instructions_removed": instructions_removed,
            "sdts_neutralized": sdts_neutralized,
            "message": f"Generated SOW with {replacements_made} placeholder replacements, {tables_filled} tables filled, {sections_filled} sections filled, {sdts_neutralized} content controls neutralized",
            "next_tools": ["word_list_tables", "word_get_section_guidance", "word_insert_table_row"]
        }

    def _replace_placeholders(self, text: str, sow_data: dict[str, Any]) -> str:
        """Replace all placeholder patterns with actual data."""
        result = text

        for pattern, var_name in self.PLACEHOLDER_PATTERNS:
            if var_name in sow_data:
                result = re.sub(pattern, str(sow_data[var_name]), result, flags=re.IGNORECASE)

        return result

    def _replace_paragraph_text(self, para, new_text: str):
        """Replace paragraph text while attempting to preserve formatting."""
        if para.runs:
            # Clear all runs except first, put all text in first run
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run.text = ""
            first_run.text = new_text
        else:
            para.text = new_text

    def _fill_data_tables(self, doc, sow_data: dict[str, Any], author: str = DEFAULT_AUTHOR) -> int:
        """Fill template tables with actual data."""
        tables_filled = 0

        for table in doc.tables:
            if not table.rows:
                continue

            # Get header to identify table purpose - use track-change-aware text extraction
            header = [_get_cell_text(cell).lower() for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)

            # Map data to tables based on purpose
            data_key_map = {
                "business_objectives": "business_objectives",
                "epics": "epics",
                "out_of_scope": "out_of_scope",
                "technology_requirements": "technology_requirements",
                "environment_requirements": "environment_requirements",
            }

            if purpose in data_key_map:
                data_key = data_key_map[purpose]
                if data_key in sow_data and sow_data[data_key]:
                    self._populate_table(table, sow_data[data_key], author)
                    tables_filled += 1

        return tables_filled

    def _populate_table(self, table, data: list[dict[str, Any]], author: str = DEFAULT_AUTHOR):
        """Populate a table with data rows, using track changes for visibility."""
        if not data or len(table.rows) < 1:
            return

        # Get column mapping from header
        header_cells = table.rows[0].cells
        col_map = {}
        for idx, cell in enumerate(header_cells):
            header_text = _get_cell_text(cell).lower()
            col_map[idx] = header_text

        # Capture old content before clearing rows
        old_content_by_row = []
        for row in table.rows[1:]:  # Skip header
            row_content = [_get_cell_text(cell) for cell in row.cells]
            old_content_by_row.append(row_content)

        # Clear existing data rows (keep header)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # Column name aliases for better matching
        column_aliases = {
            "epic": ["name", "epic", "title"],
            "desired business objectives": ["objective", "business_objective"],
            "areas out of scope": ["area", "out_of_scope"],
            "product and technology item": ["item", "technology"],
        }

        # Add new data rows with track changes
        for row_idx, item in enumerate(data):
            row = table.add_row()
            for idx, cell in enumerate(row.cells):
                # Try to match column to data key
                header_key = col_map.get(idx, "")
                value = ""

                # First try direct or partial match
                for key in item:
                    if key.lower() == header_key or key.lower() in header_key or header_key in key.lower():
                        value = str(item[key])
                        break
                else:
                    # Try alias matching
                    aliases = column_aliases.get(header_key, [])
                    for alias in aliases:
                        for key in item:
                            if key.lower() == alias:
                                value = str(item[key])
                                break
                        if value:
                            break

                # Get old value for this column if available (from first old row for deletions)
                old_value = None
                if row_idx == 0 and old_content_by_row:
                    # Show deletions from all old rows in first new row
                    old_values = [r[idx] for r in old_content_by_row if idx < len(r) and r[idx].strip()]
                    old_value = " | ".join(old_values) if old_values else None

                # Set cell content with track changes
                _set_cell_text_with_tracking(cell, value, author, old_value)

    def _strip_instructions(self, doc) -> int:
        """Remove instructional/guidance text from the document."""
        removed = 0

        # Remove paragraphs that are purely instructional
        paragraphs_to_remove = []
        for para in doc.paragraphs:
            text = _get_text_with_track_changes(para).strip()

            # Check if paragraph is instructional
            is_instruction = False

            # Pink/italic guidance text patterns
            if text.startswith("[Template") or text.startswith("[PLEASE NOTE") or "Template Guidance:" in text or text.startswith("Template guidance:"):
                is_instruction = True

            if is_instruction:
                paragraphs_to_remove.append(para)

        # Remove instructional paragraphs
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)
            removed += 1

        # Also clean inline instructions from remaining paragraphs
        for para in doc.paragraphs:
            original = _get_text_with_track_changes(para)
            cleaned = original
            for pattern in self.INSTRUCTION_PATTERNS:
                cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE | re.DOTALL)

            if cleaned != original:
                self._replace_paragraph_text(para, cleaned.strip())
                removed += 1

        return removed

    def tool_word_create_sow_from_markdown(
        self,
        output_path: str,
        markdown: str | None = None,
        template_path: str | None = None,
        markdown_file: str | None = None,
    ) -> dict[str, Any]:
        """Create a SOW document from inline markdown or markdown_file by filling a template.

        IMPORTANT: This tool requires a template to preserve document structure,
        formatting, and corporate styling. It extracts data from Markdown and
        uses generate_sow to fill the template.

        Workflow:
        1. Parse the Markdown to extract structured SOW data
        2. Load the template document
        3. Fill placeholders and tables with extracted data
        4. Save the result

        Example:
            create_sow_from_markdown(
                output_path="04. Artifacts/contoso-sow.docx",
                template_path=".github/skills/statement-of-work/templates/Agile.docx",
                markdown='''
# Contoso – Cloud Migration – Statement of Work

## 1. Engagement Overview

Customer: Contoso Ltd
Provider: Microsoft
Project: Cloud Migration Sprint 1

### 1.1 Business Objectives

| Objective | Activities | Assumptions |
|-----------|------------|-------------|
| Migrate 15 apps | Assessment, planning | Apps are containerizable |
'''
            )

        Args:
            output_path: Path for the output .docx file
            template_path: Path to the .docx template (REQUIRED)
            markdown: Markdown content of the SOW (inline)
            markdown_file: Optional path to a Markdown file. Use this for
                very large inputs to avoid MCP argument-size limits.

        Returns:
            Status dictionary with file path and extraction summary
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        markdown_text = markdown
        if markdown_file:
            resolved_md_path = resolve_office_path(markdown_file)
            md_path = Path(resolved_md_path)
            if not md_path.exists():
                return {"error": f"Markdown file not found: {markdown_file}"}
            if not md_path.is_file():
                return {"error": f"Markdown path is not a file: {markdown_file}"}
            markdown_text = md_path.read_text(encoding="utf-8")

        if markdown_text is None:
            return {"error": "Provide either 'markdown' or 'markdown_file'"}

        if not template_path:
            return {"error": "template_path is required. A template is required to preserve document structure."}

        template = Path(template_path)
        if not template.exists():
            return {"error": f"Template not found: {template_path}. A template is required to preserve document structure."}

        # Extract structured data from markdown
        sow_data = self._extract_data_from_markdown(markdown_text)

        # Use generate_sow to fill the template
        result = self.tool_word_generate_sow(
            template_path=template_path,
            output_path=output_path,
            sow_data=sow_data
        )

        if "error" in result:
            return result

        result["extracted_fields"] = list(sow_data.keys())
        result["message"] = f"Created SOW from template with {len(sow_data)} fields extracted from Markdown"

        return result

    def _extract_data_from_markdown(self, markdown: str) -> dict[str, Any]:
        """Extract structured SOW data from Markdown content."""
        data = {}
        lines = markdown.split('\n')

        # Extract header info
        for line in lines:
            line = line.strip()
            # Pattern: "Customer: Contoso Ltd"
            if line.lower().startswith('customer:'):
                data['customer_name'] = line.split(':', 1)[1].strip()
            elif line.lower().startswith('provider:'):
                data['provider_name'] = line.split(':', 1)[1].strip()
            elif line.lower().startswith('project:'):
                data['project_name'] = line.split(':', 1)[1].strip()
            elif line.lower().startswith('work order:') or line.lower().startswith('wo:'):
                data['work_order_number'] = line.split(':', 1)[1].strip()

        # Try to extract from title: "# Customer – Project – Statement of Work"
        for line in lines:
            if line.startswith('# ') and 'statement of work' in line.lower():
                parts = line[2:].split('–')
                if len(parts) >= 2:
                    if 'customer_name' not in data:
                        data['customer_name'] = parts[0].strip()
                    if 'project_name' not in data and len(parts) >= 2:
                        data['project_name'] = parts[1].strip()
                break

        # Extract tables
        data['business_objectives'] = self._extract_table_from_markdown(
            markdown, ['objective', 'activities', 'assumptions']
        )
        data['epics'] = self._extract_table_from_markdown(
            markdown, ['epic', 'description', 'assumptions']
        )
        data['technology_requirements'] = self._extract_table_from_markdown(
            markdown, ['product', 'technology', 'version', 'ready']
        )
        data['staffing'] = self._extract_table_from_markdown(
            markdown, ['role', 'count', 'responsibilities']
        )

        # Remove empty lists
        data = {k: v for k, v in data.items() if v}

        return data

    def _extract_table_from_markdown(self, markdown: str, header_keywords: list[str]) -> list[dict[str, str]]:
        """Extract a table from markdown based on header keywords."""
        lines = markdown.split('\n')
        result = []
        in_table = False
        headers = []

        for line in lines:
            stripped = line.strip()
            if not stripped.startswith('|') or not stripped.endswith('|'):
                if in_table and result:
                    break  # End of table
                in_table = False
                continue

            cells = [c.strip() for c in stripped[1:-1].split('|')]

            # Check if this is a separator row
            if all(set(c) <= set('-: ') for c in cells):
                continue

            # Check if header row matches our keywords
            if not in_table:
                header_text = ' '.join(cells).lower()
                if any(kw in header_text for kw in header_keywords):
                    headers = cells
                    in_table = True
                continue

            # Data row
            if in_table and headers:
                row = {}
                for i, cell in enumerate(cells):
                    if i < len(headers):
                        key = headers[i].lower().replace(' ', '_')
                        row[key] = cell
                if any(row.values()):
                    result.append(row)

        return result

    def tool_word_extract_sow_structure(self, file_path: str) -> dict[str, Any]:
        """Extract structured data from an existing SOW document.

        Parses a SOW document and extracts key information into
        a structured format that can be used to generate new documents.

        Example:
            extract_sow_structure(
                file_path="01. Inputs/existing-sow.docx"
            )

        Args:
            file_path: Path to the SOW document

        Returns:
            Dictionary with extracted SOW data
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        result = {
            "file": path.name,
            "sections": [],
            "tables": {},
            "extracted_data": {}
        }

        # Extract sections
        current_section = None
        for para in doc.paragraphs:
            text = _get_text_with_track_changes(para).strip()
            style = para.style.name if para.style else "Normal"

            if "Heading" in style and text:
                current_section = {
                    "title": text,
                    "style": style,
                    "content": []
                }
                result["sections"].append(current_section)
            elif current_section and text:
                current_section["content"].append(text)

        # Extract and categorize tables - use track-change-aware text extraction
        for _, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)

            rows_data = []
            for row in table.rows[1:]:  # Skip header
                row_data = [_get_cell_text(cell) for cell in row.cells]
                if any(row_data):  # Skip empty rows
                    rows_data.append(dict(zip(header, row_data)))

            if purpose != "unknown":
                result["tables"][purpose] = rows_data
                result["extracted_data"][purpose] = rows_data

        return result

    def tool_word_cleanup_sow(
        self,
        file_path: str,
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Clean a SOW document by removing all placeholder and instructional content.

        Removes:
        - Highlighted text (turquoise, yellow - template guidance markers)
        - Colored text (blue, red, purple - instructions)
        - Bracket placeholders that weren't filled: <...>, [Template Guidance: ...]
        - Instruction paragraphs containing guidance keywords

        This is the final step after generate_sow to ensure the document
        is presentation-ready with no visible template artifacts.
        All removals are tracked for auditability.

        Example:
            cleanup_sow(
                file_path="04. Artifacts/contoso-sow.docx",
                output_path="04. Artifacts/contoso-sow-final.docx"
            )

        Args:
            file_path: Path to the SOW document to clean
            output_path: Path for cleaned output (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Cleanup statistics
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = resolved_path

        try:
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            WD_COLOR_INDEX = None

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Always enable track revisions for auditability
        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        stats = {
            "highlighted_runs_cleared": 0,
            "colored_runs_cleared": 0,
            "bracket_placeholders_removed": 0,
            "instruction_paragraphs_cleared": 0,
            "table_cells_cleaned": 0
        }

        # Patterns for unfilled placeholders
        unfilled_patterns = [
            r'<[^>]+>',  # <Customer Name>, <n>, etc.
            r'\[Template Guidance:[^\]]*\]',  # [Template Guidance: ...]
            r'\[insert [^\]]+\]',  # [insert WO number]
            r'\[Customer[^\]]*\]',  # [Customer], [Customer short name...]
            r'\[\d+ days?\]',  # [3 days]
        ]

        # Instruction keywords that indicate guidance text
        instruction_keywords = [
            'template guidance', 'please note', 'assumption is that',
            'insert an introduction', 'insert a couple', 'should be',
            'must be defined', 'cannot be left blank', 'pdoc delivered',
            'example provided', 'examples provided', 'common roles provided',
            'update roles table', 'add or remove', 'delete this section',
            'use this section', 'fill in', 'do not go to'
        ]

        # Highlight colors to clear (turquoise=3, yellow=7 in WD_COLOR_INDEX)
        highlight_colors_to_clear = {3, 7}  # TURQUOISE, YELLOW

        # RGB colors to clear (instruction colors)
        instruction_rgb_patterns = [
            (15, 158, 213),   # Blue instructions
            (0, 78, 154),     # Dark blue
            (21, 96, 130),    # Teal
            (149, 0, 149),    # Purple
            (238, 0, 0),      # Red warnings
        ]

        def is_instruction_color(rgb):
            """Check if RGB color indicates instruction text."""
            if not rgb:
                return False
            for r, g, b in instruction_rgb_patterns:
                # Allow some tolerance
                if (abs(rgb[0] - r) < 20 and
                    abs(rgb[1] - g) < 20 and
                    abs(rgb[2] - b) < 20):
                    return True
            return False

        def clean_paragraph(para):
            """Clean a paragraph of placeholder and instruction content."""
            nonlocal stats
            text = _get_text_with_track_changes(para)

            # Check if entire paragraph is instructional
            text_lower = text.lower()
            # Check if this is mostly instruction (not just a mention)
            if (any(kw in text_lower for kw in instruction_keywords)
                    and len(text) > 50):  # Long paragraphs with these keywords
                if text.strip():
                    _add_tracked_deletion(para, text, author)
                for run in para.runs:
                    run.text = ""
                stats["instruction_paragraphs_cleared"] += 1
                return

            # Clean individual runs
            for run in para.runs:
                run_text = run.text
                should_clear = False
                clear_reason = None

                # Clear highlighted text
                if run.font.highlight_color:
                    hl = run.font.highlight_color
                    if WD_COLOR_INDEX:
                        if hl in (WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.YELLOW):
                            should_clear = True
                            clear_reason = "highlighted"
                    elif isinstance(hl, int) and hl in highlight_colors_to_clear:
                        should_clear = True
                        clear_reason = "highlighted"

                # Clear colored instruction text
                if not should_clear and run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    if is_instruction_color(rgb):
                        should_clear = True
                        clear_reason = "colored"

                if should_clear:
                    if run_text.strip():
                        _add_tracked_deletion(para, run_text, author)
                    run.text = ""
                    if run.font.highlight_color:
                        run.font.highlight_color = None
                    if clear_reason == "highlighted":
                        stats["highlighted_runs_cleared"] += 1
                    else:
                        stats["colored_runs_cleared"] += 1
                    continue

                # Remove unfilled bracket placeholders from run text
                original = run.text
                for pattern in unfilled_patterns:
                    run.text = re.sub(pattern, '', run.text)
                if run.text != original:
                    if original.strip():
                        # Track the removed placeholder text
                        removed_text = original.replace(run.text, '')
                        if removed_text.strip():
                            _add_tracked_deletion(para, removed_text, author)
                    stats["bracket_placeholders_removed"] += 1

        # Clean all paragraphs
        for para in doc.paragraphs:
            clean_paragraph(para)

        # Tables that should NEVER have content cleaned - they contain boilerplate
        BOILERPLATE_TABLE_PURPOSES = {"testing", "defect_definitions"}

        # Clean table cells, but SKIP boilerplate tables
        stats["boilerplate_tables_preserved"] = 0
        for table in doc.tables:
            # Identify table purpose from header row
            if len(table.rows) > 0:
                header = [cell.text.strip() for cell in table.rows[0].cells]
                purpose = self._identify_table_purpose(header)
            else:
                purpose = "unknown"

            # Skip cleaning boilerplate tables entirely
            if purpose in BOILERPLATE_TABLE_PURPOSES:
                stats["boilerplate_tables_preserved"] += 1
                continue  # Do not clean this table at all

            for row in table.rows:
                for cell in row.cells:
                    cell_cleaned = False
                    for para in cell.paragraphs:
                        original_text = para.text
                        clean_paragraph(para)
                        if para.text != original_text:
                            cell_cleaned = True
                    if cell_cleaned:
                        stats["table_cells_cleaned"] += 1

        # Save the cleaned document
        safe_save_docx(doc, output_path)

        total_cleaned = sum(stats.values())
        stats["success"] = True
        stats["file"] = output_path
        stats["total_items_cleaned"] = total_cleaned
        stats["message"] = f"Cleaned {total_cleaned} placeholder/instruction items"
        stats["track_changes"] = True
        stats["author"] = author
        stats["next_tools"] = ["word_audit_completion", "word_check_tracking"]

        return stats

    # =========================================================================
    # INTROSPECTION TOOLS - Read specific sections/tables from documents
    # =========================================================================

    def tool_word_get_section(self, file_path: str, section_title: str) -> dict[str, Any]:
        """Get content of a specific section from a SOW document.

        Retrieves a section by its heading title, including all paragraphs
        until the next heading of equal or higher level.

        Example:
            get_section(
                file_path="04. Artifacts/contoso-sow.docx",
                section_title="Executive Summary"
            )

            get_section(
                file_path="04. Artifacts/contoso-sow.docx",
                section_title="1.1 Customer desired business outcomes"
            )

        Args:
            file_path: Path to the .docx document
            section_title: Title or partial title of the section to retrieve

        Returns:
            Section content with title, level, and paragraphs
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Find the section
        found_section = None
        section_level = None
        section_content = []
        collecting = False

        for para in doc.paragraphs:
            style = para.style.name if para.style else "Normal"
            # Use track-change-aware text extraction
            text = _get_text_with_track_changes(para).strip()

            # Check if this is a heading
            is_heading = "Heading" in style
            if is_heading:
                # Extract heading level (e.g., "Heading 1" -> 1)
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 0

                # Check if this matches our target section
                if not collecting and section_title.lower() in text.lower():
                    found_section = text
                    section_level = level
                    collecting = True
                    continue

                # If we're collecting and hit same/higher level heading, stop
                if collecting and level <= section_level:
                    break

            # Collect content if we're in the target section
            if collecting and text:
                section_content.append(text)

        if not found_section:
            # List available sections for help
            available = []
            for para in doc.paragraphs:
                style = para.style.name if para.style else "Normal"
                if "Heading" in style:
                    heading_text = _get_text_with_track_changes(para).strip()
                    if heading_text:
                        available.append(heading_text[:60])

            return {
                "error": f"Section '{section_title}' not found",
                "available_sections": available[:20]
            }

        return {
            "title": found_section,
            "level": section_level,
            "content": section_content,
            "paragraph_count": len(section_content),
            "file": file_path,
            "next_tools": ["word_get_section_guidance", "word_patch_section", "word_list_sections"]
        }

    def tool_word_get_section_guidance(
        self, file_path: str, section_title: str
    ) -> dict[str, Any]:
        """Extract template guidance and instructions from a section.

        Reads a section and identifies all instructional content that tells
        you what to write there. This includes:
        - Template guidance markers: [Template Guidance: ...]
        - Highlighted text (turquoise, yellow) indicating placeholders
        - Colored text (blue, red, purple) indicating instructions
        - Insert instructions: [insert ...]

        Use this BEFORE patch_section to understand what content is expected.

        Example:
            get_section_guidance(
                file_path="04. Artifacts/contoso-sow.docx",
                section_title="Introduction"
            )

        Args:
            file_path: Path to the .docx document
            section_title: Title or partial title of the section

        Returns:
            Structured guidance including instructions, expected content hints,
            and placeholders to fill
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Instruction keywords that indicate guidance text
        instruction_keywords = [
            'template guidance', 'please note', 'assumption is that',
            'insert an introduction', 'insert a couple', 'should be',
            'must be defined', 'cannot be left blank', 'pdoc delivered',
            'example provided', 'examples provided', 'common roles provided',
            'update roles table', 'add or remove', 'delete this section',
            'use this section', 'fill in', 'do not go to', 'best practice'
        ]

        # Import WD_COLOR_INDEX for highlight detection
        try:
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            WD_COLOR_INDEX = None

        # Highlight colors indicating template guidance (turquoise=3, yellow=7)
        highlight_colors_to_check = {3, 7}

        # RGB colors indicating instruction text
        instruction_rgb_patterns = [
            (15, 158, 213),   # Blue instructions
            (0, 78, 154),     # Dark blue
            (21, 96, 130),    # Teal
            (149, 0, 149),    # Purple
            (238, 0, 0),      # Red warnings
        ]

        def is_instruction_color(rgb):
            if not rgb:
                return False
            for r, g, b in instruction_rgb_patterns:
                if (abs(rgb[0] - r) < 20 and
                    abs(rgb[1] - g) < 20 and
                    abs(rgb[2] - b) < 20):
                    return True
            return False

        # Find the section
        found_section = None
        section_level = None
        collecting = False

        guidance_items = []
        placeholders_found = []
        highlighted_text = []
        colored_instructions = []
        raw_content = []

        for para in doc.paragraphs:
            style = para.style.name if para.style else "Normal"
            text = _get_text_with_track_changes(para).strip()

            # Check if this is a heading
            is_heading = "Heading" in style
            if is_heading:
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 0

                if not collecting and section_title.lower() in text.lower():
                    found_section = text
                    section_level = level
                    collecting = True
                    continue

                if collecting and level <= section_level:
                    break

            if not collecting:
                continue

            # Collect raw content
            if text:
                raw_content.append(text)

            # Extract template guidance markers
            guidance_matches = re.findall(
                r'\[Template [Gg]uidance:[^\]]*\]', text
            )
            guidance_items.extend(guidance_matches)

            # Extract insert instructions
            insert_matches = re.findall(
                r'\[insert [^\]]+\]', text, re.IGNORECASE
            )
            guidance_items.extend(insert_matches)

            # Extract placeholders
            placeholder_matches = re.findall(r'<[^>]+>', text)
            placeholders_found.extend(placeholder_matches)

            # Check for instruction keywords in paragraph
            text_lower = text.lower()
            for keyword in instruction_keywords:
                if keyword in text_lower and text not in guidance_items:
                    # Extract the sentence containing the keyword
                    guidance_items.append(text[:200])
                    break

            # Analyze runs for highlighted/colored text
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                # Check for highlighted text
                if run.font.highlight_color:
                    hl = run.font.highlight_color
                    is_highlight = False
                    if WD_COLOR_INDEX:
                        if hl in (WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.YELLOW):
                            is_highlight = True
                    elif isinstance(hl, int) and hl in highlight_colors_to_check:
                        is_highlight = True

                    if is_highlight and run_text not in highlighted_text:
                        highlighted_text.append(run_text)

                # Check for colored instruction text
                if (run.font.color and run.font.color.rgb
                        and is_instruction_color(run.font.color.rgb)
                        and run_text not in colored_instructions):
                    colored_instructions.append(run_text)

        if not found_section:
            available = []
            for para in doc.paragraphs:
                style = para.style.name if para.style else "Normal"
                if "Heading" in style:
                    heading_text = _get_text_with_track_changes(para).strip()
                    if heading_text:
                        available.append(heading_text[:60])

            return {
                "error": f"Section '{section_title}' not found",
                "available_sections": available[:20]
            }

        # Deduplicate
        guidance_items = list(dict.fromkeys(guidance_items))
        placeholders_found = list(dict.fromkeys(placeholders_found))

        # Build structured guidance
        guidance = {
            "section": found_section,
            "level": section_level,
            "file": file_path,
            "template_instructions": guidance_items,
            "placeholders": placeholders_found,
            "highlighted_guidance": highlighted_text,
            "colored_instructions": colored_instructions[:10],  # Limit
            "raw_content_preview": raw_content[:5],  # First 5 paragraphs
            "summary": {
                "instruction_count": len(guidance_items),
                "placeholder_count": len(placeholders_found),
                "highlighted_count": len(highlighted_text),
                "colored_count": len(colored_instructions)
            }
        }

        # Add interpretation hints
        hints = []
        if guidance_items:
            hints.append("Review template_instructions for what content is expected")
        if placeholders_found:
            hints.append(f"Replace these placeholders: {', '.join(placeholders_found[:5])}")
        if highlighted_text:
            hints.append("Highlighted text indicates template guidance to replace")
        if not guidance_items and not highlighted_text:
            hints.append("Section appears to have standard boilerplate - may only need placeholder replacement")

        guidance["writing_hints"] = hints
        guidance["next_tools"] = ["word_insert_table_row", "word_patch_table_row", "word_patch_section", "word_list_tables"]

        return guidance

    def tool_word_get_table(self, file_path: str, table_identifier: str) -> dict[str, Any]:
        """Get content of a specific table from a SOW document.

        Retrieves a table by index (0-based) or by identifying header content.
        Returns the table structure with headers and all rows as dictionaries.

        Example:
            get_table(
                file_path="04. Artifacts/contoso-sow.docx",
                table_identifier="0"  # First table
            )

            get_table(
                file_path="04. Artifacts/contoso-sow.docx",
                table_identifier="business_objectives"  # By purpose
            )

            get_table(
                file_path="04. Artifacts/contoso-sow.docx",
                table_identifier="Role"  # By header content
            )

        Args:
            file_path: Path to the .docx document
            table_identifier: Table index (as string) or header keyword/purpose

        Returns:
            Table content with headers and rows
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Try to parse as index
        target_index = None
        with contextlib.suppress(ValueError):
            target_index = int(table_identifier)

        # Find the table
        found_table = None
        found_index = None

        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            # Get header - use track-change-aware text extraction
            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)
            header_text = " ".join(header).lower()

            # Match by index
            if target_index is not None and idx == target_index:
                found_table = table
                found_index = idx
                break

            # Match by purpose
            if target_index is None and table_identifier.lower() == purpose:
                found_table = table
                found_index = idx
                break

            # Match by header content
            if target_index is None and table_identifier.lower() in header_text:
                found_table = table
                found_index = idx
                break

        if not found_table:
            # List available tables
            available = []
            for idx, table in enumerate(doc.tables):
                if table.rows:
                    header = [_get_cell_text(cell)[:20] for cell in table.rows[0].cells]
                    purpose = self._identify_table_purpose(header)
                    available.append({
                        "index": idx,
                        "purpose": purpose,
                        "header_preview": header[:3]
                    })

            return {
                "error": f"Table '{table_identifier}' not found",
                "available_tables": available
            }

        # Extract table content - use track-change-aware text extraction
        header = [_get_cell_text(cell) for cell in found_table.rows[0].cells]
        rows = []
        for row in found_table.rows[1:]:
            row_data = {}
            for i, cell in enumerate(row.cells):
                if i < len(header):
                    row_data[header[i]] = _get_cell_text(cell)
            if any(row_data.values()):  # Skip empty rows
                rows.append(row_data)

        return {
            "index": found_index,
            "purpose": self._identify_table_purpose(header),
            "header": header,
            "rows": rows,
            "row_count": len(rows),
            "file": file_path
        }

    def tool_word_list_sections(self, file_path: str) -> dict[str, Any]:
        """List all sections (headings) in a SOW document.

        Returns a hierarchical view of all document sections with their
        heading levels, useful for understanding document structure.

        Example:
            list_sections(file_path="04. Artifacts/contoso-sow.docx")

        Args:
            file_path: Path to the .docx document

        Returns:
            List of sections with titles and levels
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        sections = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else "Normal"
            text = _get_text_with_track_changes(para).strip()

            if "Heading" in style and text:
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 0

                sections.append({
                    "title": text,
                    "level": level,
                    "style": style
                })

        return {
            "sections": sections,
            "count": len(sections),
            "file": file_path,
            "next_tools": ["word_get_section_guidance", "word_get_section", "word_patch_section"]
        }

    def tool_word_list_tables(self, file_path: str) -> dict[str, Any]:
        """List all tables in a SOW document with their purposes.

        Returns metadata about each table including identified purpose,
        headers, and row counts.

        Example:
            list_tables(file_path="04. Artifacts/contoso-sow.docx")

        Args:
            file_path: Path to the .docx document

        Returns:
            List of tables with metadata
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        tables = []
        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            # Use track-change-aware text extraction
            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)

            tables.append({
                "index": idx,
                "purpose": purpose,
                "header": header,
                "rows": len(table.rows) - 1,  # Exclude header
                "columns": len(header)
            })

        return {
            "tables": tables,
            "count": len(tables),
            "file": file_path
        }

    # =========================================================================
    # PATCHING TOOLS - Update specific sections/tables in documents
    # =========================================================================

    def tool_word_patch_section(
        self,
        file_path: str,
        section_title: str,
        new_content: list[str],
        output_path: str | None = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Update the content of a specific section in a SOW document.

        Replaces all paragraphs in a section (between its heading and the
        next heading of equal or higher level) with new content.
        All changes are tracked for auditability.

        Example:
            patch_section(
                file_path="04. Artifacts/contoso-sow.docx",
                section_title="Executive Summary",
                new_content=[
                    "Contoso Ltd engages Microsoft for a strategic cloud migration initiative.",
                    "This 12-week engagement will migrate 15 legacy applications to Azure.",
                    "The project follows an Agile sprint-based delivery approach."
                ]
            )

        Args:
            file_path: Path to the .docx document
            section_title: Title of the section to update
            new_content: List of paragraphs to replace section content
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with paragraphs replaced count
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = resolved_path

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Enable track revisions
        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        # Find the section boundaries
        section_start = None
        section_end = None
        section_level = None

        paragraphs = list(doc.paragraphs)

        for i, para in enumerate(paragraphs):
            style = para.style.name if para.style else "Normal"
            text = _get_text_with_track_changes(para).strip()

            is_heading = "Heading" in style
            if is_heading:
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 0

                # Found our target section
                if section_start is None and section_title.lower() in text.lower():
                    section_start = i
                    section_level = level
                    continue

                # Found the end (next heading of same/higher level)
                if section_start is not None and level <= section_level:
                    section_end = i
                    break

        if section_start is None:
            return {"error": f"Section '{section_title}' not found"}

        if section_end is None:
            section_end = len(paragraphs)

        # Clear existing content (but keep the heading)
        paragraphs_removed = 0
        for i in range(section_start + 1, section_end):
            old_text = _get_text_with_track_changes(paragraphs[i]).strip()
            if old_text:
                # Mark old content as deleted
                for run in paragraphs[i].runs:
                    run.text = ""
                _add_tracked_deletion(paragraphs[i], old_text, author)
                paragraphs_removed += 1

        # Insert new content after the heading
        heading_para = paragraphs[section_start]
        heading_element = heading_para._element
        parent = heading_element.getparent()

        # Create new paragraphs
        paragraphs_added = 0
        insert_position = list(parent).index(heading_element) + 1

        for content_text in new_content:
            new_para = doc.add_paragraph()
            # Add as tracked insertion
            _add_tracked_insertion(new_para, content_text, author)
            # Move to correct position
            new_element = new_para._element
            parent.remove(new_element)
            parent.insert(insert_position, new_element)
            insert_position += 1
            paragraphs_added += 1

        safe_save_docx(doc, output_path)

        return {
            "success": True,
            "file": output_path,
            "section": section_title,
            "paragraphs_cleared": paragraphs_removed,
            "paragraphs_added": paragraphs_added,
            "track_changes": True,
            "author": author,
            "message": f"Updated section '{section_title}' with {paragraphs_added} paragraphs (tracked by '{author}')",
            "next_tools": ["word_add_comment", "word_get_section_guidance", "word_list_tables", "word_insert_table_row"]
        }

    def tool_word_insert_table_row(
        self,
        file_path: str,
        table_identifier: str,
        row_data: dict[str, str],
        position: str = "end",
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Insert a single row into a table at a specific position.

        Adds exactly one row without affecting existing content. This is the
        preferred tool for adding data to tables as it preserves boilerplate
        rows that templates may contain.

        Example:
            insert_table_row(
                file_path="04. Artifacts/contoso-sow.docx",
                table_identifier="definitions",
                row_data={"Term / acronym": "WISMO", "Description": "Where Is My Order"},
                position="end"  # or "start" or row index like "2"
            )

        Args:
            file_path: Path to the .docx document
            table_identifier: Table index (as string) or purpose/header keyword
            row_data: Dictionary with column values (keys must match headers)
            position: "start" (after header), "end", or row index (as string)
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with new row position
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = resolved_path

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        # Find the table
        target_index = None
        with contextlib.suppress(ValueError):
            target_index = int(table_identifier)

        found_table = None
        found_index = None

        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)
            header_text = " ".join(header).lower()

            if target_index is not None and idx == target_index:
                found_table = table
                found_index = idx
                break
            if target_index is None and table_identifier.lower() == purpose:
                found_table = table
                found_index = idx
                break
            if target_index is None and table_identifier.lower() in header_text:
                found_table = table
                found_index = idx
                break

        if not found_table:
            return {"error": f"Table '{table_identifier}' not found"}

        header = [_get_cell_text(cell) for cell in found_table.rows[0].cells]

        # Determine insert position
        if position == "start":
            insert_idx = 1  # After header
        elif position == "end":
            insert_idx = len(found_table.rows)
        else:
            try:
                insert_idx = int(position)
                if insert_idx < 1:
                    insert_idx = 1
                if insert_idx > len(found_table.rows):
                    insert_idx = len(found_table.rows)
            except ValueError:
                insert_idx = len(found_table.rows)

        # Add the row at end first, then move if needed
        new_row = found_table.add_row()

        for i, col_name in enumerate(header):
            if i < len(new_row.cells):
                value = row_data.get(col_name, "")
                if not value:
                    for key, val in row_data.items():
                        if key.lower() == col_name.lower():
                            value = val
                            break
                _set_cell_text_with_tracking(new_row.cells[i], value, author, None)

        # Move row to correct position if not at end
        if insert_idx < len(found_table.rows) - 1:
            tbl = found_table._tbl
            tr = new_row._tr
            tbl.remove(tr)
            tbl.insert(insert_idx, tr)

        safe_save_docx(doc, output_path)

        return {
            "success": True,
            "file": output_path,
            "table_index": found_index,
            "table_purpose": self._identify_table_purpose(header),
            "row_inserted_at": insert_idx,
            "total_rows": len(found_table.rows),
            "track_changes": True,
            "author": author,
            "message": f"Inserted row at position {insert_idx} in table",
            "next_tools": ["word_add_comment", "word_insert_table_row", "word_patch_table_row"]
        }

    def tool_word_patch_table_row(
        self,
        file_path: str,
        table_identifier: str,
        row_index: int,
        updates: dict[str, str],
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Update specific cells in an existing table row.

        Modifies only the specified columns in a single row, preserving
        all other content. Row index is 1-based (row 1 is first data row,
        row 0 is header).

        Example:
            patch_table_row(
                file_path="04. Artifacts/contoso-sow.docx",
                table_identifier="staffing",
                row_index=3,
                updates={"Responsibilities / notes": "Updated role description"}
            )

        Args:
            file_path: Path to the .docx document
            table_identifier: Table index (as string) or purpose/header keyword
            row_index: 1-based row index (1 = first data row)
            updates: Dictionary of column names to new values
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with cells updated
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        # Find the table
        target_index = None
        with contextlib.suppress(ValueError):
            target_index = int(table_identifier)

        found_table = None
        found_index = None

        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)
            header_text = " ".join(header).lower()

            if target_index is not None and idx == target_index:
                found_table = table
                found_index = idx
                break
            if target_index is None and table_identifier.lower() == purpose:
                found_table = table
                found_index = idx
                break
            if target_index is None and table_identifier.lower() in header_text:
                found_table = table
                found_index = idx
                break

        if not found_table:
            return {"error": f"Table '{table_identifier}' not found"}

        header = [_get_cell_text(cell) for cell in found_table.rows[0].cells]

        # Validate row index
        if row_index < 1 or row_index >= len(found_table.rows):
            return {
                "error": f"Row index {row_index} out of range. Table has {len(found_table.rows) - 1} data rows (1-{len(found_table.rows) - 1})"
            }

        target_row = found_table.rows[row_index]
        cells_updated = 0

        for col_name, new_value in updates.items():
            # Find column index
            col_idx = None
            for i, h in enumerate(header):
                if h == col_name or h.lower() == col_name.lower():
                    col_idx = i
                    break

            if col_idx is None:
                continue

            if col_idx < len(target_row.cells):
                old_value = _get_cell_text(target_row.cells[col_idx])
                _set_cell_text_with_tracking(
                    target_row.cells[col_idx],
                    new_value,
                    author,
                    old_value
                )
                cells_updated += 1

        doc.save(output_path)

        return {
            "success": True,
            "file": output_path,
            "table_index": found_index,
            "table_purpose": self._identify_table_purpose(header),
            "row_index": row_index,
            "cells_updated": cells_updated,
            "track_changes": True,
            "author": author,
            "message": f"Updated {cells_updated} cells in row {row_index}",
            "next_tools": ["word_add_comment", "word_patch_table_row", "word_insert_table_row"]
        }

    def tool_word_duplicate_table_structure(
        self,
        file_path: str,
        source_table: str,
        insert_after_paragraph: str = None,
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Create a new empty table with the same structure as an existing one.

        Duplicates the column structure (headers, widths, formatting) of a
        source table without copying its data rows. Useful for creating
        additional tables that match template styling.

        Example:
            duplicate_table_structure(
                file_path="04. Artifacts/contoso-sow.docx",
                source_table="staffing",
                insert_after_paragraph="Customer Team Structure"
            )

        Args:
            file_path: Path to the .docx document
            source_table: Table index (as string) or purpose/header keyword
            insert_after_paragraph: Text to find - new table inserted after this
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with new table index
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        _enable_track_revisions(doc)

        # Find source table
        target_index = None
        with contextlib.suppress(ValueError):
            target_index = int(source_table)

        found_table = None

        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)
            header_text = " ".join(header).lower()

            if target_index is not None and idx == target_index:
                found_table = table
                break
            if target_index is None and source_table.lower() == purpose:
                found_table = table
                break
            if target_index is None and source_table.lower() in header_text:
                found_table = table
                break

        if not found_table:
            return {"error": f"Source table '{source_table}' not found"}

        # Get source structure
        source_header = [_get_cell_text(cell) for cell in found_table.rows[0].cells]
        num_cols = len(source_header)

        # Find insert position
        insert_para = None
        if insert_after_paragraph:
            for para in doc.paragraphs:
                if insert_after_paragraph.lower() in _get_text_with_track_changes(para).lower():
                    insert_para = para
                    break

        # Create new table at end of document
        new_table = doc.add_table(rows=1, cols=num_cols)

        # Copy header
        for i, header_text in enumerate(source_header):
            _set_cell_text_with_tracking(new_table.rows[0].cells[i], header_text, author, None)

        # Try to copy table style
        if found_table.style:
            new_table.style = found_table.style

        # Move table to correct position if specified
        new_table_tbl = new_table._tbl  # Store reference before save
        if insert_para:
            para_elem = insert_para._element
            para_elem.addnext(new_table_tbl)

        doc.save(output_path)

        # Recalculate correct table index by finding our table in the saved document
        # Re-open to get accurate index after XML manipulation
        try:
            doc_reloaded = Document(output_path)
            new_table_index = -1
            for idx, tbl in enumerate(doc_reloaded.tables):
                if tbl.rows and len(tbl.rows[0].cells) == num_cols:
                    # Check if header matches
                    tbl_header = [_get_cell_text(cell) for cell in tbl.rows[0].cells]
                    if tbl_header == source_header and len(tbl.rows) == 1:
                        # This is likely our newly created table (only header row)
                        new_table_index = idx
                        break
            if new_table_index == -1:
                new_table_index = len(doc_reloaded.tables) - 1  # fallback
        except Exception:
            new_table_index = len(doc.tables) - 1  # fallback to original method

        return {
            "success": True,
            "file": output_path,
            "new_table_index": new_table_index,
            "columns": source_header,
            "column_count": num_cols,
            "inserted_after": insert_after_paragraph if insert_para else None,
            "track_changes": True,
            "author": author,
            "message": f"Created new table with {num_cols} columns matching source structure" +
                      (f" after paragraph containing '{insert_after_paragraph}'" if insert_para else " at end of document")
        }

    def tool_word_create_new_table(
        self,
        file_path: str,
        headers: list[str],
        rows: list[dict[str, str]] = None,
        insert_after_section: str = None,
        insert_before_section: str = None,
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Create a new table with custom headers and optional initial rows.

        Creates a table at a specific location in the document, either after
        or before a named section heading. This is the preferred method for
        adding structured content like timelines, milestones, or custom data.

        Unlike duplicate_table_structure, this tool:
        - Creates tables with custom headers (not copied from existing tables)
        - Properly positions the table relative to section headings
        - Can pre-populate with initial rows
        - Returns accurate table index after insertion

        Example:
            create_new_table(
                file_path="04. Artifacts/contoso-sow.docx",
                headers=["Phase", "Start Date", "End Date", "Key Deliverables"],
                rows=[
                    {"Phase": "Discovery", "Start Date": "Week 1", "End Date": "Week 4", "Key Deliverables": "Assessment report"},
                    {"Phase": "Build", "Start Date": "Week 5", "End Date": "Week 16", "Key Deliverables": "MVP release"}
                ],
                insert_after_section="Timeline"
            )

        Args:
            file_path: Path to the .docx document
            headers: List of column header names for the new table
            rows: Optional list of row dictionaries with column values
            insert_after_section: Section heading after which to insert the table
            insert_before_section: Section heading before which to insert the table
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes

        Returns:
            Status with new table index and column information
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if not headers or len(headers) == 0:
            return {"error": "Headers list is required and cannot be empty"}

        if insert_after_section and insert_before_section:
            return {"error": "Specify either insert_after_section or insert_before_section, not both"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        _enable_track_revisions(doc)

        # Find the target section heading
        target_section = insert_after_section or insert_before_section
        insert_after = insert_after_section is not None
        target_para = None
        target_level = None

        if target_section:
            paragraphs = list(doc.paragraphs)
            for i, para in enumerate(paragraphs):
                style = para.style.name if para.style else "Normal"
                text = _get_text_with_track_changes(para).strip()

                if "Heading" in style and target_section.lower() in text.lower():
                    target_para = para
                    try:
                        target_level = int(style.split()[-1])
                    except (ValueError, IndexError):
                        target_level = 1

                    # If inserting before, we use this paragraph
                    # If inserting after, find the end of the section
                    if insert_after:
                        # Find the next heading of same/higher level or end
                        for j in range(i + 1, len(paragraphs)):
                            next_style = paragraphs[j].style.name if paragraphs[j].style else "Normal"
                            if "Heading" in next_style:
                                try:
                                    next_level = int(next_style.split()[-1])
                                    if next_level <= target_level:
                                        # Insert before this next section
                                        target_para = paragraphs[j]
                                        insert_after = False  # We found where to insert before
                                        break
                                except (ValueError, IndexError):
                                    pass
                            # Also check for the last paragraph before next heading
                            elif j == len(paragraphs) - 1:
                                target_para = paragraphs[j]
                    break

            if target_para is None:
                return {"error": f"Section '{target_section}' not found in document"}

        # Create the new table
        num_cols = len(headers)
        num_rows = 1 + (len(rows) if rows else 0)
        new_table = doc.add_table(rows=num_rows, cols=num_cols)

        # Apply a default table style (Table Grid is commonly available)
        with contextlib.suppress(Exception):
            new_table.style = 'Table Grid'

        # Set headers with track changes
        for i, header_text in enumerate(headers):
            _set_cell_text_with_tracking(new_table.rows[0].cells[i], header_text, author, None)

        # Populate rows if provided
        rows_added = 0
        if rows:
            for row_idx, row_data in enumerate(rows):
                table_row = new_table.rows[row_idx + 1]  # +1 to skip header
                for col_idx, header in enumerate(headers):
                    cell_value = row_data.get(header, "")
                    if cell_value:
                        _set_cell_text_with_tracking(table_row.cells[col_idx], cell_value, author, None)
                rows_added += 1

        # Move table to correct position
        new_table_tbl = new_table._tbl
        insertion_point = None

        if target_para:
            target_element = target_para._element

            if insert_after:
                # Insert after the target paragraph
                target_element.addnext(new_table_tbl)
                insertion_point = "after"
            else:
                # Insert before the target paragraph
                target_element.addprevious(new_table_tbl)
                insertion_point = "before"

        doc.save(output_path)

        # Recalculate correct table index
        try:
            doc_reloaded = Document(output_path)
            new_table_index = -1
            for idx, tbl in enumerate(doc_reloaded.tables):
                if tbl.rows and len(tbl.rows[0].cells) == num_cols:
                    tbl_header = [_get_cell_text(cell) for cell in tbl.rows[0].cells]
                    if tbl_header == headers:
                        new_table_index = idx
                        break
            if new_table_index == -1:
                new_table_index = len(doc_reloaded.tables) - 1
        except Exception:
            new_table_index = len(doc.tables) - 1

        position_msg = ""
        if insert_after_section:
            position_msg = f" after section '{insert_after_section}'"
        elif insert_before_section:
            position_msg = f" before section '{insert_before_section}'"
        else:
            position_msg = " at end of document"

        return {
            "success": True,
            "file": output_path,
            "new_table_index": new_table_index,
            "headers": headers,
            "column_count": num_cols,
            "rows_added": rows_added,
            "position": insertion_point or "end",
            "target_section": target_section,
            "track_changes": True,
            "author": author,
            "message": f"Created new table with {num_cols} columns and {rows_added} data rows{position_msg}",
            "next_tools": ["word_insert_table_row", "word_patch_table_row", "word_add_comment"]
        }

    def tool_word_analyze_template_formatting(
        self,
        file_path: str
    ) -> dict[str, Any]:
        """Analyze a template to identify formatting patterns for different content types.

        Examines the document to categorize content by formatting:
        - **Boilerplate (preserve)**: Standard methodology, governance, process text
        - **Placeholders (replace)**: `<...>`, `[...]` markers needing customer data
        - **Guidance (remove)**: Highlighted/colored instruction text for authors
        - **Tables by type**: Which tables are boilerplate vs. engagement-specific

        Use this before editing to understand what should be preserved vs. modified.

        Example:
            analyze_template_formatting(
                file_path=".github/skills/statement-of-work/templates/Agile.docx"
            )

        Args:
            file_path: Path to the .docx template

        Returns:
            Comprehensive formatting analysis with handling recommendations
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        try:
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            WD_COLOR_INDEX = None

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        analysis = {
            "file": path.name,
            "content_types": {
                "boilerplate_sections": [],
                "placeholder_sections": [],
                "guidance_to_remove": [],
            },
            "table_handling": [],
            "formatting_patterns": {
                "highlight_colors": {},
                "text_colors": {},
                "placeholder_patterns": [],
            },
            "recommendations": {}
        }

        # Patterns
        placeholder_patterns = [
            (r'<[^>]+>', 'angle_bracket'),
            (r'\[insert [^\]]+\]', 'insert_instruction'),
            (r'\[Template [Gg]uidance:[^\]]*\]', 'template_guidance'),
            (r'\[TBD\]', 'tbd'),
            (r'\[Customer[^\]]*\]', 'customer_ref'),
        ]

        # Guidance keywords indicating boilerplate methodology
        boilerplate_keywords = [
            'scrum', 'sprint', 'agile', 'methodology', 'governance',
            'escalation', 'change management', 'acceptance criteria',
            'defect', 'severity', 'priority', 'testing types'
        ]

        # Guidance keywords indicating author instructions
        instruction_keywords = [
            'template guidance', 'please note', 'insert an introduction',
            'insert a couple', 'delete this section', 'add or remove',
            'example provided', 'fill in', 'update roles'
        ]

        # Analyze paragraphs
        for para in doc.paragraphs:
            text = _get_text_with_track_changes(para).strip()
            style = para.style.name if para.style else "Normal"
            text_lower = text.lower()

            # Check for placeholders
            has_placeholder = False
            for pattern, ptype in placeholder_patterns:
                if re.search(pattern, text):
                    has_placeholder = True
                    if ptype not in analysis["formatting_patterns"]["placeholder_patterns"]:
                        analysis["formatting_patterns"]["placeholder_patterns"].append(ptype)

            # Check for boilerplate methodology
            is_boilerplate = any(kw in text_lower for kw in boilerplate_keywords)

            # Check for author instructions
            is_instruction = any(kw in text_lower for kw in instruction_keywords)

            # Analyze run formatting
            for run in para.runs:
                # Track highlight colors
                if run.font.highlight_color:
                    hl = run.font.highlight_color
                    hl_name = str(hl)
                    if WD_COLOR_INDEX:
                        if hl == WD_COLOR_INDEX.TURQUOISE:
                            hl_name = "TURQUOISE (template guidance)"
                        elif hl == WD_COLOR_INDEX.YELLOW:
                            hl_name = "YELLOW (placeholder)"
                    analysis["formatting_patterns"]["highlight_colors"][hl_name] = \
                        analysis["formatting_patterns"]["highlight_colors"].get(hl_name, 0) + 1

                # Track text colors
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    rgb_str = f"RGB({rgb[0]},{rgb[1]},{rgb[2]})"
                    analysis["formatting_patterns"]["text_colors"][rgb_str] = \
                        analysis["formatting_patterns"]["text_colors"].get(rgb_str, 0) + 1

            # Categorize section
            if "Heading" in style and text:
                if is_instruction:
                    analysis["content_types"]["guidance_to_remove"].append({
                        "section": text[:60],
                        "action": "Remove guidance text, may need content"
                    })
                elif has_placeholder and not is_boilerplate:
                    analysis["content_types"]["placeholder_sections"].append({
                        "section": text[:60],
                        "action": "Fill with engagement-specific content"
                    })
                elif is_boilerplate:
                    analysis["content_types"]["boilerplate_sections"].append({
                        "section": text[:60],
                        "action": "Preserve - standard methodology"
                    })

        # Analyze tables
        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)
            row_count = len(table.rows) - 1

            # Determine handling recommendation
            if purpose in ['testing', 'defect_definitions']:
                action = "PRESERVE - standard definitions, do not modify"
                mode = "preserve"
            elif purpose == 'definitions':
                action = "APPEND ONLY - keep standard terms, add project-specific"
                mode = "append"
            elif purpose in ['staffing']:
                action = "MODIFY - keep standard roles, update counts, add specific roles"
                mode = "modify"
            elif purpose in ['business_objectives', 'epics', 'out_of_scope',
                            'delivery_approach', 'technology_requirements',
                            'environment_requirements']:
                action = "REPLACE - engagement-specific content"
                mode = "replace"
            else:
                action = "REVIEW - determine if boilerplate or engagement-specific"
                mode = "review"

            analysis["table_handling"].append({
                "index": idx,
                "purpose": purpose,
                "header": header[:3],
                "existing_rows": row_count,
                "action": action,
                "mode": mode
            })

        # Build recommendations
        analysis["recommendations"] = {
            "preserve_these_sections": [
                s["section"] for s in analysis["content_types"]["boilerplate_sections"]
            ][:10],
            "fill_these_sections": [
                s["section"] for s in analysis["content_types"]["placeholder_sections"]
            ][:10],
            "tables_to_replace": [
                t["purpose"] for t in analysis["table_handling"] if t["mode"] == "replace"
            ],
            "tables_to_append": [
                t["purpose"] for t in analysis["table_handling"] if t["mode"] == "append"
            ],
            "tables_to_preserve": [
                t["purpose"] for t in analysis["table_handling"] if t["mode"] == "preserve"
            ],
            "formatting_to_remove": list(analysis["formatting_patterns"]["highlight_colors"].keys()),
        }

        analysis["summary"] = {
            "boilerplate_sections": len(analysis["content_types"]["boilerplate_sections"]),
            "placeholder_sections": len(analysis["content_types"]["placeholder_sections"]),
            "tables_total": len(analysis["table_handling"]),
            "tables_to_preserve": len([t for t in analysis["table_handling"] if t["mode"] == "preserve"]),
            "tables_to_replace": len([t for t in analysis["table_handling"] if t["mode"] == "replace"]),
        }

        analysis["next_tools"] = ["word_copy_template", "word_list_tables", "word_get_section_guidance"]

        return analysis

    def tool_word_add_comment(
        self,
        file_path: str,
        target_text: str,
        comment_text: str,
        author: str = DEFAULT_AUTHOR,
        output_path: str = None
    ) -> dict[str, Any]:
        """Add a comment to a specific portion of text in the document.

        Attaches a Word comment to the first occurrence of target_text.
        Useful for explaining changes, flagging items for review, or
        providing context on decisions made during document generation.

        Comments appear in Word's Review pane and print in the margin.

        Example:
            add_comment(
                file_path="04. Artifacts/contoso-sow.docx",
                target_text="15 legacy applications",
                comment_text="Count verified in discovery session 2024-01-15"
            )

            add_comment(
                file_path="04. Artifacts/contoso-sow.docx",
                target_text="$1.2M",
                comment_text="ROM estimate - requires detailed estimation phase",
                author="Solution Architect"
            )

        Args:
            file_path: Path to the .docx document
            target_text: Text to attach the comment to (first match used)
            comment_text: The comment content
            author: Author name for the comment (default: "Solution Architect Agent")
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status with comment location details
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        author = author or getattr(self, "_comment_author", DEFAULT_AUTHOR)

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        # Use ZIP-based approach for reliable comment insertion
        # python-docx doesn't properly support creating new OPC parts
        import shutil
        import tempfile
        import zipfile

        from lxml import etree

        # Define namespaces
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
        REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
        COMMENTS_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        COMMENTS_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

        NSMAP = {'w': W_NS, 'w14': W14_NS}

        # Generate unique comment ID
        global _revision_id
        _revision_id += 1
        comment_id = str(_revision_id)

        # Create temp directory for extraction
        temp_dir = tempfile.mkdtemp()

        try:
            # Extract the docx
            with zipfile.ZipFile(file_path, 'r') as zf:
                zf.extractall(temp_dir)

            # Read document.xml
            doc_xml_path = Path(temp_dir) / 'word' / 'document.xml'
            doc_tree = etree.parse(str(doc_xml_path))
            doc_root = doc_tree.getroot()

            # Find the target text in document.xml
            found_para = None
            found_run = None
            found_section = None
            found_table_info = None

            # Search in paragraphs
            for para in doc_root.iter(f'{{{W_NS}}}p'):
                para_text = ''.join(para.itertext())
                if target_text in para_text:
                    found_para = para
                    # Find which run contains the text
                    cumulative = ''
                    for run in para.iter(f'{{{W_NS}}}r'):
                        run_text = ''.join(run.itertext())
                        cumulative += run_text
                        if target_text in cumulative:
                            found_run = run
                            break
                    break

            # Search in tables if not found
            if found_para is None:
                for tbl_idx, tbl in enumerate(doc_root.iter(f'{{{W_NS}}}tbl')):
                    for row_idx, row in enumerate(tbl.iter(f'{{{W_NS}}}tr')):
                        for cell_idx, cell in enumerate(row.iter(f'{{{W_NS}}}tc')):
                            cell_text = ''.join(cell.itertext())
                            if target_text in cell_text:
                                for para in cell.iter(f'{{{W_NS}}}p'):
                                    found_para = para
                                    for run in para.iter(f'{{{W_NS}}}r'):
                                        run_text = ''.join(run.itertext())
                                        if target_text in run_text or target_text in ''.join(para.itertext()):
                                            found_run = run
                                            break
                                    break
                                found_table_info = {
                                    "table_index": tbl_idx,
                                    "row": row_idx,
                                    "cell": cell_idx
                                }
                                break
                        if found_para is not None:
                            break
                    if found_para is not None:
                        break

            if found_para is None:
                shutil.rmtree(temp_dir)
                return {"error": f"Target text '{target_text[:50]}...' not found in document"}

            # Insert comment markers in document.xml
            # Create comment range start
            range_start = etree.Element(f'{{{W_NS}}}commentRangeStart')
            range_start.set(f'{{{W_NS}}}id', comment_id)

            # Create comment range end
            range_end = etree.Element(f'{{{W_NS}}}commentRangeEnd')
            range_end.set(f'{{{W_NS}}}id', comment_id)

            # Create comment reference run
            comment_ref_run = etree.Element(f'{{{W_NS}}}r')
            comment_ref = etree.SubElement(comment_ref_run, f'{{{W_NS}}}commentReference')
            comment_ref.set(f'{{{W_NS}}}id', comment_id)

            # Insert markers
            if found_run is not None:
                # Insert before the run
                run_parent = found_run.getparent()
                run_idx = list(run_parent).index(found_run)
                run_parent.insert(run_idx, range_start)
                # Insert end after the run
                new_run_idx = list(run_parent).index(found_run)
                run_parent.insert(new_run_idx + 1, range_end)
                # Insert reference after end
                end_idx = list(run_parent).index(range_end)
                run_parent.insert(end_idx + 1, comment_ref_run)
            else:
                # Insert at paragraph level
                children = list(found_para)
                if children:
                    found_para.insert(0, range_start)
                else:
                    found_para.append(range_start)
                found_para.append(range_end)
                found_para.append(comment_ref_run)

            # Save modified document.xml
            doc_tree.write(str(doc_xml_path), xml_declaration=True, encoding='UTF-8', standalone=True)

            # Check if comments.xml exists
            comments_xml_path = Path(temp_dir) / 'word' / 'comments.xml'
            rels_path = Path(temp_dir) / 'word' / '_rels' / 'document.xml.rels'
            content_types_path = Path(temp_dir) / '[Content_Types].xml'

            # Create or update comments.xml
            comment_date = datetime.utcnow().isoformat() + 'Z'
            initials = ''.join(w[0].upper() for w in author.split()[:2])

            if comments_xml_path.exists():
                # Parse existing comments.xml
                comments_tree = etree.parse(str(comments_xml_path))
                comments_root = comments_tree.getroot()
            else:
                # Create new comments.xml
                comments_root = etree.Element(f'{{{W_NS}}}comments', nsmap=NSMAP)
                comments_tree = etree.ElementTree(comments_root)

            # Create comment element
            comment_elem = etree.SubElement(comments_root, f'{{{W_NS}}}comment')
            comment_elem.set(f'{{{W_NS}}}id', comment_id)
            comment_elem.set(f'{{{W_NS}}}author', author)
            comment_elem.set(f'{{{W_NS}}}date', comment_date)
            comment_elem.set(f'{{{W_NS}}}initials', initials)

            # Add paragraph with comment text
            comment_para = etree.SubElement(comment_elem, f'{{{W_NS}}}p')
            comment_run = etree.SubElement(comment_para, f'{{{W_NS}}}r')
            comment_text_elem = etree.SubElement(comment_run, f'{{{W_NS}}}t')
            comment_text_elem.text = comment_text

            # Save comments.xml
            comments_tree.write(str(comments_xml_path), xml_declaration=True, encoding='UTF-8', standalone=True)

            # Update relationships if comments.xml is new
            if not rels_path.exists():
                # Create _rels directory and document.xml.rels
                rels_dir = rels_path.parent
                rels_dir.mkdir(parents=True, exist_ok=True)
                rels_root = etree.Element('Relationships', xmlns=REL_NS)
                rels_tree = etree.ElementTree(rels_root)
            else:
                rels_tree = etree.parse(str(rels_path))
                rels_root = rels_tree.getroot()

            # Check if comments relationship exists
            comments_rel_exists = False
            for rel in rels_root.iter(f'{{{REL_NS}}}Relationship'):
                if rel.get('Type') == COMMENTS_REL_TYPE:
                    comments_rel_exists = True
                    break

            if not comments_rel_exists:
                # Generate unique relationship ID
                existing_ids = [rel.get('Id') for rel in rels_root.iter(f'{{{REL_NS}}}Relationship')]
                rel_id = 'rId1'
                counter = 1
                while rel_id in existing_ids:
                    counter += 1
                    rel_id = f'rId{counter}'

                # Add relationship
                new_rel = etree.SubElement(rels_root, f'{{{REL_NS}}}Relationship')
                new_rel.set('Id', rel_id)
                new_rel.set('Type', COMMENTS_REL_TYPE)
                new_rel.set('Target', 'comments.xml')

                rels_tree.write(str(rels_path), xml_declaration=True, encoding='UTF-8', standalone=True)

            # Update [Content_Types].xml if comments part is new
            ct_tree = etree.parse(str(content_types_path))
            ct_root = ct_tree.getroot()
            CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'

            # Check if comments content type exists
            comments_ct_exists = False
            for override in ct_root.iter(f'{{{CT_NS}}}Override'):
                if override.get('PartName') == '/word/comments.xml':
                    comments_ct_exists = True
                    break

            if not comments_ct_exists:
                override = etree.SubElement(ct_root, f'{{{CT_NS}}}Override')
                override.set('PartName', '/word/comments.xml')
                override.set('ContentType', COMMENTS_CONTENT_TYPE)
                ct_tree.write(str(content_types_path), xml_declaration=True, encoding='UTF-8', standalone=True)

            # Repack the docx
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for root, _dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path_full = Path(root) / file
                        arcname = file_path_full.relative_to(temp_dir)
                        zf.write(file_path_full, arcname)

            result = {
                "success": True,
                "file": output_path,
                "comment_id": comment_id,
                "target_text": target_text[:60] + ("..." if len(target_text) > 60 else ""),
                "comment": comment_text[:100] + ("..." if len(comment_text) > 100 else ""),
                "author": author,
                "message": f"Added comment to '{target_text[:30]}...'"
            }

            if found_section:
                result["section"] = found_section
            if found_table_info:
                result["table_location"] = found_table_info

            return result

        except Exception as e:
            return {"error": f"Failed to add comment: {e}"}
        finally:
            # Clean up temp directory
            shutil.rmtree(temp_dir, ignore_errors=True)

    def tool_word_patch_placeholder(
        self,
        file_path: str,
        placeholder: str,
        value: str,
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Replace a specific placeholder throughout a SOW document.

        Finds and replaces all instances of a placeholder pattern with
        the provided value. Works in both paragraphs and table cells.
        By default, changes are tracked for auditability.

        Example:
            patch_placeholder(
                file_path="04. Artifacts/contoso-sow.docx",
                placeholder="<Customer Name>",
                value="Contoso Ltd"
            )

            patch_placeholder(
                file_path="04. Artifacts/contoso-sow.docx",
                placeholder="[TBD]",
                value="Q1 2026"
            )

        Args:
            file_path: Path to the .docx document
            placeholder: The placeholder text to find (e.g., "<Customer Name>")
            value: The replacement value
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with replacement count
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        # Enable track revisions
        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        replacements = 0

        def replace_in_paragraph(para):
            nonlocal replacements
            para_text = _get_text_with_track_changes(para)
            if (placeholder in para_text
                    and _replace_with_track_changes(para, placeholder, value, author)):
                replacements += 1

        # Replace in paragraphs
        for para in doc.paragraphs:
            replace_in_paragraph(para)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)

        if replacements == 0:
            return {
                "success": True,
                "file": output_path,
                "placeholder": placeholder,
                "value": value,
                "replacements": 0,
                "track_changes": True,
                "author": author,
                "message": "No matching placeholders found; file not modified.",
                "next_tools": ["word_audit_sow", "word_fix_split_placeholders"]
            }

        doc.save(output_path)

        return {
            "success": True,
            "file": output_path,
            "placeholder": placeholder,
            "value": value,
            "replacements": replacements,
            "track_changes": True,
            "author": author,
            "message": f"Replaced {replacements} instances of '{placeholder}' with track changes by '{author}'",
            "next_tools": ["word_add_comment", "word_list_tables", "word_insert_table_row", "word_get_section_guidance"]
        }

    def tool_word_audit_sow(
        self,
        file_path: str,
        patterns: list[str] = None
    ) -> dict[str, Any]:
        """Audit a SOW document for unfilled placeholders, including those split across runs.

        This tool detects placeholders that standard patch_placeholder might miss
        because Word has split the text across multiple runs (due to formatting changes).

        Checks for:
        - Angle bracket placeholders: <Customer Name>, <Project Name>, etc.
        - Square bracket placeholders: [TBD], [insert ...], [Template Guidance: ...]
        - Split-run issues: placeholders visible in paragraph.text but not in any single run

        Example:
            audit_sow(file_path="04. Artifacts/contoso-sow.docx")

            audit_sow(
                file_path="04. Artifacts/contoso-sow.docx",
                patterns=["<Customer Name>", "[TBD]", "Contoso"]
            )

        Args:
            file_path: Path to the .docx document to audit
            patterns: Optional list of specific patterns to search for.
                      If not provided, uses default placeholder patterns.

        Returns:
            Audit report with found placeholders and split-run issues
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        # Default patterns to search for
        if patterns is None:
            patterns = [
                r'<[^>]+>',           # <Customer Name>, <n>, etc.
                r'\[TBD\]',           # [TBD]
                r'\[insert [^\]]+\]', # [insert WO number]
                r'\[Template [^\]]+\]', # [Template Guidance: ...]
            ]

        findings = {
            "placeholders_in_runs": [],      # Can be fixed with patch_placeholder
            "placeholders_split_runs": [],   # Need fix_split_placeholders
            "summary": {}
        }

        def check_paragraph(para, location: str):
            """Check a paragraph for placeholders."""
            full_text = _get_text_with_track_changes(para)

            for pattern in patterns:
                # Check if pattern is a regex (contains metacharacters like +, *, etc.)
                # or a literal string
                is_regex = any(c in pattern for c in r'+*?{}|^$\\')

                if is_regex:
                    # Regex pattern
                    matches = re.findall(pattern, full_text)
                    for match in matches:
                        in_single_run = any(match in run.text for run in para.runs)

                        finding = {
                            "location": location,
                            "placeholder": match,
                            "paragraph_text": full_text[:100] + ("..." if len(full_text) > 100 else ""),
                            "in_single_run": in_single_run
                        }

                        if in_single_run:
                            findings["placeholders_in_runs"].append(finding)
                        else:
                            finding["runs"] = [run.text for run in para.runs]
                            findings["placeholders_split_runs"].append(finding)
                else:
                    # Literal pattern
                    if pattern in full_text:
                        # Check if it's in a single run
                        in_single_run = any(pattern in run.text for run in para.runs)

                        finding = {
                            "location": location,
                            "placeholder": pattern,
                            "paragraph_text": full_text[:100] + ("..." if len(full_text) > 100 else ""),
                            "in_single_run": in_single_run
                        }

                        if in_single_run:
                            findings["placeholders_in_runs"].append(finding)
                        else:
                            finding["runs"] = [run.text for run in para.runs]
                            findings["placeholders_split_runs"].append(finding)

        # Check body paragraphs
        for idx, para in enumerate(doc.paragraphs):
            check_paragraph(para, f"paragraph_{idx}")

        # Check tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        check_paragraph(para, f"table_{t_idx}_row_{r_idx}_cell_{c_idx}_para_{p_idx}")

        # Check headers and footers
        for section in doc.sections:
            for h_idx, para in enumerate(section.header.paragraphs):
                check_paragraph(para, f"header_para_{h_idx}")
            for f_idx, para in enumerate(section.footer.paragraphs):
                check_paragraph(para, f"footer_para_{f_idx}")

        # Build summary
        findings["summary"] = {
            "total_placeholders": len(findings["placeholders_in_runs"]) + len(findings["placeholders_split_runs"]),
            "fixable_with_patch_placeholder": len(findings["placeholders_in_runs"]),
            "need_fix_split_placeholders": len(findings["placeholders_split_runs"]),
            "unique_placeholders": list({
                f["placeholder"] for f in findings["placeholders_in_runs"] + findings["placeholders_split_runs"]
            })
        }

        if findings["placeholders_split_runs"]:
            findings["recommendation"] = "Use fix_split_placeholders to repair split-run issues before using patch_placeholder"
            findings["next_tools"] = ["word_fix_split_placeholders", "word_patch_placeholder"]
        elif findings["placeholders_in_runs"]:
            findings["recommendation"] = "Use patch_placeholder to replace remaining placeholders"
            findings["next_tools"] = ["word_patch_placeholder", "word_insert_table_row"]
        else:
            findings["recommendation"] = "Document is clean - no placeholders found. Review content before final cleanup."
            findings["next_tools"] = ["word_list_tables", "word_check_tracking", "word_audit_completion"]

        return findings

    def tool_word_fix_split_placeholders(
        self,
        file_path: str,
        replacements: dict[str, str],
        output_path: str = None,
        author: str = DEFAULT_AUTHOR
    ) -> dict[str, Any]:
        """Fix placeholders that are split across multiple Word runs.

        Word often splits text across runs when formatting changes or edits occur.
        This makes placeholders like <Customer Name> invisible to simple run-based
        search and replace. This tool reconstructs paragraph text and replaces
        placeholders even when split. All changes are tracked for auditability.

        Example:
            fix_split_placeholders(
                file_path="04. Artifacts/contoso-sow.docx",
                replacements={
                    "<Customer Name>": "Contoso Ltd",
                    "<Project Name>": "Cloud Migration",
                    "[TBD]": "Q1 2026"
                }
            )

        Args:
            file_path: Path to the .docx document
            replacements: Dictionary mapping placeholders to their replacement values
            output_path: Optional output path (defaults to overwriting input)
            author: Author name for tracked changes (default: "Solution Architect Agent")

        Returns:
            Status with replacement counts per placeholder
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        # Enable track revisions
        _enable_track_revisions(doc)
        global _revision_id
        _revision_id = 0

        stats = dict.fromkeys(replacements, 0)

        def fix_paragraph(para) -> int:
            """Fix placeholders in a paragraph, handling split runs."""
            fixed = 0
            full_text = _get_text_with_track_changes(para)

            # Check if any replacement is needed
            needs_fix = False
            for placeholder, _value in replacements.items():
                if placeholder in full_text:
                    needs_fix = True
                    stats[placeholder] += full_text.count(placeholder)
                    fixed += full_text.count(placeholder)

            if needs_fix:
                # Use track changes for each replacement
                for placeholder, value in replacements.items():
                    if placeholder in full_text:
                        _replace_with_track_changes(para, placeholder, value, author)
                        full_text = para.text  # Update after replacement

            return fixed

        total_fixed = 0

        # Fix body paragraphs
        for para in doc.paragraphs:
            total_fixed += fix_paragraph(para)

        # Fix tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total_fixed += fix_paragraph(para)

        # Fix headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                total_fixed += fix_paragraph(para)
            for para in section.footer.paragraphs:
                total_fixed += fix_paragraph(para)

        if total_fixed == 0:
            return {
                "success": True,
                "file": output_path,
                "total_replacements": 0,
                "by_placeholder": stats,
                "track_changes": True,
                "author": author,
                "message": "No matching placeholders found; file not modified.",
                "next_tools": ["word_audit_sow", "word_fix_split_placeholders"]
            }

        doc.save(output_path)

        return {
            "success": True,
            "file": output_path,
            "total_replacements": total_fixed,
            "by_placeholder": stats,
            "track_changes": True,
            "author": author,
            "message": f"Fixed {total_fixed} placeholders across split runs (tracked by '{author}')",
            "next_tools": ["word_list_tables", "word_get_section_guidance", "word_insert_table_row"]
        }

    def tool_word_patch_with_track_changes(
        self,
        file_path: str,
        replacements: dict[str, str],
        author: str = DEFAULT_AUTHOR,
        output_path: str = None
    ) -> dict[str, Any]:
        """Replace text in a document with Word Track Changes enabled.

        Creates revision marks (insertions/deletions) that appear in Word's
        review mode. Old text is marked as deleted (red strikethrough) and
        new text is marked as inserted (green underline).

        This is useful for:
        - Auditable document changes
        - Review workflows where changes need approval
        - Comparing before/after states in Word

        Example:
            patch_with_track_changes(
                file_path="04. Artifacts/contoso-sow.docx",
                replacements={
                    "<Customer Name>": "Contoso Ltd",
                    "<Project Name>": "Cloud Migration",
                    "[TBD]": "Q1 2026"
                },
                author="Solution Architect"
            )

        Args:
            file_path: Path to the .docx document
            replacements: Dictionary mapping old text to new text
            author: Name to attribute changes to (appears in Word's review pane)
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status with replacement counts
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        # Enable track revisions in settings.xml so Word shows the changes
        _enable_track_revisions(doc)

        # Reset revision ID counter for this document
        global _revision_id
        _revision_id = 0

        stats = dict.fromkeys(replacements, 0)

        def patch_paragraph(para) -> int:
            """Patch a paragraph with track changes."""
            patched = 0
            full_text = _get_text_with_track_changes(para)

            for old_text, new_text in replacements.items():
                if (old_text in full_text
                        and _replace_with_track_changes(para, old_text, new_text, author)):
                    stats[old_text] += 1
                    patched += 1
                    # Update full_text for next iteration
                    full_text = _get_text_with_track_changes(para)

            return patched

        total_patched = 0

        # Patch body paragraphs
        for para in doc.paragraphs:
            total_patched += patch_paragraph(para)

        # Patch tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total_patched += patch_paragraph(para)

        # Patch headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                total_patched += patch_paragraph(para)
            for para in section.footer.paragraphs:
                total_patched += patch_paragraph(para)

        doc.save(output_path)

        return {
            "success": True,
            "file": output_path,
            "total_changes": total_patched,
            "by_replacement": stats,
            "author": author,
            "message": f"Made {total_patched} tracked changes attributed to '{author}'",
            "next_tools": ["word_add_comment", "word_check_tracking", "word_audit_completion"]
        }

    def tool_word_enable_track_changes(
        self,
        file_path: str,
        output_path: str = None
    ) -> dict[str, Any]:
        """Enable Track Changes mode in a Word document.

        Sets the document settings so that Word will track subsequent
        changes when the document is opened and edited.

        Note: This sets the tracking flag, but changes made by python-docx
        after this point will NOT be automatically tracked. Use
        patch_with_track_changes for programmatic tracked edits.

        Example:
            enable_track_changes(file_path="04. Artifacts/contoso-sow.docx")

        Args:
            file_path: Path to the .docx document
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status message
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        # Access settings element
        settings = doc.settings.element

        # Create or update trackRevisions element
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings.append(track_revisions)

        # Enable tracking
        track_revisions.set(qn('w:val'), 'true')

        doc.save(output_path)

        return {
            "success": True,
            "file": output_path,
            "message": "Track Changes enabled. Subsequent edits in Word will be tracked."
        }

    def tool_word_copy_template(
        self,
        template_name: str,
        output_path: str,
        template_dir: str = None
    ) -> dict[str, Any]:
        """Copy a SOW template to start a new document.

        Copies a template file from the templates directory to the specified
        output path, ready for editing with track changes.

        Available templates can be found in .github/skills/statement-of-work/templates/

        Example:
            copy_template(
                template_name="Agile.docx",
                output_path="04. Artifacts/contoso-sow.docx"
            )

        Args:
            template_name: Name of the template file (e.g., "Agile.docx")
            output_path: Destination path for the new document
            template_dir: Optional custom template directory path

        Returns:
            Status with source and destination paths
        """
        import shutil

        # Default template directory
        if template_dir is None:
            # Try to find templates relative to this file
            script_dir = Path(__file__).parent.parent.parent
            template_dir = script_dir / "skills" / "statement-of-work" / "templates"

        template_path = Path(template_dir) / template_name

        if not template_path.exists():
            # List available templates
            available = []
            if Path(template_dir).exists():
                available = [f.name for f in Path(template_dir).glob("*.docx")]
            return {
                "error": f"Template not found: {template_name}",
                "template_dir": str(template_dir),
                "available_templates": available
            }

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        try:
            shutil.copy2(template_path, output)
        except Exception as e:
            return {"error": f"Failed to copy template: {e}"}

        return {
            "success": True,
            "template": str(template_path),
            "output": str(output),
            "message": f"Copied template '{template_name}' to '{output_path}'",
            "next_tools": ["word_parse_sow_template", "word_list_tables", "word_list_sections"]
        }

    def tool_word_check_tracking(
        self,
        file_path: str
    ) -> dict[str, Any]:
        """Get a summary of track changes in a Word document.

        Returns counts of insertions and deletions, broken down by location
        (body text vs tables), with sample content for review.

        Example:
            check_tracking(file_path="04. Artifacts/contoso-sow.docx")

        Args:
            file_path: Path to the .docx document

        Returns:
            Track changes summary with counts and samples
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        body = doc.element.body

        # Check if track revisions is enabled
        settings = doc.settings.element
        track_rev = settings.find(qn("w:trackRevisions"))
        tracking_enabled = track_rev is not None

        # Count insertions and deletions in body (non-table)
        all_ins = body.findall('.//' + qn('w:ins'))
        all_del = body.findall('.//' + qn('w:del'))

        # Analyze by location
        table_stats = {}
        body_ins = 0
        body_del = 0

        for idx, table in enumerate(doc.tables):
            tbl = table._tbl
            t_ins = tbl.findall('.//' + qn('w:ins'))
            t_del = tbl.findall('.//' + qn('w:del'))
            if t_ins or t_del:
                # Get table purpose - use track-change-aware text extraction
                header = [_get_cell_text(cell)[:20] for cell in table.rows[0].cells] if table.rows else []
                purpose = self._identify_table_purpose(header) if header else "unknown"
                table_stats[f"table_{idx}"] = {
                    "purpose": purpose,
                    "insertions": len(t_ins),
                    "deletions": len(t_del),
                    "header_preview": header[:3]
                }

        # Body stats = total - table stats
        table_ins_total = sum(t["insertions"] for t in table_stats.values())
        table_del_total = sum(t["deletions"] for t in table_stats.values())
        body_ins = len(all_ins) - table_ins_total
        body_del = len(all_del) - table_del_total

        # Get sample changes
        samples = {"insertions": [], "deletions": []}
        for ins in all_ins[:5]:
            text_els = ins.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in text_els)[:50]
            author = ins.get(qn('w:author'), 'unknown')
            if text.strip():
                samples["insertions"].append({"text": text, "author": author})

        for d in all_del[:5]:
            text_els = d.findall('.//' + qn('w:delText'))
            text = ''.join(t.text or '' for t in text_els)[:50]
            author = d.get(qn('w:author'), 'unknown')
            if text.strip():
                samples["deletions"].append({"text": text, "author": author})

        return {
            "success": True,
            "file": file_path,
            "tracking_enabled": tracking_enabled,
            "summary": {
                "total_insertions": len(all_ins),
                "total_deletions": len(all_del),
                "body_insertions": body_ins,
                "body_deletions": body_del
            },
            "tables": table_stats,
            "samples": samples,
            "has_changes": len(all_ins) > 0 or len(all_del) > 0
        }

    def tool_word_audit_completion(
        self,
        file_path: str
    ) -> dict[str, Any]:
        """Audit a SOW document for completion status.

        Comprehensive check for:
        - Unfilled placeholders (angle brackets, square brackets, TBD)
        - Empty required sections
        - Empty table cells in key tables
        - Template instruction remnants
        - Track changes pending review

        Returns a completion score and detailed findings.

        Example:
            audit_completion(file_path="04. Artifacts/contoso-sow.docx")

        Args:
            file_path: Path to the .docx document

        Returns:
            Completion audit with score, issues, and recommendations
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": f"Failed to open document: {e}"}

        issues = {
            "placeholders": [],
            "empty_sections": [],
            "empty_table_cells": [],
            "instruction_remnants": [],
            "pending_track_changes": False
        }

        # Placeholder patterns
        placeholder_patterns = [
            r'<[^>]+>',  # <Customer Name>
            r'\[TBD\]',
            r'\[insert[^\]]*\]',
            r'\[Template[^\]]*\]',
            r'\[PLACEHOLDER[^\]]*\]',
        ]

        # Instruction keywords
        instruction_keywords = [
            'template guidance',
            'replace with',
            'insert your',
            'add your',
            'example:',
            'delete this',
        ]

        # Check paragraphs
        current_section = "Document Start"
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else "Normal"

            # Track sections
            if "Heading" in style:
                # Check if previous section was empty
                if current_section and not any(section_content):
                    issues["empty_sections"].append(current_section)
                current_section = text[:50] if text else "Unnamed Section"
                section_content = []
            else:
                if text:
                    section_content.append(text)

            # Check for placeholders
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    issues["placeholders"].append({
                        "text": match,
                        "location": current_section,
                        "context": text[:80]
                    })

            # Check for instruction remnants
            text_lower = text.lower()
            for keyword in instruction_keywords:
                if keyword in text_lower:
                    issues["instruction_remnants"].append({
                        "keyword": keyword,
                        "location": current_section,
                        "context": text[:80]
                    })

        # Check tables for empty cells in key columns - use track-change-aware text extraction
        key_tables = ["business_objectives", "epics", "staffing", "definitions"]
        for idx, table in enumerate(doc.tables):
            if not table.rows:
                continue
            header = [_get_cell_text(cell) for cell in table.rows[0].cells]
            purpose = self._identify_table_purpose(header)

            if purpose in key_tables:
                for row_idx, row in enumerate(table.rows[1:], 1):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = _get_cell_text(cell)
                        # Check for empty or placeholder content
                        if not cell_text or cell_text in ['', '-', 'N/A', 'TBD']:
                            issues["empty_table_cells"].append({
                                "table": purpose,
                                "table_index": idx,
                                "row": row_idx,
                                "column": header[col_idx] if col_idx < len(header) else f"Col {col_idx}"
                            })

        # Check for pending track changes
        body = doc.element.body
        ins_count = len(body.findall('.//' + qn('w:ins')))
        del_count = len(body.findall('.//' + qn('w:del')))
        if ins_count > 0 or del_count > 0:
            issues["pending_track_changes"] = {
                "insertions": ins_count,
                "deletions": del_count,
                "action": "Review and accept/reject changes before finalizing"
            }

        # Calculate completion score
        total_issues = (
            len(issues["placeholders"]) +
            len(issues["empty_sections"]) +
            len(issues["empty_table_cells"]) +
            len(issues["instruction_remnants"])
        )

        # Score: 100 = perfect, deduct points for issues
        score = max(0, 100 - (total_issues * 5))

        # Determine status
        if score >= 95 and not issues["pending_track_changes"]:
            status = "READY"
            recommendation = "Document appears complete and ready for review."
        elif score >= 80:
            status = "NEEDS_REVIEW"
            recommendation = "Minor issues found. Review flagged items before finalizing."
        elif score >= 50:
            status = "IN_PROGRESS"
            recommendation = "Significant content still missing. Continue filling sections."
        else:
            status = "DRAFT"
            recommendation = "Document is still a draft. Many sections need content."

        # Determine next tools based on status
        if status == "READY":
            next_tools = ["word_cleanup_sow"]
        elif issues["placeholders"]:
            next_tools = ["word_audit_sow", "word_fix_split_placeholders", "word_patch_placeholder"]
        elif issues["empty_sections"]:
            next_tools = ["word_patch_section", "word_get_section"]
        elif issues["instruction_remnants"]:
            next_tools = ["word_cleanup_sow"]
        else:
            next_tools = ["word_cleanup_sow"]

        return {
            "success": True,
            "file": file_path,
            "status": status,
            "score": score,
            "recommendation": recommendation,
            "summary": {
                "placeholders_found": len(issues["placeholders"]),
                "empty_sections": len(issues["empty_sections"]),
                "empty_table_cells": len(issues["empty_table_cells"]),
                "instruction_remnants": len(issues["instruction_remnants"]),
                "pending_changes": bool(issues["pending_track_changes"])
            },
            "issues": issues,
            "next_tools": next_tools
        }

    def tool_word_replace_global_variables(
        self,
        file_path: str,
        replacements: dict[str, str],
        output_path: str = None
    ) -> dict[str, Any]:
        """Replace global variable placeholders throughout a Word document.

        This tool performs AGGRESSIVE text replacement at the XML level,
        handling text split across multiple runs within paragraphs. It:
        - Reconstructs full paragraph text from all w:t elements
        - Replaces placeholders even when split across runs
        - Removes SDT (content control) structures entirely to prevent reset
        - Processes headers, footers, and document body

        Use this for placeholders like <Customer Name> and <Project Name> that
        Word templates often store in content controls or split across runs.

        Example:
            replace_global_variables(
                file_path="04. Artifacts/contoso-sow.docx",
                replacements={
                    "<Customer Name>": "Contoso",
                    "<Project Name>": "Customer Care Transformation",
                    "<Microsoft OR Partner name>": "Microsoft"
                }
            )

        Args:
            file_path: Path to the .docx document
            replacements: Dictionary mapping placeholders to replacement values
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status with replacement counts by location
        """
        import shutil
        import tempfile
        import zipfile

        from lxml import etree

        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        if output_path is None:
            output_path = file_path

        stats = {
            "document_xml": 0,
            "headers": 0,
            "footers": 0,
            "sdts_unwrapped": 0,
            "paragraphs_fixed": 0,
            "by_placeholder": dict.fromkeys(replacements, 0)
        }

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        def replace_in_xml(xml_content: bytes) -> tuple:
            """Replace placeholders in XML content with aggressive cross-run handling."""
            root = etree.fromstring(xml_content)
            total_replaced = 0
            sdts_unwrapped = 0
            paragraphs_fixed = 0

            # Step 1: UNWRAP all SDTs (Structured Document Tags / Content Controls)
            # Instead of just removing placeholder properties, completely unwrap SDTs
            # by moving their content to the parent and removing the SDT wrapper.
            # This prevents any SDT-related reset behavior.
            for sdt in root.findall('.//w:sdt', ns):
                parent = sdt.getparent()
                if parent is None:
                    continue
                idx = list(parent).index(sdt)

                # Find the content element (w:sdtContent)
                sdt_content = sdt.find('w:sdtContent', ns)
                if sdt_content is not None:
                    # Move all children of sdtContent to parent
                    for child in list(sdt_content):
                        parent.insert(idx, child)
                        idx += 1

                # Remove the SDT element
                parent.remove(sdt)
                sdts_unwrapped += 1

            # Step 2: Process each paragraph - reconstruct text and replace across runs
            for para in root.findall('.//w:p', ns):
                # Get all text elements in this paragraph
                text_elements = para.findall('.//w:t', ns)
                if not text_elements:
                    continue

                # Reconstruct full paragraph text
                full_text = ''.join(t.text or '' for t in text_elements)

                # Check if any replacements are needed
                new_text = full_text
                para_replacements = 0
                for placeholder, value in replacements.items():
                    if placeholder in new_text:
                        count = new_text.count(placeholder)
                        new_text = new_text.replace(placeholder, value)
                        stats["by_placeholder"][placeholder] += count
                        para_replacements += count
                        total_replaced += count

                if para_replacements > 0:
                    paragraphs_fixed += 1
                    # Rewrite the text elements
                    # Put all text in the first element, clear the rest
                    if text_elements:
                        text_elements[0].text = new_text
                        for t in text_elements[1:]:
                            t.text = ''

            # Step 3: Also do simple replacement for any remaining w:t elements
            # (in case some are outside paragraphs)
            for t in root.findall('.//w:t', ns):
                if t.text:
                    original = t.text
                    modified = original
                    for placeholder, value in replacements.items():
                        if placeholder in modified:
                            count = modified.count(placeholder)
                            modified = modified.replace(placeholder, value)
                            # Don't double-count, only count if not already processed
                    if modified != original:
                        t.text = modified

            stats["sdts_unwrapped"] = stats.get("sdts_unwrapped", 0) + sdts_unwrapped
            stats["paragraphs_fixed"] = stats.get("paragraphs_fixed", 0) + paragraphs_fixed
            return etree.tostring(root, xml_declaration=True, encoding='UTF-8'), total_replaced

        # Create a temporary file to work with
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp_path = tmp.name

        try:
            # Copy original to temp
            shutil.copy2(file_path, tmp_path)

            # Open as zip and modify XML files
            with (zipfile.ZipFile(tmp_path, 'r') as zin,
                  zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout):
                for item in zin.namelist():
                    content = zin.read(item)

                    if item == 'word/document.xml':
                        content, count = replace_in_xml(content)
                        stats["document_xml"] = count
                    elif 'header' in item and item.endswith('.xml'):
                        content, count = replace_in_xml(content)
                        stats["headers"] += count
                    elif 'footer' in item and item.endswith('.xml'):
                        content, count = replace_in_xml(content)
                        stats["footers"] += count

                    zout.writestr(item, content)

        finally:
            # Clean up temp file
            Path(tmp_path).unlink(missing_ok=True)

        total = stats["document_xml"] + stats["headers"] + stats["footers"]

        return {
            "success": True,
            "file": output_path,
            "total_replacements": total,
            "by_location": {
                "document_body": stats["document_xml"],
                "headers": stats["headers"],
                "footers": stats["footers"]
            },
            "by_placeholder": stats["by_placeholder"],
            "sdts_unwrapped": stats.get("sdts_unwrapped", 0),
            "paragraphs_fixed": stats.get("paragraphs_fixed", 0),
            "message": f"Replaced {total} placeholders, unwrapped {stats.get('sdts_unwrapped', 0)} content controls, fixed {stats.get('paragraphs_fixed', 0)} paragraphs"
        }

    # =========================================================================
    # MCP PROMPTS - Discoverable workflow guides
    # =========================================================================

    def prompt_sow_generation(self) -> dict[str, Any]:
        """Generate a complete Statement of Work from a template.

        Category: sow-workflow

        This prompt guides the complete SOW generation workflow from
        copying a template through filling all content and final cleanup.

        Returns:
            Prompt definition with workflow steps
        """
        return {
            "name": "sow_generation",
            "description": "Complete SOW generation workflow from template to final document",
            "arguments": [
                {"name": "customer_name", "description": "Customer/company name", "required": True},
                {"name": "project_name", "description": "Project or engagement name", "required": True},
                {"name": "template", "description": "Template to use (default: Agile.docx)", "required": False}
            ],
            "instructions": """
## SOW Generation Workflow

Follow these steps to generate a complete Statement of Work:

### Step 1: Copy Template
Use `copy_template` to create a new document from the template.
```
copy_template(
    template_name="Agile.docx",
    output_path="04. Artifacts/{customer}-sow.docx"
)
```

### Step 2: Parse Template Structure
Use `parse_sow_template` to understand what needs to be filled.
```
parse_sow_template(template_path="04. Artifacts/{customer}-sow.docx")
```

### Step 3: Generate Base Content
Use `generate_sow` with structured data to fill placeholders and tables.
```
generate_sow(
    template_path="04. Artifacts/{customer}-sow.docx",
    output_path="04. Artifacts/{customer}-sow.docx",
    sow_data={...}
)
```

### Step 4: Fill Prose Sections
For each prose section (Introduction, Executive Summary, etc.):

**4a. Read template guidance first:**
```
get_section_guidance(
    file_path="04. Artifacts/{customer}-sow.docx",
    section_title="Introduction"
)
```
This extracts what the template expects you to write.

**4b. Then patch with real content:**
```
patch_section(
    file_path="04. Artifacts/{customer}-sow.docx",
    section_title="Introduction",
    new_content=["Introduction paragraph 1...", "Introduction paragraph 2..."]
)
```

### Step 5: Audit and Clean
- Use `audit_completion` to check for unfilled placeholders
- Use `fix_split_placeholders` for any split-run issues
- Use `cleanup_sow` to remove template instructions

### Step 6: Review Track Changes
- Use `check_tracking` to see all tracked changes
- Review document in Word to verify all placeholders are replaced
""",
            "next_tools": ["word_copy_template", "word_parse_sow_template", "word_generate_sow", "word_get_section_guidance"]
        }

    def prompt_section_editing(self) -> dict[str, Any]:
        """Edit prose sections in a SOW document.

        Category: sow-workflow

        This prompt guides reading and editing narrative sections
        like Introduction, Executive Summary, and Business Context.

        Returns:
            Prompt definition with section editing workflow
        """
        return {
            "name": "section_editing",
            "description": "Read and edit prose sections in a SOW document",
            "arguments": [
                {"name": "file_path", "description": "Path to SOW document", "required": True},
                {"name": "section_name", "description": "Section to edit", "required": True}
            ],
            "instructions": """
## Section Editing Workflow

### Step 1: List Sections
See all sections in the document:
```
list_sections(file_path="04. Artifacts/{customer}-sow.docx")
```

### Step 2: Read Template Guidance
**Before writing, understand what's expected.** Extract guidance for the section:
```
get_section_guidance(
    file_path="04. Artifacts/{customer}-sow.docx",
    section_title="Introduction"
)
```

This returns:
- `template_instructions`: What the template says to write
- `placeholders`: Variables like `<Customer Name>` to replace
- `highlighted_guidance`: Highlighted text indicating template hints
- `writing_hints`: Actionable suggestions

### Step 3: Read Current Content (if needed)
See the full section text:
```
get_section(
    file_path="04. Artifacts/{customer}-sow.docx",
    section_title="Introduction"
)
```

### Step 4: Prepare New Content
Craft replacement paragraphs that:
- Follow the guidance from Step 2
- Address customer-specific context
- Reference specific business objectives
- Align with the engagement approach

### Step 5: Update Section
Replace section content with track changes:
```
patch_section(
    file_path="04. Artifacts/{customer}-sow.docx",
    section_title="Introduction",
    new_content=["Paragraph 1...", "Paragraph 2..."],
    track_changes=True,
    author="Solution Architect"
)
```

### Step 6: Verify
Read the section again to confirm changes.
""",
            "next_tools": ["word_list_sections", "word_get_section_guidance", "word_get_section", "word_patch_section"]
        }

    def prompt_document_audit(self) -> dict[str, Any]:
        """Audit a SOW document for completion and quality.

        Category: sow-workflow

        This prompt guides comprehensive document auditing to ensure
        all placeholders are filled and the document is ready.

        Returns:
            Prompt definition with audit workflow
        """
        return {
            "name": "document_audit",
            "description": "Comprehensive SOW document audit and cleanup",
            "arguments": [
                {"name": "file_path", "description": "Path to SOW document", "required": True}
            ],
            "instructions": """
## Document Audit Workflow

### Step 1: Check Completion Status
Get overall completion score and issues:
```
audit_completion(file_path="04. Artifacts/{customer}-sow.docx")
```

### Step 2: Find Placeholders
Identify all unfilled placeholders including split-run issues:
```
audit_sow(file_path="04. Artifacts/{customer}-sow.docx")
```

### Step 3: Fix Issues
For placeholders in single runs:
```
patch_placeholder(
    file_path="04. Artifacts/{customer}-sow.docx",
    placeholder="<Customer Name>",
    value="Contoso Ltd"
)
```

For split-run placeholders:
```
fix_split_placeholders(
    file_path="04. Artifacts/{customer}-sow.docx",
    replacements={"<Customer Name>": "Contoso Ltd", ...}
)
```

### Step 4: Review Track Changes
See all changes made:
```
check_tracking(file_path="04. Artifacts/{customer}-sow.docx")
```

### Step 5: Cleanup Template Artifacts
Remove instructions and formatting artifacts:
```
cleanup_sow(
    file_path="04. Artifacts/{customer}-sow.docx",
    track_changes=True
)
```

### Step 6: Final Review
Open document in Word to verify all placeholders are correctly replaced.
""",
            "next_tools": ["word_audit_completion", "word_audit_sow", "word_fix_split_placeholders", "word_cleanup_sow"]
        }

    def prompt_table_editing(self) -> dict[str, Any]:
        """Edit tables in a SOW document.

        Category: sow-workflow

        This prompt guides reading and editing structured tables
        like Business Objectives, Epics, and Staffing.

        Returns:
            Prompt definition with table editing workflow
        """
        return {
            "name": "table_editing",
            "description": "Read and edit tables in a SOW document",
            "arguments": [
                {"name": "file_path", "description": "Path to SOW document", "required": True},
                {"name": "table_type", "description": "Type of table to edit", "required": True}
            ],
            "instructions": """
## Table Editing Workflow

### Step 1: List Tables
See all tables and their purposes:
```
list_tables(file_path="04. Artifacts/{customer}-sow.docx")
```

### Step 2: Read Table Content
Get current table content by identifier or index:
```
get_table(
    file_path="04. Artifacts/{customer}-sow.docx",
    table_identifier="business_objectives"
)
```

### Step 3: Prepare Row Data
Structure rows as dictionaries matching header columns.

### Step 4: Update Table
Replace or append rows:
```
patch_table(
    file_path="04. Artifacts/{customer}-sow.docx",
    table_identifier="business_objectives",
    rows=[
        {"Desired Business Objectives": "...", "Provider activities": "...", "Assumptions": "..."}
    ],
    append=False,  # Replace all rows
    track_changes=True
)
```

### Step 5: Verify
Read the table again to confirm changes.
""",
            "next_tools": ["word_list_tables", "word_get_table", "word_insert_table_row"]
        }
