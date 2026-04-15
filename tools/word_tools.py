#!/usr/bin/env python3
"""
word_tools.py - MCP tools for Word document processing

Provides tools to parse, extract content from, and generate Word (.docx) files.
Supports full GitHub Flavored Markdown input with headings, bullets, bold/italic,
strikethrough, tables, code blocks, and task lists.
"""

import os
import re
from pathlib import Path
from typing import Any, Literal

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .markdown_parser import (
    CodeBlock,
    HorizontalRule,
    Paragraph,
    Table,
    TextRun,
    parse_markdown_to_nodes,
)
from .save_utils import open_docx_with_retries, resolve_office_path, safe_save_docx

DEFAULT_COMMENT_AUTHOR = os.environ.get("MCP_AUTHOR", "Solution Architect Agent")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID_NS = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
PKG_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
REL_COMMENTS_EXTENDED = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
CT_COMMENTS_EXTENDED = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"


def _is_done_value(value: str | None) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "yes"}


def _qname_attr(element, *attr_names: str) -> str | None:
    for name in attr_names:
        if name in element.attrib and element.attrib[name] is not None:
            return str(element.attrib[name])
    return None


def _parse_comments_ids_xml(comments_ids_bytes, etree_module) -> tuple[dict[str, str], list[str]]:
    """Parse commentsIds.xml and return explicit+ordered paraId mappings.

    Returns:
        (explicit_map, ordered_para_ids)
        explicit_map maps comment id-like attributes to paraId when present.
        ordered_para_ids preserves declaration order for index-based fallback.
    """
    explicit: dict[str, str] = {}
    ordered: list[str] = []

    try:
        root = etree_module.fromstring(comments_ids_bytes)
    except Exception:
        return explicit, ordered

    for node in root.iter():
        local_name = etree_module.QName(node).localname.lower()
        if local_name != "commentid":
            continue

        para_id = _qname_attr(
            node,
            f"{{{W16CID_NS}}}paraId",
            f"{{{W14_NS}}}paraId",
            "paraId",
        )
        if not para_id:
            continue
        ordered.append(str(para_id))

        cid = _qname_attr(
            node,
            f"{{{W_NS}}}id",
            f"{{{W16CID_NS}}}commentId",
            f"{{{W16CID_NS}}}id",
            "commentId",
            "id",
        )
        if cid is not None and str(cid).strip() != "":
            explicit[str(cid)] = str(para_id)

    return explicit, ordered


def _ensure_comments_extended_package_bits(parts: dict[str, bytes], etree_module) -> None:
    """Ensure content type + relationship entries exist for commentsExtended.xml."""
    if "[Content_Types].xml" in parts:
        ct_root = etree_module.fromstring(parts["[Content_Types].xml"])
        overrides = ct_root.findall(f".//{{{PKG_CT_NS}}}Override")
        has_override = any(
            node.get("PartName") == "/word/commentsExtended.xml"
            for node in overrides
        )
        if not has_override:
            override = etree_module.SubElement(ct_root, f"{{{PKG_CT_NS}}}Override")
            override.set("PartName", "/word/commentsExtended.xml")
            override.set("ContentType", CT_COMMENTS_EXTENDED)
            parts["[Content_Types].xml"] = etree_module.tostring(
                ct_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            )

    rels_key = "word/_rels/document.xml.rels"
    if rels_key in parts:
        rels_root = etree_module.fromstring(parts[rels_key])
        rels = rels_root.findall(f".//{{{PKG_REL_NS}}}Relationship")
        has_rel = any(node.get("Type") == REL_COMMENTS_EXTENDED for node in rels)
        if not has_rel:
            used_ids = {
                node.get("Id")
                for node in rels
                if node.get("Id")
            }
            next_idx = 1
            while f"rId{next_idx}" in used_ids:
                next_idx += 1
            rel = etree_module.SubElement(rels_root, f"{{{PKG_REL_NS}}}Relationship")
            rel.set("Id", f"rId{next_idx}")
            rel.set("Type", REL_COMMENTS_EXTENDED)
            rel.set("Target", "commentsExtended.xml")
            parts[rels_key] = etree_module.tostring(
                rels_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            )


class WordTools:
    """MCP tool mixin for Word document processing.

    Note: The extract/to_markdown methods are kept for internal use by
    office_unified_tools but are not exposed as separate MCP tools.
    They are filtered out in office_server.py via DEPRECATED_TOOLS.

    Public MCP tools:
    - word_from_markdown: Convert Markdown to Word document

    Internal methods (used by unified tools):
    - tool_word_extract: Extract text and structure
    - tool_word_to_markdown: Convert to Markdown format

    See WordAdvancedTools for SOW-specific and track-changes tools.
    """

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """Extract heading level from Word style names.

        Handles variants like "Heading 1", "Heading 1 NUM", and localized
        custom names that still include an integer heading level.
        """
        if not style_name.startswith("Heading"):
            return 1

        match = re.search(r"\b(\d+)\b", style_name)
        if not match:
            return 1

        try:
            return int(match.group(1))
        except ValueError:
            return 1

    def tool_word_extract(self, file_path: str) -> dict[str, Any]:
        """Extract text and structure from a Word document.

        NOTE: This method is internal - use office_read() instead.

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary with paragraphs, headings, tables, and metadata
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"error": f"Expected .docx file, got: {path.suffix}"}

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return {"error": f"Failed to open document: {open_error}"}

        paragraphs = []
        headings = []
        tables = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith("Heading"):
                level = self._get_heading_level(style_name)
                headings.append({"level": level, "text": text})

            paragraphs.append({
                "text": text,
                "style": style_name
            })

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)

        return {
            "file": path.name,
            "paragraphs": paragraphs,
            "headings": headings,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables)
        }

    def tool_word_to_markdown(self, file_path: str) -> str:
        """Convert a Word document to Markdown format.

        NOTE: This method is internal - use office_read(output_format="markdown") instead.

        Args:
            file_path: Path to the .docx file

        Returns:
            Markdown string representation of the document
        """
        if not HAS_DOCX:
            return "Error: python-docx not installed. Run: pip install python-docx"

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return f"Error: File not found: {file_path}"

        doc, _resolved_path, open_error = open_docx_with_retries(file_path)
        if open_error:
            return f"Error: Failed to open document: {open_error}"
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith("Heading"):
                level = self._get_heading_level(style_name)
                lines.append(f"{'#' * level} {text}")
            elif style_name == "List Bullet":
                lines.append(f"- {text}")
            elif style_name == "List Number":
                lines.append(f"1. {text}")
            else:
                lines.append(text)
            lines.append("")

        # Process tables
        for table in doc.tables:
            if not table.rows:
                continue

            # Header row
            header = [cell.text.strip() for cell in table.rows[0].cells]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")

            # Data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        return "\n".join(lines)

    def tool_word_from_markdown(
        self,
        output_path: str,
        markdown: str | None = None,
        markdown_file: str | None = None,
    ) -> dict[str, Any]:
        """Convert Markdown to a Word document from inline content or markdown_file.

        This is the primary tool for creating Word documents from text content.

        Supports GitHub Flavored Markdown:
        - Headings (# ## ### ####)
        - Bullet lists (- or *)
        - Numbered lists (1. 2. 3.)
        - Task lists (- [ ] and - [x])
        - Bold (**text**) and italic (*text*) inline formatting
        - Strikethrough (~~text~~)
        - Inline code (`code`) rendered in Consolas font
        - Tables (| col | col |) with proper formatting
        - Code blocks (```) rendered in monospace with language hints
        - Horizontal rules (---)

        Example:
            word_from_markdown(
                output_path="04. Artifacts/report.docx",
                markdown='''
# Project Status Report

## Executive Summary
The project is **on track** for Q4 delivery with ~~no~~ minor delays.

## Key Metrics
| Metric | Value | Status |
|--------|-------|--------|
| Budget | $120,000 | On track |
| Timeline | Q4 2026 | Green |
| Quality | 95% | Exceeds |

## Next Steps
- [x] Complete UAT testing
- [ ] Finalize documentation
- [ ] Schedule go-live review

1. Phase 1 complete
2. Phase 2 in progress
3. Phase 3 planned
'''
            )

        Args:
            output_path: Path for the output .docx file
            markdown: Full GitHub Flavored Markdown content (inline)
            markdown_file: Optional path to a Markdown file. Use this for
                very large documents to avoid MCP argument-size limits.

        Returns:
            Status dictionary with file path
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        markdown_text = markdown
        if markdown_file:
            resolved_md_path = resolve_office_path(markdown_file)
            md_path = Path(resolved_md_path)
            if not md_path.exists():
                return {"error": f"Markdown file not found: {markdown_file}"}
            if not md_path.is_file():
                return {"error": f"Markdown path is not a file: {markdown_file}"}
            markdown_text = md_path.read_text(encoding="utf-8")

        if markdown_text is None:
            return {"error": "Provide either 'markdown' or 'markdown_file'"}

        doc = Document()

        # Parse Markdown to structured nodes using GFM parser
        nodes = parse_markdown_to_nodes(markdown_text)

        def get_list_style(style_name: str, fallback: str = None) -> str | None:
            """Get a list style name, falling back to alternatives if not found."""
            # Try the requested style first
            try:
                doc.styles[style_name]
                return style_name
            except KeyError:
                pass

            # Try common alternatives
            alternatives = {
                "List Bullet": ["Bullet 1", "Bullet list", "List Paragraph"],
                "List Number": ["Number 1", "Numbered list", "List Paragraph"],
            }
            for alt in alternatives.get(style_name, []):
                try:
                    doc.styles[alt]
                    return alt
                except KeyError:
                    pass

            return fallback

        def add_runs_to_paragraph(para, runs: list[TextRun]):
            """Add formatted text runs to a paragraph."""
            for run_data in runs:
                run = para.add_run(run_data.text)
                if run_data.bold:
                    run.bold = True
                if run_data.italic:
                    run.italic = True
                if run_data.code:
                    run.font.name = "Consolas"
                if run_data.strikethrough:
                    run.font.strike = True

        def add_table(table: Table):
            """Add a table to the document."""
            if not table.rows:
                return

            num_rows = len(table.rows)
            num_cols = table.column_count

            doc_table = doc.add_table(rows=num_rows, cols=num_cols)
            doc_table.style = "Table Grid"

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < num_cols:
                        doc_cell = doc_table.rows[row_idx].cells[col_idx]
                        para = doc_cell.paragraphs[0]
                        para.clear()
                        add_runs_to_paragraph(para, cell.runs)

            # Add spacing after table
            doc.add_paragraph()

        def add_code_block(code_block: CodeBlock):
            """Add a code block to the document."""
            para = doc.add_paragraph()
            run = para.add_run(code_block.code)
            run.font.name = "Consolas"

        # Get available list styles
        bullet_style = get_list_style("List Bullet")
        number_style = get_list_style("List Number")

        for node in nodes:
            if isinstance(node, Paragraph):
                # Heading
                if node.level > 0:
                    doc.add_heading(node.text, level=min(node.level, 4))
                # Bullet list
                elif node.list_type == "bullet":
                    # Handle task lists
                    if node.task_checked is not None:
                        prefix = "☑ " if node.task_checked else "☐ "
                        para = doc.add_paragraph(style=bullet_style) if bullet_style else doc.add_paragraph()
                        if not bullet_style:
                            para.add_run("• ")
                        para.add_run(prefix)
                        add_runs_to_paragraph(para, node.runs)
                    else:
                        para = doc.add_paragraph(style=bullet_style) if bullet_style else doc.add_paragraph()
                        if not bullet_style:
                            para.add_run("• ")
                        add_runs_to_paragraph(para, node.runs)
                # Numbered list
                elif node.list_type == "ordered":
                    para = doc.add_paragraph(style=number_style) if number_style else doc.add_paragraph()
                    add_runs_to_paragraph(para, node.runs)
                # Regular paragraph
                else:
                    if node.runs:
                        para = doc.add_paragraph()
                        add_runs_to_paragraph(para, node.runs)

            elif isinstance(node, Table):
                add_table(node)

            elif isinstance(node, CodeBlock):
                add_code_block(node)

            elif isinstance(node, HorizontalRule):
                doc.add_paragraph()

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        safe_save_docx(doc, str(output))
        return {"success": True, "file": output_path, "message": "Created Word document from Markdown"}

    def _word_collect_comment_records(self, comments_root, comments_ids_bytes, comments_ex_root, etree_module) -> list[dict[str, Any]]:
        """Collect normalized comment records from OOXML parts."""
        from datetime import datetime

        explicit_ids: dict[str, str] = {}
        ordered_para_ids: list[str] = []
        if comments_ids_bytes:
            explicit_ids, ordered_para_ids = _parse_comments_ids_xml(comments_ids_bytes, etree_module)

        ext_by_para: dict[str, dict[str, Any]] = {}
        if comments_ex_root is not None:
            para_attr = f'{{{W15_NS}}}paraId'
            done_attr = f'{{{W15_NS}}}done'
            parent_attr = f'{{{W15_NS}}}paraIdParent'
            for node in comments_ex_root.findall(f'.//{{{W15_NS}}}commentEx'):
                para_id = node.get(para_attr)
                if not para_id:
                    continue
                ext_by_para[str(para_id)] = {
                    "done": _is_done_value(node.get(done_attr)),
                    "parent_para_id": node.get(parent_attr),
                }

        comment_nodes = comments_root.findall(f'.//{{{W_NS}}}comment')
        records: list[dict[str, Any]] = []

        para_attr = f'{{{W14_NS}}}paraId'
        parent_para_attr = f'{{{W14_NS}}}paraIdParent'

        for idx, comment_elem in enumerate(comment_nodes):
            comment_id = str(comment_elem.get(f'{{{W_NS}}}id') or "")
            author = comment_elem.get(f'{{{W_NS}}}author', 'Unknown')
            date_str = comment_elem.get(f'{{{W_NS}}}date', '')
            initials = comment_elem.get(f'{{{W_NS}}}initials', '')

            para_id = None
            para_parent = None
            first_para = comment_elem.find(f'{{{W_NS}}}p')
            if first_para is not None:
                para_id = first_para.get(para_attr)
                para_parent = first_para.get(parent_para_attr)

            if not para_id and comment_id in explicit_ids:
                para_id = explicit_ids[comment_id]
            elif not para_id and idx < len(ordered_para_ids):
                para_id = ordered_para_ids[idx]

            text = ''.join(comment_elem.itertext()).strip()

            date_formatted = date_str
            if date_str:
                try:
                    dt = datetime.fromisoformat(str(date_str).replace('Z', '+00:00'))
                    date_formatted = dt.strftime('%Y-%m-%d %H:%M')
                except (ValueError, TypeError):
                    pass

            records.append({
                "id": comment_id,
                "author": author,
                "initials": initials,
                "date": date_formatted,
                "text": text,
                "para_id": para_id,
                "_parent_para_id": para_parent,
            })

        para_to_comment_id = {
            str(record["para_id"]): record["id"]
            for record in records
            if record.get("para_id")
        }

        for record in records:
            para_id = record.get("para_id")
            ext = ext_by_para.get(str(para_id)) if para_id else None
            parent_para_id = record.get("_parent_para_id")
            if not parent_para_id and ext:
                parent_para_id = ext.get("parent_para_id")

            record["done"] = bool(ext.get("done")) if ext else False
            record["is_reply"] = bool(parent_para_id)
            record["parent_id"] = para_to_comment_id.get(str(parent_para_id)) if parent_para_id else None
            record.pop("_parent_para_id", None)

        return records

    def _word_build_threads(
        self,
        flat_comments: list[dict[str, Any]],
        all_comments: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """Build thread groups ({root, replies}) from flat comments."""
        all_by_id = {
            str(comment.get("id")): comment
            for comment in all_comments
            if comment.get("id") is not None
        }

        def root_comment_id(comment_id: str) -> str:
            current = str(comment_id)
            seen: set[str] = set()
            while current in all_by_id and current not in seen:
                seen.add(current)
                parent_id = all_by_id[current].get("parent_id")
                if not parent_id or str(parent_id) not in all_by_id:
                    break
                current = str(parent_id)
            return current

        threads: dict[str, dict[str, Any]] = {}
        ordered_roots: list[str] = []

        for comment in flat_comments:
            cid = str(comment.get("id"))
            rid = root_comment_id(cid)
            if rid not in threads:
                root_obj = all_by_id.get(rid, comment)
                threads[rid] = {
                    "root": dict(root_obj),
                    "replies": [],
                }
                ordered_roots.append(rid)

            if cid == rid:
                threads[rid]["root"] = dict(comment)
            else:
                threads[rid]["replies"].append(dict(comment))

        return [threads[root_id] for root_id in ordered_roots]

    def tool_word_get_comments(
        self,
        file_path: str,
        filter: Literal["open", "resolved", "mine", "all"] = "all",
        author: str | None = None,
        format: Literal["flat", "threaded"] = "flat",
    ) -> dict[str, Any]:
        """Extract comments from a Word document with metadata and thread context.

        Args:
            file_path: Path to the .docx file
            filter: Optional filter (all/open/resolved/mine)
            author: Author name for filter='mine' (defaults to comment identity)
            format: flat (default) or threaded

        Returns:
            Dictionary with comment list/count and optional thread groups
        """
        import zipfile

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"error": f"Expected .docx file, got: {path.suffix}"}

        try:
            from lxml import etree
        except ImportError:
            return {"error": "lxml not installed. Run: pip install lxml"}

        try:
            with zipfile.ZipFile(resolved_path, 'r') as zf:
                if 'word/comments.xml' not in zf.namelist():
                    empty = {
                        "file": path.name,
                        "comments": [],
                        "comment_count": 0,
                        "note": "No comments in document",
                    }
                    if format == "threaded":
                        empty["flat"] = []
                        empty["threads"] = []
                        empty["thread_count"] = 0
                    return empty

                comments_root = etree.fromstring(zf.read('word/comments.xml'))
                comments_ids_bytes = zf.read('word/commentsIds.xml') if 'word/commentsIds.xml' in zf.namelist() else None
                comments_ex_root = (
                    etree.fromstring(zf.read('word/commentsExtended.xml'))
                    if 'word/commentsExtended.xml' in zf.namelist()
                    else None
                )

                all_comments = self._word_collect_comment_records(
                    comments_root=comments_root,
                    comments_ids_bytes=comments_ids_bytes,
                    comments_ex_root=comments_ex_root,
                    etree_module=etree,
                )

            selected_comments = list(all_comments)
            filter_mode = str(filter or "all").lower()
            if filter_mode == "open":
                selected_comments = [comment for comment in selected_comments if not comment.get("done")]
            elif filter_mode == "resolved":
                selected_comments = [comment for comment in selected_comments if comment.get("done")]
            elif filter_mode == "mine":
                mine = (author or getattr(self, "_comment_author", DEFAULT_COMMENT_AUTHOR) or "").strip().lower()
                selected_comments = [
                    comment for comment in selected_comments
                    if str(comment.get("author", "")).strip().lower() == mine
                ]
            elif filter_mode != "all":
                return {"error": f"Unsupported filter: {filter}. Expected one of: all, open, resolved, mine"}

            result: dict[str, Any] = {
                "file": path.name,
                "comments": selected_comments,
                "comment_count": len(selected_comments),
                "filter": filter_mode,
            }

            if format == "threaded":
                threads = self._word_build_threads(selected_comments, all_comments)
                result["flat"] = selected_comments
                result["threads"] = threads
                result["thread_count"] = len(threads)
            elif format != "flat":
                return {"error": f"Unsupported format: {format}. Expected 'flat' or 'threaded'"}

            return result

        except Exception as e:
            return {"error": f"Failed to extract comments: {e}"}

    def tool_word_resolve_comment(
        self,
        file_path: str,
        comment_id: str,
        resolved: bool,
        output_path: str | None = None,
    ) -> dict[str, Any]:
        """Mark a Word comment thread as resolved or open.

        Resolution state is stored in word/commentsExtended.xml (w15:commentEx@w15:done).
        If a reply comment ID is supplied, the root thread comment is updated.

        Args:
            file_path: Path to the .docx file
            comment_id: Comment ID from word_get_comments
            resolved: True to resolve, False to reopen
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status dictionary with thread resolution details
        """
        import tempfile
        import zipfile
        from shutil import move

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"error": f"Expected .docx file, got: {path.suffix}"}

        target_comment_id = str(comment_id)

        try:
            from lxml import etree
        except ImportError:
            return {"error": "lxml not installed. Run: pip install lxml"}

        try:
            with zipfile.ZipFile(resolved_path, 'r') as zf_in:
                parts = {name: zf_in.read(name) for name in zf_in.namelist()}

            if 'word/comments.xml' not in parts:
                return {"error": "No comments.xml found in document"}

            comments_root = etree.fromstring(parts['word/comments.xml'])
            comments_ids_bytes = parts.get('word/commentsIds.xml')
            comments_ex_root = (
                etree.fromstring(parts['word/commentsExtended.xml'])
                if 'word/commentsExtended.xml' in parts
                else None
            )

            comments = self._word_collect_comment_records(
                comments_root=comments_root,
                comments_ids_bytes=comments_ids_bytes,
                comments_ex_root=comments_ex_root,
                etree_module=etree,
            )
            comments_by_id = {
                str(comment.get('id')): comment
                for comment in comments
                if comment.get('id') is not None
            }

            if target_comment_id not in comments_by_id:
                valid_ids = sorted(comments_by_id.keys(), key=lambda x: (0, int(x)) if x.isdigit() else (1, x))
                return {
                    "error": f"Comment {target_comment_id} not found. Valid IDs: {valid_ids}",
                    "valid_comment_ids": valid_ids,
                }

            # Resolve to root thread comment when target is a reply.
            root_comment_id = target_comment_id
            seen: set[str] = set()
            while root_comment_id in comments_by_id and root_comment_id not in seen:
                seen.add(root_comment_id)
                parent_id = comments_by_id[root_comment_id].get('parent_id')
                if not parent_id or str(parent_id) not in comments_by_id:
                    break
                root_comment_id = str(parent_id)

            root_comment = comments_by_id.get(root_comment_id)
            root_para_id = root_comment.get('para_id') if root_comment else None
            if not root_para_id:
                # Legacy docs may omit w14:paraId. Synthesize one on the root comment
                # so resolution metadata can be attached and future lookups are stable.
                import random

                para_id_attr = f'{{{W14_NS}}}paraId'
                existing_para_ids: set[str] = {
                    para.get(para_id_attr)
                    for para in comments_root.findall(f'.//{{{W_NS}}}p')
                    if para.get(para_id_attr)
                }

                while True:
                    candidate = f"{random.getrandbits(32):08X}"
                    if candidate not in existing_para_ids:
                        root_para_id = candidate
                        break

                comment_node = comments_root.find(
                    f".//{{{W_NS}}}comment[@w:id='{root_comment_id}']",
                    namespaces={"w": W_NS},
                )
                if comment_node is None:
                    return {
                        "error": f"Comment {root_comment_id} could not be located in comments.xml",
                        "comment_id": root_comment_id,
                    }

                first_para = comment_node.find(f'{{{W_NS}}}p')
                if first_para is None:
                    first_para = etree.SubElement(comment_node, f'{{{W_NS}}}p')
                    first_run = etree.SubElement(first_para, f'{{{W_NS}}}r')
                    first_text = etree.SubElement(first_run, f'{{{W_NS}}}t')
                    first_text.text = ""

                first_para.set(para_id_attr, str(root_para_id))
                parts['word/comments.xml'] = etree.tostring(
                    comments_root,
                    xml_declaration=True,
                    encoding='UTF-8',
                    standalone='yes',
                )

            created_comments_extended = False
            if comments_ex_root is None:
                comments_ex_root = etree.Element(f'{{{W15_NS}}}commentsEx', nsmap={'w15': W15_NS})
                created_comments_extended = True

            para_attr = f'{{{W15_NS}}}paraId'
            done_attr = f'{{{W15_NS}}}done'
            entry = None
            for node in comments_ex_root.findall(f'.//{{{W15_NS}}}commentEx'):
                if node.get(para_attr) == str(root_para_id):
                    entry = node
                    break

            if entry is None:
                entry = etree.SubElement(comments_ex_root, f'{{{W15_NS}}}commentEx')
                entry.set(para_attr, str(root_para_id))

            entry.set(done_attr, '1' if resolved else '0')

            parts['word/commentsExtended.xml'] = etree.tostring(
                comments_ex_root,
                xml_declaration=True,
                encoding='UTF-8',
                standalone='yes',
            )

            if created_comments_extended:
                _ensure_comments_extended_package_bits(parts, etree)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp_path = tmp.name

            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
                for name, blob in parts.items():
                    zf_out.writestr(name, blob)

            destination = output_path or resolved_path
            move(tmp_path, destination)

            return {
                "success": True,
                "file": destination,
                "comment_id": target_comment_id,
                "thread_root_comment_id": root_comment_id,
                "resolved": bool(resolved),
                "done": bool(resolved),
                "message": "Comment thread marked resolved" if resolved else "Comment thread reopened",
            }

        except Exception as e:
            return {"error": f"Failed to set comment resolution state: {e}"}

    def tool_word_reply_to_comment(
        self,
        file_path: str,
        comment_id: str,
        text: str,
        author: str | None = None,
        output_path: str | None = None,
        auto_resolve: bool = False,
    ) -> dict[str, Any]:
        """Add a threaded reply to an existing Word comment.

        Creates a new ``w:comment`` entry in ``word/comments.xml`` and links it
        to the parent thread using ``w14:paraIdParent`` on the reply's first
        paragraph. If the parent comment has no ``w14:paraId`` (older docs), a
        synthetic one is added and reused.

        Args:
            file_path: Path to the .docx file
            comment_id: ID of the parent comment (from word_get_comments)
            text: Reply text
            author: Reply author (defaults to office_set_comment_identity/env)
            output_path: Optional output path (defaults to overwriting input)
            auto_resolve: Mark the thread resolved after adding the reply

        Returns:
            Status dictionary with reply details
        """
        import random
        import tempfile
        import zipfile
        from datetime import datetime, timezone

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"error": f"Expected .docx file, got: {path.suffix}"}

        target_comment_id = str(comment_id)
        reply_text = str(text or "").strip()
        if not reply_text:
            return {"error": "text is required"}

        reply_author = (author or "").strip() or getattr(self, "_comment_author", DEFAULT_COMMENT_AUTHOR)
        reply_initials = (getattr(self, "_comment_initials", "") or "").strip().upper()
        if not reply_initials:
            tokens = [part for part in reply_author.split() if part]
            reply_initials = "".join(part[0].upper() for part in tokens[:2]) if tokens else "SA"

        try:
            from lxml import etree
        except ImportError:
            return {"error": "lxml not installed. Run: pip install lxml"}

        def next_para_id(existing_ids: set[str]) -> str:
            while True:
                candidate = f"{random.getrandbits(32):08X}"
                if candidate not in existing_ids:
                    existing_ids.add(candidate)
                    return candidate

        try:
            with zipfile.ZipFile(resolved_path, 'r') as zf_in:
                if 'word/comments.xml' not in zf_in.namelist():
                    return {"error": "No comments.xml found in document"}

                comments_root = etree.fromstring(zf_in.read('word/comments.xml'))
                comment_nodes = comments_root.findall(f'.//{{{W_NS}}}comment')

                valid_ids = []
                max_id = -1
                target_comment = None
                for node in comment_nodes:
                    node_id = node.get(f'{{{W_NS}}}id')
                    if node_id is not None:
                        valid_ids.append(node_id)
                        with_temp = node_id.strip()
                        if with_temp.isdigit():
                            max_id = max(max_id, int(with_temp))
                    if node_id == target_comment_id:
                        target_comment = node

                if target_comment is None:
                    valid_ids_sorted = sorted(valid_ids, key=lambda x: (0, int(x)) if x.isdigit() else (1, x))
                    return {
                        "error": f"Comment {target_comment_id} not found. Valid IDs: {valid_ids_sorted}",
                        "valid_comment_ids": valid_ids_sorted,
                    }

                # Collect existing paraIds across comments to keep them unique.
                para_id_attr = f'{{{W14_NS}}}paraId'
                para_parent_attr = f'{{{W14_NS}}}paraIdParent'
                existing_para_ids: set[str] = set()
                for para in comments_root.findall(f'.//{{{W_NS}}}p'):
                    para_id = para.get(para_id_attr)
                    if para_id:
                        existing_para_ids.add(para_id)

                # Parent paragraph and paraId (create one if missing).
                parent_para = target_comment.find(f'{{{W_NS}}}p')
                if parent_para is None:
                    parent_para = etree.SubElement(target_comment, f'{{{W_NS}}}p')
                    parent_run = etree.SubElement(parent_para, f'{{{W_NS}}}r')
                    parent_t = etree.SubElement(parent_run, f'{{{W_NS}}}t')
                    parent_t.text = ""

                parent_para_id = parent_para.get(para_id_attr)
                if not parent_para_id:
                    parent_para_id = next_para_id(existing_para_ids)
                    parent_para.set(para_id_attr, parent_para_id)

                new_comment_id = str(max_id + 1 if max_id >= 0 else 0)
                while new_comment_id in valid_ids:
                    max_id += 1
                    new_comment_id = str(max_id)

                reply_para_id = next_para_id(existing_para_ids)
                timestamp = datetime.now(timezone.utc).isoformat().replace('+00:00', 'Z')

                # Create reply comment.
                new_comment = etree.SubElement(comments_root, f'{{{W_NS}}}comment')
                new_comment.set(f'{{{W_NS}}}id', new_comment_id)
                new_comment.set(f'{{{W_NS}}}author', reply_author)
                new_comment.set(f'{{{W_NS}}}date', timestamp)
                new_comment.set(f'{{{W_NS}}}initials', reply_initials)

                reply_para = etree.SubElement(new_comment, f'{{{W_NS}}}p')
                reply_para.set(para_id_attr, reply_para_id)
                reply_para.set(para_parent_attr, parent_para_id)
                reply_run = etree.SubElement(reply_para, f'{{{W_NS}}}r')
                reply_t = etree.SubElement(reply_run, f'{{{W_NS}}}t')
                reply_t.text = reply_text

                comments_bytes = etree.tostring(
                    comments_root,
                    xml_declaration=True,
                    encoding='UTF-8',
                    standalone='yes',
                )

                comments_extended_bytes = None
                if 'word/commentsExtended.xml' in zf_in.namelist():
                    comments_ex_root = etree.fromstring(zf_in.read('word/commentsExtended.xml'))
                    para_attr = f'{{{W15_NS}}}paraId'
                    parent_attr = f'{{{W15_NS}}}paraIdParent'
                    done_attr = f'{{{W15_NS}}}done'

                    existing_comment_ex = {
                        node.get(para_attr)
                        for node in comments_ex_root.findall(f'.//{{{W15_NS}}}commentEx')
                        if node.get(para_attr)
                    }

                    if parent_para_id not in existing_comment_ex:
                        parent_ex = etree.SubElement(comments_ex_root, f'{{{W15_NS}}}commentEx')
                        parent_ex.set(para_attr, parent_para_id)
                        parent_ex.set(done_attr, '0')

                    reply_ex = etree.SubElement(comments_ex_root, f'{{{W15_NS}}}commentEx')
                    reply_ex.set(para_attr, reply_para_id)
                    reply_ex.set(parent_attr, parent_para_id)
                    reply_ex.set(done_attr, '0')

                    comments_extended_bytes = etree.tostring(
                        comments_ex_root,
                        xml_declaration=True,
                        encoding='UTF-8',
                        standalone='yes',
                    )

                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp_path = tmp.name

                with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
                    for item in zf_in.namelist():
                        if item == 'word/comments.xml':
                            zf_out.writestr(item, comments_bytes)
                        elif item == 'word/commentsExtended.xml' and comments_extended_bytes is not None:
                            zf_out.writestr(item, comments_extended_bytes)
                        else:
                            zf_out.writestr(item, zf_in.read(item))

            from shutil import move
            destination = output_path or resolved_path
            move(tmp_path, destination)

            response = {
                "success": True,
                "file": destination,
                "parent_comment_id": target_comment_id,
                "reply_comment_id": new_comment_id,
                "author": reply_author,
                "auto_resolve": bool(auto_resolve),
                "message": "Added threaded reply to comment",
            }

            if auto_resolve:
                resolve_result = self.tool_word_resolve_comment(
                    file_path=destination,
                    comment_id=target_comment_id,
                    resolved=True,
                    output_path=destination,
                )
                if resolve_result.get("error"):
                    return {
                        "error": (
                            "Added threaded reply but failed to auto-resolve thread: "
                            f"{resolve_result.get('error')}"
                        ),
                        "reply_comment_id": new_comment_id,
                        "parent_comment_id": target_comment_id,
                    }
                response["resolved"] = True
                response["thread_root_comment_id"] = resolve_result.get("thread_root_comment_id")
                response["message"] = "Added threaded reply and marked thread resolved"

            return response

        except Exception as e:
            return {"error": f"Failed to reply to comment: {e}"}

    def tool_word_reply_comment(
        self,
        file_path: str,
        comment_id: str,
        reply_text: str,
        author: str | None = None,
        output_path: str | None = None,
        auto_resolve: bool = False,
    ) -> dict[str, Any]:
        """Backward-compatible alias for word_reply_to_comment."""
        return self.tool_word_reply_to_comment(
            file_path=file_path,
            comment_id=comment_id,
            text=reply_text,
            author=author,
            output_path=output_path,
            auto_resolve=auto_resolve,
        )

    def tool_word_delete_comment(
        self,
        file_path: str,
        comment_id: str,
        output_path: str | None = None,
    ) -> dict[str, Any]:
        """Delete a comment from a Word document by comment ID.

        Removes the comment from comments.xml and strips associated reference
        markers from document.xml.

        Args:
            file_path: Path to the .docx file
            comment_id: Comment ID from word_get_comments output
            output_path: Optional output path (defaults to overwriting input)

        Returns:
            Status dictionary with deletion details
        """
        import tempfile
        import zipfile

        resolved_path = resolve_office_path(file_path)
        path = Path(resolved_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"error": f"Expected .docx file, got: {path.suffix}"}

        try:
            from lxml import etree
        except ImportError:
            return {"error": "lxml not installed. Run: pip install lxml"}

        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        comment_id = str(comment_id)
        try:
            with zipfile.ZipFile(resolved_path, 'r') as zf_in:
                if 'word/comments.xml' not in zf_in.namelist():
                    return {"error": "No comments.xml found in document"}

                comments_root = etree.fromstring(zf_in.read('word/comments.xml'))
                comment_nodes = comments_root.findall(f'.//{{{W_NS}}}comment')
                removed_nodes = [
                    node for node in comment_nodes
                    if node.get(f'{{{W_NS}}}id') == comment_id
                ]

                if not removed_nodes:
                    return {"error": f"Comment ID not found: {comment_id}"}

                for node in removed_nodes:
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)

                comments_bytes = etree.tostring(
                    comments_root,
                    xml_declaration=True,
                    encoding='UTF-8',
                    standalone='yes',
                )

                doc_root = etree.fromstring(zf_in.read('word/document.xml'))
                ref_tags = (
                    f'{{{W_NS}}}commentRangeStart',
                    f'{{{W_NS}}}commentRangeEnd',
                    f'{{{W_NS}}}commentReference',
                )
                removed_refs = 0
                for tag in ref_tags:
                    for node in list(doc_root.findall(f'.//{tag}')):
                        if node.get(f'{{{W_NS}}}id') == comment_id:
                            parent = node.getparent()
                            if parent is not None:
                                parent.remove(node)
                                removed_refs += 1

                doc_bytes = etree.tostring(
                    doc_root,
                    xml_declaration=True,
                    encoding='UTF-8',
                    standalone='yes',
                )

                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp_path = tmp.name

                with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
                    for item in zf_in.namelist():
                        if item == 'word/comments.xml':
                            zf_out.writestr(item, comments_bytes)
                        elif item == 'word/document.xml':
                            zf_out.writestr(item, doc_bytes)
                        else:
                            zf_out.writestr(item, zf_in.read(item))

            if output_path:
                from shutil import move
                move(tmp_path, output_path)
            else:
                from shutil import move
                move(tmp_path, resolved_path)

            return {
                "success": True,
                "file": output_path or resolved_path,
                "comment_id": comment_id,
                "removed_comment_nodes": len(removed_nodes),
                "removed_reference_nodes": removed_refs,
                "message": "Comment deleted",
            }
        except Exception as e:
            return {"error": f"Failed to delete comment: {e}"}
